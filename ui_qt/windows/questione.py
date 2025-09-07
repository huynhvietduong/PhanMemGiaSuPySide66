from __future__ import annotations
import json
import os
import base64
import io
from typing import List, Dict, Optional
from datetime import datetime

# Import PySide6
from PySide6 import QtCore, QtGui, QtWidgets
from PySide6.QtCore import Qt, QByteArray, QBuffer, QIODevice, Signal
from PySide6.QtWidgets import (
    QMenu, QFileDialog, QMessageBox, QDialog,
    QVBoxLayout, QHBoxLayout, QFormLayout,
    QWidget, QLabel, QLineEdit, QTextEdit,
    QPushButton, QComboBox, QSpinBox,
    QToolBar, QGroupBox, QSplitter,
    QStackedWidget, QScrollArea,
    QTableWidget, QTableWidgetItem,
    QTreeWidget, QTreeWidgetItem,
    QHeaderView, QAbstractItemView
, QSlider, QWidget, QApplication)

from PySide6.QtGui import (
    QKeySequence, QShortcut, QPixmap, QImage,
    QTextDocument, QTextCursor, QColor, QBrush,
    QAction, QIcon, QFont
)
from PySide6.QtPrintSupport import QPrintPreviewDialog, QPrinter


# #(Custom QTextBrowser để load ảnh từ database resources)
class CustomHTMLViewer(QtWidgets.QTextBrowser):
    """Custom QTextBrowser với khả năng load ảnh từ database"""

    def __init__(self, db_manager, question_id=None, parent=None):
        super().__init__(parent)
        self.db = db_manager
        self.question_id = question_id
        self.setOpenExternalLinks(False)

    def loadResource(self, type_hint, name):
        """Override để load resource từ database hoặc convert HTML"""
        # Nếu là QTextDocument.ImageResource
        if type_hint == QtGui.QTextDocument.ImageResource:
            try:
                resource_id = str(name.toString()) if hasattr(name, 'toString') else str(name)

                # Nếu là data URI, decode và return
                if resource_id.startswith('data:image/'):
                    return self._decode_data_uri(resource_id)

                # Nếu là số thuần túy (resource ID từ QTextEdit)
                if resource_id.isdigit():
                    # #(Tìm ảnh trong database dựa trên resource pattern)
                    image_from_db = self._load_image_from_database(resource_id)
                    if image_from_db:
                        return image_from_db

                    return self._create_placeholder_image(f"Resource ID: {resource_id}")

            except Exception as e:
                print(f"⚠️ Lỗi load resource: {e}")

        # Fallback về parent implementation
        return super().loadResource(type_hint, name)

    # #(Phương thức load ảnh từ database dựa trên resource ID hoặc pattern matching)

    # #(Helper method load pixmap từ nhiều định dạng data)
    def _try_load_pixmap_from_data(self, data):
        """Thử load pixmap từ data với nhiều format"""
        if not data:
            return None

        try:
            pixmap = QtGui.QPixmap()

            # Nếu là binary data
            if isinstance(data, (bytes, bytearray)):
                if pixmap.loadFromData(data):
                    return pixmap

            # Nếu là string (base64 hoặc data URI)
            elif isinstance(data, str):
                import base64

                # Data URI format
                if data.startswith('data:image'):
                    try:
                        header, b64_data = data.split(',', 1)
                        decoded_data = base64.b64decode(b64_data)
                        if pixmap.loadFromData(decoded_data):
                            return pixmap
                    except:
                        pass

                # Pure base64
                else:
                    try:
                        decoded_data = base64.b64decode(data)
                        if pixmap.loadFromData(decoded_data):
                            return pixmap
                    except:
                        pass

            return None

        except Exception as e:
            print(f"⚠️ Lỗi load pixmap từ data: {e}")
            return None

    # THAY THẾ phương thức _create_placeholder_image():
    def _create_placeholder_image(self, info_text="Ảnh không tìm thấy"):
        """Tạo placeholder image với thông tin chi tiết"""
        pixmap = QtGui.QPixmap(300, 150)
        pixmap.fill(QtGui.QColor("#f8f9fa"))

        painter = QtGui.QPainter(pixmap)
        painter.setPen(QtGui.QColor("#666"))
        painter.setFont(QtGui.QFont("Arial", 10))

        # Vẽ border
        painter.drawRect(pixmap.rect().adjusted(2, 2, -2, -2))

        # Vẽ icon ảnh
        painter.setFont(QtGui.QFont("Arial", 24))
        painter.drawText(pixmap.rect(), Qt.AlignCenter | Qt.AlignTop, "🖼️")

        # Vẽ text thông tin
        painter.setFont(QtGui.QFont("Arial", 9))
        text_rect = pixmap.rect().adjusted(10, 50, -10, -10)

        full_text = f"Đáp án chứa ảnh\n{info_text}\nẢnh được nhúng từ text editor"
        painter.drawText(text_rect, Qt.AlignCenter | Qt.TextWordWrap, full_text)

        painter.end()
        return pixmap

    def _decode_data_uri(self, data_uri):
        """Decode data URI thành QPixmap"""
        try:
            import base64
            if ',base64,' in data_uri:
                header, data = data_uri.split(',base64,', 1)
                image_data = base64.b64decode(data)

                pixmap = QtGui.QPixmap()
                if pixmap.loadFromData(image_data):
                    return pixmap
        except Exception as e:
            print(f"⚠️ Lỗi decode data URI: {e}")

        return self._create_placeholder_image()


class QuestionBankWindowQt(QtWidgets.QWidget):
    # ========== NHÓM 1: KHỞI TẠO VÀ THIẾT LẬP ========== #
    def __init__(self, db_manager, parent=None):
        super().__init__(parent)
        self.db = db_manager
        self.setObjectName("QuestionBankWindowQt")
        self.setWindowTitle("Ngân hàng câu hỏi")
        self.showMaximized()

        # Đảm bảo bảng tồn tại
        self._ensure_tables()

        self.current_question_id: int | None = None
        self.tree_nodes: Dict[str, int] = {}

        root = QtWidgets.QVBoxLayout(self)

        # Tạo toolbar chính
        main_toolbar = QtWidgets.QToolBar()
        main_toolbar.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        main_toolbar.setMovable(False)
        main_toolbar.setStyleSheet("""
            QToolBar { 
                background: #f8f9fa; 
                border: 1px solid #dee2e6; 
                spacing: 8px; 
                padding: 4px;
            }
            QToolBar::separator { 
                background: #dee2e6; 
                width: 2px; 
                margin: 0 4px;
            }
        """)
        root.addWidget(main_toolbar)

        # Nhóm quản lý cây
        toggle_action = main_toolbar.addAction("🌲 Ẩn/Hiện cây")
        toggle_action.triggered.connect(self.toggle_tree_panel)

        manage_action = main_toolbar.addAction("⚙️ Quản lý cây")
        manage_action.triggered.connect(self.open_tree_manager)

        main_toolbar.addSeparator()

        # Nhóm tìm kiếm
        search_widget = QtWidgets.QWidget()
        search_layout = QtWidgets.QHBoxLayout(search_widget)
        search_layout.setContentsMargins(0, 0, 0, 0)

        search_layout.addWidget(QtWidgets.QLabel("🔍"))

        self.search_edit = QtWidgets.QLineEdit()
        self.search_edit.setPlaceholderText("Tìm kiếm câu hỏi...")
        self.search_edit.setMinimumWidth(200)
        self.search_edit.setStyleSheet("padding: 4px; border: 1px solid #ced4da; border-radius: 4px;")
        search_layout.addWidget(self.search_edit)

        main_toolbar.addWidget(search_widget)

        search_action = main_toolbar.addAction("Tìm")
        search_action.triggered.connect(self.search_questions)

        main_toolbar.addSeparator()

        # Thêm nút tạo câu hỏi nổi bật
        add_question_action = main_toolbar.addAction("➕ Thêm câu hỏi")
        add_question_action.triggered.connect(self.open_add_question_dialog)
        # Style cho nút
        add_btn_widget = main_toolbar.widgetForAction(add_question_action)
        if add_btn_widget:
            add_btn_widget.setStyleSheet("""
                QToolButton {
                    background: #28a745;
                    color: white;
                    font-weight: bold;
                    padding: 6px 12px;
                    border-radius: 4px;
                }
                QToolButton:hover {
                    background: #218838;
                }
            """)
        # Nhóm tạo mới
        new_action = main_toolbar.addAction("➕ Tạo mới")
        new_action.triggered.connect(self.new_question)

        template_action = main_toolbar.addAction("📋 Template")
        template_action.triggered.connect(self.show_template_dialog)

        main_toolbar.addSeparator()

        # Nhóm import/export
        import_action = main_toolbar.addAction("📥 Import Word")
        import_action.triggered.connect(self.import_from_word)

        export_action = main_toolbar.addAction("📤 Export Word")
        export_action.triggered.connect(self.export_to_word)

        export_pdf_action = main_toolbar.addAction("📄 Export PDF")
        export_pdf_action.triggered.connect(self.export_to_pdf)

        main_toolbar.addSeparator()

        # Toolbar phụ cho filters
        filter_toolbar = QtWidgets.QToolBar()
        filter_toolbar.setStyleSheet("QToolBar { background: #e9ecef; border: 1px solid #dee2e6; }")
        root.addWidget(filter_toolbar)

        self._create_filter_controls(filter_toolbar)

        # Splitter 3 cột
        split = QtWidgets.QSplitter(Qt.Horizontal)
        root.addWidget(split, 1)

        # --- Cột trái: Cây ---
        left = QtWidgets.QWidget()
        left_l = QtWidgets.QVBoxLayout(left)
        left_l.setContentsMargins(6, 6, 6, 6)

        left_l.addWidget(QtWidgets.QLabel("Cây thư mục"))
        self.tree = QtWidgets.QTreeWidget()
        self.tree.setHeaderHidden(True)
        self.tree.itemSelectionChanged.connect(self.on_tree_select)
        left_l.addWidget(self.tree, 1)
        split.addWidget(left)

        # --- Cột giữa: Danh sách câu hỏi với header controls ---
        mid = QtWidgets.QWidget()
        mid_l = QtWidgets.QVBoxLayout(mid)
        mid_l.setContentsMargins(6, 6, 6, 6)

        # Header controls với thống kê
        header_widget = QtWidgets.QWidget()
        header_layout = QtWidgets.QVBoxLayout(header_widget)
        header_layout.setSpacing(8)

        title_row = QtWidgets.QHBoxLayout()
        title_label = QtWidgets.QLabel("📋 Danh sách câu hỏi")
        title_label.setStyleSheet("font-weight: bold; font-size: 14px;")
        self.stats_label = QtWidgets.QLabel("0 câu hỏi")
        self.stats_label.setStyleSheet("color: #666; font-size: 12px;")
        title_row.addWidget(title_label)
        title_row.addStretch()
        title_row.addWidget(self.stats_label)
        header_layout.addLayout(title_row)

        # Dòng điều khiển nhanh
        controls_row = QtWidgets.QHBoxLayout()

        # Tìm kiếm nhanh
        self.quick_search = QtWidgets.QLineEdit()
        self.quick_search.setPlaceholderText("🔍 Tìm theo nội dung, tags...")
        self.quick_search.setMaximumWidth(250)
        self.quick_search.textChanged.connect(self.on_quick_search)
        controls_row.addWidget(self.quick_search)

        # Bộ lọc loại content
        self.content_type_filter = QtWidgets.QComboBox()
        self.content_type_filter.addItems(["Tất cả loại", "📝 Text", "🖼️ Image", "📄 PDF", "📘 Word"])
        self.content_type_filter.currentTextChanged.connect(self.apply_filters)
        controls_row.addWidget(self.content_type_filter)

        # Bộ lọc độ khó
        self.difficulty_filter = QtWidgets.QComboBox()
        self.difficulty_filter.addItems(["Tất cả độ khó", "🟢 Dễ", "🟡 Trung bình", "🔴 Khó"])
        self.difficulty_filter.currentTextChanged.connect(self.apply_filters)
        controls_row.addWidget(self.difficulty_filter)

        # Nút tùy chọn hiển thị
        self.view_options_btn = QtWidgets.QPushButton("⚙️")
        self.view_options_btn.setToolTip("Tùy chọn hiển thị")
        self.view_options_btn.setMaximumWidth(30)
        self.view_options_btn.clicked.connect(self.show_view_options)
        controls_row.addWidget(self.view_options_btn)

        controls_row.addStretch()
        header_layout.addLayout(controls_row)
        mid_l.addWidget(header_widget)

        # Bảng câu hỏi tối ưu (9 cột thay vì 8)
        # Bảng câu hỏi với 10 cột
        self.q_table = QtWidgets.QTableWidget(0, 9)
        headers = ["☑️", "ID", "📊 Loại", "🎯 Độ khó", "✅ Đáp án", "📁 Chủ đề", "🏷️ Tags", "📊 Sử dụng",
                   "📅 Ngày tạo"]
        self.q_table.setHorizontalHeaderLabels(headers)

        # Cấu hình resize mode tối ưu
        header = self.q_table.horizontalHeader()
        header.setSectionResizeMode(0, QtWidgets.QHeaderView.Fixed)  # Checkbox
        header.setSectionResizeMode(1, QtWidgets.QHeaderView.Fixed)  # ID
        header.setSectionResizeMode(2, QtWidgets.QHeaderView.Fixed)  # Loại
        header.setSectionResizeMode(3, QtWidgets.QHeaderView.Fixed)  # Độ khó
        header.setSectionResizeMode(4, QtWidgets.QHeaderView.ResizeToContents)  # Đáp án
        header.setSectionResizeMode(5, QtWidgets.QHeaderView.ResizeToContents)  # Chủ đề
        header.setSectionResizeMode(6, QtWidgets.QHeaderView.ResizeToContents)  # Tags
        header.setSectionResizeMode(7, QtWidgets.QHeaderView.Fixed)  # Ngày tạo

        # Đặt chiều rộng cố định cho các cột
        self.q_table.setColumnWidth(0, 40)  # Checkbox
        self.q_table.setColumnWidth(1, 40)  # ID
        self.q_table.setColumnWidth(3, 40)  # Loại
        self.q_table.setColumnWidth(4, 40)  # Độ khó
        self.q_table.setColumnWidth(8, 50)  # Ngày tạo

        self.q_table.setSortingEnabled(True)
        self.q_table.setContextMenuPolicy(Qt.CustomContextMenu)
        self.q_table.customContextMenuRequested.connect(self.show_enhanced_context_menu)
        self.q_table.setAlternatingRowColors(True)
        self.q_table.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        self.q_table.setSelectionMode(QtWidgets.QAbstractItemView.ExtendedSelection)
        self.q_table.itemSelectionChanged.connect(self.on_question_select_enhanced)

        # Style cải tiến cho bảng
        self.q_table.setStyleSheet("""
            QTableWidget {
                border: 1px solid #ddd;
                border-radius: 8px;
                background-color: white;
                gridline-color: #f0f0f0;
                selection-background-color: #e3f2fd;
            }
            QTableWidget::item {
                padding: 8px;
                border-bottom: 1px solid #f0f0f0;
            }
            QTableWidget::item:selected {
                background-color: #e3f2fd;
                color: #1976d2;
            }
            QHeaderView::section {
                background-color: #f8f9fa;
                padding: 10px;
                border: none;
                border-right: 1px solid #dee2e6;
                font-weight: bold;
                color: #495057;
            }
        """)

        mid_l.addWidget(self.q_table, 1)
        split.addWidget(mid)
        # ========== TẠO PREVIEW PANEL ĐƠN GIẢN - KHÔNG CONFLICT ========== #
        # Tạo widget phải đơn giản
        right_widget = QtWidgets.QWidget()
        right_layout = QtWidgets.QVBoxLayout(right_widget)
        right_layout.setContentsMargins(6, 6, 6, 6)

        # Preview Group đơn giản
        self.preview_group = QtWidgets.QGroupBox("📋 PREVIEW CÂU HỎI")
        self.preview_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #dee2e6;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
            }
        """)
        preview_layout = QtWidgets.QVBoxLayout(self.preview_group)

        # Thông tin cơ bản
        self.info_widget = QtWidgets.QWidget()
        info_layout = QtWidgets.QFormLayout(self.info_widget)
        info_layout.setSpacing(5)

        self.preview_id_label = QtWidgets.QLabel("-")
        self.preview_topic_label = QtWidgets.QLabel("-")
        self.preview_difficulty_label = QtWidgets.QLabel("-")
        self.preview_date_label = QtWidgets.QLabel("-")

        info_layout.addRow("🆔 ID:", self.preview_id_label)
        info_layout.addRow("📁 Chủ đề:", self.preview_topic_label)
        info_layout.addRow("🎯 Độ khó:", self.preview_difficulty_label)
        info_layout.addRow("📅 Ngày tạo:", self.preview_date_label)

        preview_layout.addWidget(self.info_widget)

        # Separator
        separator = QtWidgets.QFrame()
        separator.setFrameShape(QtWidgets.QFrame.HLine)
        separator.setFrameShadow(QtWidgets.QFrame.Sunken)
        preview_layout.addWidget(separator)
        # Tạo preview container adaptive
        self.preview_container = QtWidgets.QStackedWidget()

        # Text preview (giữ nguyên)
        self.preview_text = QtWidgets.QLabel("Chọn câu hỏi để xem preview...")
        self.preview_text.setWordWrap(True)
        self.preview_text.setStyleSheet("padding: 10px; background: #f8f9fa; border-radius: 4px; min-height: 300px;")
        self.preview_text.setAlignment(Qt.AlignTop | Qt.AlignLeft)
        self.preview_container.addWidget(self.preview_text)

        # Image preview adaptive
        self.preview_image = AdaptiveImageViewer()
        self.preview_image.set_size_limits(600, 500, 250)  # Nhỏ hơn cho preview
        self.preview_image.enable_zoom_controls()
        self.setup_preview_interactions()
        self.preview_image.setToolTip("Double-click để xem ảnh fullscreen\nDùng nút +/- để zoom")
        image_scroll_area = QtWidgets.QScrollArea()
        image_scroll_area.setWidget(self.preview_image)
        image_scroll_area.setWidgetResizable(True)
        image_scroll_area.setMaximumHeight(350)
        image_scroll_area.setStyleSheet("QScrollArea { border: none; }")
        self.preview_container.addWidget(self.preview_image)

        preview_layout.addWidget(self.preview_container)
        right_layout.addWidget(self.preview_group)
        split.addWidget(right_widget)

        # Khởi tạo biến đơn giản
        self.current_zoom = 1.0
        self.original_pixmap = None

        # Thiết lập tỷ lệ splitter: Tree(20%) - Questions(50%) - Preview(30%)
        split.setSizes([200, 200, 1000])

        self.q_table.itemSelectionChanged.connect(self.on_question_select)
        self.q_table.cellDoubleClicked.connect(self.on_cell_double_clicked)
        # Init dữ liệu
        self.refresh_tree()
        self.load_available_subjects()
        self.load_available_grades()

        # Signal cho combobox
        self.subject_cb.currentIndexChanged.connect(self.load_available_topics)
        self.grade_cb.currentIndexChanged.connect(self.load_available_topics)
        self.topic_cb.currentIndexChanged.connect(self.load_available_types)

        # Keyboard shortcuts
        QShortcut("Ctrl+N", self, self.new_question)
        QShortcut("Ctrl+S", self, self.save_question)
        QShortcut("Ctrl+F", self, self.focus_search)
        QShortcut("Delete", self, self.delete_question)
        QShortcut("F5", self, self.refresh_all)

        self.setAcceptDrops(True)
        self.q_table.setDragDropMode(QtWidgets.QAbstractItemView.InternalMove)

        self._setup_tree_management()

        self.apply_default_compact_mode()

    # ========== HELPER FUNCTIONS CHO DATABASE ROW ========== #
    def apply_default_compact_mode(self):
        """Áp dụng chế độ thu gọn làm mặc định khi khởi động"""

        def set_compact():
            try:
                self.toggle_compact_mode()
                print("✅ Đã áp dụng chế độ thu gọn làm mặc định")
            except Exception as e:
                print(f"⚠️ Lỗi áp dụng chế độ thu gọn: {e}")

        # Delay 200ms để đảm bảo UI đã render xong
        QtCore.QTimer.singleShot(200, set_compact)

    def safe_get(self, row, column, default=None):
        """Truy cập an toàn dữ liệu từ sqlite3.Row hoặc dict - SỬA LỖI"""
        if row is None:
            return default

        try:
            # Nếu là dict, dùng .get()
            if isinstance(row, dict):
                return row.get(column, default)

            # Nếu là sqlite3.Row, truy cập trực tiếp
            if hasattr(row, 'keys') and column in row.keys():
                value = row[column]
                return value if value is not None else default

            # Thử truy cập như attribute
            if hasattr(row, column):
                value = getattr(row, column)
                return value if value is not None else default

            return default

        except (KeyError, IndexError, TypeError, AttributeError) as e:
            print(f"⚠️ safe_get error for column '{column}': {e}")
            return default

    # #(Helper method load pixmap từ answer_data binary)
    def load_pixmap_from_answer_data(self, answer_data):
        """Load QPixmap từ answer_data (binary hoặc base64)"""
        if not answer_data:
            return None

        try:
            pixmap = QtGui.QPixmap()

            # Nếu là bytes/binary data
            if isinstance(answer_data, (bytes, bytearray)):
                if pixmap.loadFromData(answer_data):
                    return pixmap

            # Nếu là string (có thể là base64)
            elif isinstance(answer_data, str):
                import base64
                try:
                    if answer_data.startswith('data:image'):
                        # Data URL format
                        header, data = answer_data.split(',', 1)
                        decoded_data = base64.b64decode(data)
                    else:
                        # Pure base64
                        decoded_data = base64.b64decode(answer_data)

                    if pixmap.loadFromData(decoded_data):
                        return pixmap
                except:
                    pass

            return None

        except Exception as e:
            print(f"⚠️ Lỗi load pixmap: {e}")
            return None

    # #(Helper method kiểm tra HTML có chứa ảnh không)
    def _has_images_in_html(self, html_content):
        """Kiểm tra HTML có chứa ảnh không"""
        if not html_content:
            return False

        # Kiểm tra các tag ảnh HTML
        html_lower = html_content.lower()

        # Kiểm tra tag <img>
        if '<img' in html_lower:
            return True

        # Kiểm tra data URI cho ảnh
        if 'data:image/' in html_lower:
            return True

        # Kiểm tra các định dạng ảnh base64
        if 'base64' in html_lower and ('png' in html_lower or 'jpg' in html_lower or 'jpeg' in html_lower):
            return True

        return False

    # #(Phương thức tiền xử lý HTML để convert resource IDs thành data URIs)
    def _process_html_for_display(self, html_content, question_id):
        """Tiền xử lý HTML để convert resource IDs thành displayable format"""
        if not html_content:
            return html_content

        try:
            import re

            # Tìm và thay thế tất cả thẻ img với src là số
            def replace_img_tag(match):
                resource_id = match.group(1)
                placeholder_html = (
                    '<div style="border: 2px dashed #ccc; padding: 20px; text-align: center; '
                    'margin: 10px; background: #f9f9f9;">'
                    '<p style="color: #666; margin: 0;"><strong>🖼️ Hình ảnh đáp án</strong></p>'
                    f'<p style="color: #999; font-size: 12px; margin: 5px 0 0 0;">Resource ID: {resource_id}</p>'
                    '<p style="color: #777; font-size: 11px; margin: 5px 0 0 0;">'
                    '<em>Ảnh được nhúng từ text editor</em></p>'
                    '</div>'
                )
                return placeholder_html

            # Pattern tìm thẻ img với src là số
            img_pattern = r'<img[^>]*src=["\'](\d+)["\'][^>]*>'
            processed_html = re.sub(img_pattern, replace_img_tag, html_content, flags=re.IGNORECASE)

            return processed_html

        except Exception as e:
            print(f"⚠️ Lỗi process HTML: {e}")
            return html_content

    def row_to_dict(self, row):
        """Chuyển đổi sqlite3.Row thành dictionary an toàn - SỬA LỖI RECURSION"""
        if row is None:
            return {}

        try:
            # Kiểm tra nếu đã là dict thì trả về luôn
            if isinstance(row, dict):
                return row

            # Nếu là sqlite3.Row, chuyển đổi sang dict
            if hasattr(row, 'keys'):
                return {key: row[key] for key in row.keys()}

            # Nếu là tuple hoặc list, thử convert
            if isinstance(row, (tuple, list)) and hasattr(row, '_fields'):
                return dict(zip(row._fields, row))

            # Fallback: thử dict() constructor
            return dict(row)

        except Exception as e:
            print(f"⚠️ Lỗi row_to_dict: {e}, type: {type(row)}")
            return {}

    def open_add_question_dialog(self):
        """Mở dialog thêm câu hỏi mới"""
        tree_id = self._current_tree_id()
        if not tree_id:
            QtWidgets.QMessageBox.warning(self, "Chưa chọn thư mục",
                                          "Vui lòng chọn thư mục trong cây để lưu câu hỏi")
            return

        dialog = QuestionEditDialog(self.db, tree_id=tree_id, parent=self)
        if dialog.exec() == QtWidgets.QDialog.Accepted:
            # Refresh danh sách câu hỏi
            self.on_tree_select()

    def _ensure_tables(self):
        """Đảm bảo các bảng tồn tại với schema mới"""
        self.db.upgrade_question_bank_schema()
        print("✅ Đã đảm bảo schema ngân hàng câu hỏi")

    def _insert_sample_tree_data(self):
        """Thêm dữ liệu mẫu cho cây thư mục"""
        sample_data = [
            (None, "Toán", "Môn", "Môn Toán học"),
            (None, "Lý", "Môn", "Môn Vật lý"),
            (None, "Hóa", "Môn", "Môn Hóa học"),
            (1, "Lớp 10", "Lớp", "Toán lớp 10"),
            (1, "Lớp 11", "Lớp", "Toán lớp 11"),
            (1, "Lớp 12", "Lớp", "Toán lớp 12"),
            (4, "Mệnh đề - Tập hợp", "Chủ đề", "Chương 1"),
            (4, "Hàm số", "Chủ đề", "Chương 2"),
            (7, "Nhận biết", "Mức độ", "Câu hỏi nhận biết cơ bản"),
            (7, "Thông hiểu", "Mức độ", "Câu hỏi thông hiểu"),
            (7, "Vận dụng", "Mức độ", "Câu hỏi vận dụng"),
        ]

        for parent_id, name, level, description in sample_data:
            self.db.execute_query(
                "INSERT INTO exercise_tree (parent_id, name, level, description) VALUES (?, ?, ?, ?)",
                (parent_id, name, level, description)
            )

    # ========== NHÓM 2: TẠO GIAO DIỆN ========== #
    def _create_filter_controls(self, toolbar):
        """Tạo các combobox filter"""
        toolbar.addWidget(QtWidgets.QLabel("Môn:"))
        self.subject_cb = QtWidgets.QComboBox()
        self.subject_cb.setMinimumWidth(120)
        toolbar.addWidget(self.subject_cb)

        toolbar.addWidget(QtWidgets.QLabel("Lớp:"))
        self.grade_cb = QtWidgets.QComboBox()
        self.grade_cb.setMinimumWidth(100)
        toolbar.addWidget(self.grade_cb)

        toolbar.addWidget(QtWidgets.QLabel("Chủ đề:"))
        self.topic_cb = QtWidgets.QComboBox()
        self.topic_cb.setMinimumWidth(150)
        toolbar.addWidget(self.topic_cb)

        toolbar.addWidget(QtWidgets.QLabel("Dạng:"))
        self.type_cb = QtWidgets.QComboBox()
        self.type_cb.setMinimumWidth(120)
        toolbar.addWidget(self.type_cb)

        toolbar.addWidget(QtWidgets.QLabel("Mức độ:"))
        self.level_cb = QtWidgets.QComboBox()
        self.level_cb.addItems(["", "Nhận biết", "Thông hiểu", "Vận dụng", "Vận dụng cao"])
        self.level_cb.setMinimumWidth(120)
        toolbar.addWidget(self.level_cb)

        toolbar.addSeparator()

        filter_btn = toolbar.addAction("🔽 Lọc")
        filter_btn.triggered.connect(self.filter_by_combobox)

        clear_filter_btn = toolbar.addAction("🔄 Xóa lọc")
        clear_filter_btn.triggered.connect(self.clear_filters)

    def _create_preview_tab_content(self, layout):
        """Tạo nội dung tab preview"""
        layout.addWidget(QtWidgets.QLabel("📋 Xem trước câu hỏi:"))

        self.preview_widget = QtWidgets.QTextEdit()
        self.preview_widget.setReadOnly(True)
        self.preview_widget.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 2px solid #e9ecef;
                border-radius: 8px;
                padding: 16px;
                font-size: 14px;
            }
        """)
        layout.addWidget(self.preview_widget)

    def _create_stats_tab_content(self, layout):
        """Tạo nội dung tab thống kê"""
        layout.addWidget(QtWidgets.QLabel("📊 Thống kê ngân hàng câu hỏi:"))

        self.stats_widget = QtWidgets.QTextEdit()
        self.stats_widget.setReadOnly(True)
        layout.addWidget(self.stats_widget)

        update_stats_btn = QtWidgets.QPushButton("🔄 Cập nhật thống kê")
        update_stats_btn.clicked.connect(self.update_statistics)
        layout.addWidget(update_stats_btn)

    def _create_history_tab_content(self, layout):
        """Tạo nội dung tab lịch sử"""
        layout.addWidget(QtWidgets.QLabel("📜 Lịch sử chỉnh sửa:"))

        self.history_table = QtWidgets.QTableWidget(0, 4)
        self.history_table.setHorizontalHeaderLabels(["Thời gian", "Hành động", "Nội dung cũ", "Nội dung mới"])
        header = self.history_table.horizontalHeader()
        header.setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeToContents)
        header.setSectionResizeMode(1, QtWidgets.QHeaderView.ResizeToContents)
        header.setSectionResizeMode(2, QtWidgets.QHeaderView.Stretch)
        header.setSectionResizeMode(3, QtWidgets.QHeaderView.Stretch)

        layout.addWidget(self.history_table)

    # ========== NHÓM 3: QUẢN LÝ CÂY THƯ MỤC ========== #
    def refresh_tree(self):
        """Làm mới cây thư mục"""
        try:
            self.tree.clear()
            self.tree_nodes.clear()

            rows = self.db.execute_query(
                "SELECT id, parent_id, name, level FROM exercise_tree ORDER BY parent_id, level, name",
                fetch='all'
            ) or []

            if not rows:
                self._insert_sample_tree_data()
                rows = self.db.execute_query(
                    "SELECT id, parent_id, name, level FROM exercise_tree ORDER BY parent_id, level, name",
                    fetch='all'
                ) or []

            children: Dict[int | None, list] = {}
            for r in rows:
                children.setdefault(r["parent_id"], []).append(r)

            def build(parent_db_id: int | None, parent_item: QtWidgets.QTreeWidgetItem | None):
                for node in children.get(parent_db_id, []):
                    icon_text = self._get_level_icon(node["level"])
                    item_text = f"{icon_text} {node['name']}"

                    item = QtWidgets.QTreeWidgetItem([item_text])
                    item.setData(0, Qt.UserRole, node["id"])
                    item.setToolTip(0, f"Level: {node['level']}\nID: {node['id']}")

                    self.tree_nodes[str(id(item))] = node["id"]

                    if parent_item is None:
                        self.tree.addTopLevelItem(item)
                    else:
                        parent_item.addChild(item)
                    build(node["id"], item)

            build(None, None)
            self.tree.expandAll()

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể tải cây thư mục: {e}")

    def _get_level_icon(self, level: str) -> str:
        """Trả về icon cho từng level"""
        icons = {
            "Môn": "📚",
            "Lớp": "🎓",
            "Chủ đề": "📖",
            "Dạng": "📝",
            "Mức độ": "⭐"
        }
        return icons.get(level, "📁")

    def on_tree_select(self):
        """Xử lý khi chọn node trên cây"""
        items = self.tree.selectedItems()
        if not items:
            return
        tree_id = items[0].data(0, Qt.UserRole)
        if not tree_id:
            return

        rows = self.db.execute_query("SELECT * FROM question_bank WHERE tree_id=?", (tree_id,), fetch="all") or []
        self._load_question_rows(rows)

    def toggle_tree_panel(self):
        """Ẩn/hiện panel cây"""
        w = self.tree.parentWidget()
        w.setVisible(not w.isVisible())

    def open_tree_manager(self):
        """Mở cửa sổ quản lý cây"""
        QtWidgets.QMessageBox.information(self, "Thông tin", "Chức năng quản lý cây đang phát triển.")

    def _setup_tree_management(self):
        """Thiết lập chức năng quản lý cây thư mục"""
        # Thêm context menu cho tree
        self.tree.setContextMenuPolicy(Qt.CustomContextMenu)
        self.tree.customContextMenuRequested.connect(self._show_tree_context_menu)

        # Thêm double-click để edit
        self.tree.itemDoubleClicked.connect(self._edit_tree_node)

        # Thêm keyboard shortcuts
        self._setup_tree_shortcuts()

    def _setup_tree_shortcuts(self):
        """Thiết lập keyboard shortcuts cho tree"""
        # F2 để edit node được chọn
        edit_shortcut = QShortcut(QKeySequence("F2"), self.tree)
        edit_shortcut.activated.connect(self._edit_selected_tree_node)

        # Delete để xóa node
        delete_shortcut = QShortcut(QKeySequence("Delete"), self.tree)
        delete_shortcut.activated.connect(self._delete_selected_tree_node)

        # Ctrl+N để thêm node mới
        add_shortcut = QShortcut(QKeySequence("Ctrl+Shift+N"), self.tree)
        add_shortcut.activated.connect(self._add_tree_node)

    def _show_tree_context_menu(self, position):
        """Hiển thị context menu cho tree"""
        item = self.tree.itemAt(position)

        menu = QtWidgets.QMenu(self)

        # Thêm node mới
        add_action = menu.addAction("➕ Thêm nhánh mới")
        add_action.triggered.connect(lambda: self._add_tree_node(item))

        if item:  # Nếu click vào node
            menu.addSeparator()

            # Thêm node con
            add_child_action = menu.addAction("📁 Thêm nhánh con")
            add_child_action.triggered.connect(lambda: self._add_child_node(item))

            # Sửa node
            edit_action = menu.addAction("✏️ Sửa tên nhánh")
            edit_action.triggered.connect(lambda: self._edit_tree_node(item))

            # Sao chép node
            copy_action = menu.addAction("📋 Sao chép nhánh")
            copy_action.triggered.connect(lambda: self._copy_tree_node(item))

            menu.addSeparator()

            # Xóa node
            delete_action = menu.addAction("🗑️ Xóa nhánh")
            delete_action.triggered.connect(lambda: self._delete_tree_node(item))

            menu.addSeparator()

            # Thông tin node
            info_action = menu.addAction("ℹ️ Thông tin")
            info_action.triggered.connect(lambda: self._show_node_info(item))

        # Làm mới cây
        menu.addSeparator()
        refresh_action = menu.addAction("🔄 Làm mới")
        refresh_action.triggered.connect(self.refresh_tree)

        # Hiển thị menu
        menu.exec(self.tree.mapToGlobal(position))

    def _add_tree_node(self, parent_item=None):
        """Thêm node mới"""
        try:
            dialog = TreeNodeDialog(self.db, mode="add", parent=self)

            # Nếu có parent item, set làm parent
            parent_id = None
            if parent_item:
                parent_id = parent_item.data(0, Qt.UserRole)
                if parent_id:
                    dialog.set_parent_id(parent_id)

            if dialog.exec() == QtWidgets.QDialog.Accepted:
                # Refresh tree sau khi thêm
                self.refresh_tree()

                # Tìm lại parent item sau khi refresh
                if parent_id:
                    self._expand_node_by_id(parent_id)

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể thêm node: {e}")

    def _add_child_node(self, parent_item):
        """Thêm node con"""
        if not parent_item:
            return

        parent_id = parent_item.data(0, Qt.UserRole)
        if not parent_id:
            return

        try:
            dialog = TreeNodeDialog(self.db, mode="add", parent=self)
            dialog.set_parent_id(parent_id)

            if dialog.exec() == QtWidgets.QDialog.Accepted:
                self.refresh_tree()
                # Tìm lại và expand parent sau khi refresh
                self._expand_node_by_id(parent_id)

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể thêm node con: {e}")

    def _edit_tree_node(self, item):
        """Sửa node"""
        if not item:
            return

        node_id = item.data(0, Qt.UserRole)
        if not node_id:
            return

        try:
            dialog = TreeNodeDialog(self.db, mode="edit", node_id=node_id, parent=self)

            if dialog.exec() == QtWidgets.QDialog.Accepted:
                self.refresh_tree()
                # Tìm lại và select node sau khi refresh
                self._select_node_by_id(node_id)

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể sửa node: {e}")

    def _edit_selected_tree_node(self):
        """Sửa node được chọn"""
        selected_items = self.tree.selectedItems()
        if selected_items:
            self._edit_tree_node(selected_items[0])

    def _copy_tree_node(self, item):
        """Sao chép node"""
        if not item:
            return

        node_id = item.data(0, Qt.UserRole)
        if not node_id:
            return

        try:
            # Lấy thông tin node gốc
            row = self.db.execute_query(
                "SELECT name, level, description, parent_id FROM exercise_tree WHERE id = ?",
                (node_id,), fetch="one"
            )

            if row:
                new_name = f"{row['name']} (Sao chép)"

                # Tạo node mới
                description = row.get('description', '') if row.get('description') else ''

                self.db.execute_query(
                    "INSERT INTO exercise_tree (parent_id, name, level, description) VALUES (?, ?, ?, ?)",
                    (row['parent_id'], new_name, row['level'], description)
                )

                self.refresh_tree()
                QtWidgets.QMessageBox.information(self, "Thành công", f"Đã sao chép '{new_name}'")

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể sao chép node: {e}")

    def _delete_tree_node(self, item):
        """Xóa node với xác nhận"""
        if not item:
            return

        node_id = item.data(0, Qt.UserRole)
        node_name = item.text(0)

        if not node_id:
            return

        try:
            # Kiểm tra node con
            children_count = self.db.execute_query(
                "SELECT COUNT(*) as count FROM exercise_tree WHERE parent_id = ?",
                (node_id,), fetch="one"
            )

            # Kiểm tra câu hỏi trong node
            questions_count = self.db.execute_query(
                "SELECT COUNT(*) as count FROM question_bank WHERE tree_id = ?",
                (node_id,), fetch="one"
            )

            warning_msg = f"Bạn có chắc muốn xóa nhánh '{node_name}'?"

            if children_count and children_count["count"] > 0:
                warning_msg += f"\n\n⚠️ Nhánh này có {children_count['count']} nhánh con."

            if questions_count and questions_count["count"] > 0:
                warning_msg += f"\n⚠️ Nhánh này chứa {questions_count['count']} câu hỏi."
                warning_msg += "\n\nTất cả dữ liệu sẽ bị xóa vĩnh viễn!"

            reply = QtWidgets.QMessageBox.question(
                self, "Xác nhận xóa",
                warning_msg,
                QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No,
                QtWidgets.QMessageBox.No
            )

            if reply == QtWidgets.QMessageBox.Yes:
                # Xóa node và tất cả con
                self.db.execute_query("DELETE FROM exercise_tree WHERE id = ?", (node_id,))
                self.refresh_tree()
                QtWidgets.QMessageBox.information(self, "Thành công", f"Đã xóa nhánh '{node_name}'")

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể xóa node: {e}")

    def _delete_selected_tree_node(self):
        """Xóa node được chọn"""
        selected_items = self.tree.selectedItems()
        if selected_items:
            self._delete_tree_node(selected_items[0])

    def _show_node_info(self, item):
        """Hiển thị thông tin node"""
        if not item:
            return

        node_id = item.data(0, Qt.UserRole)
        if not node_id:
            return

        try:
            # Lấy thông tin node
            node = self.db.execute_query(
                "SELECT * FROM exercise_tree WHERE id = ?",
                (node_id,), fetch="one"
            )

            if node:
                # Đếm số lượng con
                children_count = self.db.execute_query(
                    "SELECT COUNT(*) as count FROM exercise_tree WHERE parent_id = ?",
                    (node_id,), fetch="one"
                )["count"]

                # Đếm số câu hỏi
                questions_count = self.db.execute_query(
                    "SELECT COUNT(*) as count FROM question_bank WHERE tree_id = ?",
                    (node_id,), fetch="one"
                )["count"]

                info_text = f"""
                📁 THÔNG TIN NHÁNH

                ID: {node['id']}
                Tên: {node['name']}
                Cấp độ: {node['level']}
                Mô tả: {node.get('description', 'Không có')}

                📊 THỐNG KÊ:
                - Số nhánh con: {children_count}
                - Số câu hỏi: {questions_count}

                🕐 Ngày tạo: {node.get('created_at', 'Không rõ')}
                """

                QtWidgets.QMessageBox.information(self, "Thông tin nhánh", info_text)

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể lấy thông tin: {e}")

    def _expand_node_by_id(self, node_id):
        """Tìm và expand node theo ID"""
        try:
            root = self.tree.invisibleRootItem()
            self._find_and_expand_recursive(root, node_id)
        except Exception:
            pass

    def _find_and_expand_recursive(self, parent_item, target_id):
        """Đệ quy tìm và expand node"""
        for i in range(parent_item.childCount()):
            child = parent_item.child(i)
            if child and child.data(0, Qt.UserRole) == target_id:
                child.setExpanded(True)
                return True

            if self._find_and_expand_recursive(child, target_id):
                return True

        return False

    def _select_node_by_id(self, node_id):
        """Tìm và select node theo ID"""
        try:
            root = self.tree.invisibleRootItem()
            self._find_and_select_recursive(root, node_id)
        except Exception:
            pass

    def _find_and_select_recursive(self, parent_item, target_id):
        """Đệ quy tìm và select node"""
        for i in range(parent_item.childCount()):
            child = parent_item.child(i)
            if child and child.data(0, Qt.UserRole) == target_id:
                self.tree.setCurrentItem(child)
                return True

            if self._find_and_select_recursive(child, target_id):
                return True

        return False

    def open_tree_manager(self):
        """Mở cửa sổ quản lý cây nâng cao"""
        dialog = TreeManagerDialog(self.db, parent=self)
        dialog.exec()
        self.refresh_tree()

    # ========== NHÓM 4.5: TÍNH NĂNG HIỂN THỊ NÂNG CAO ========== #

    def on_quick_search(self):
        """Tìm kiếm nhanh theo nội dung"""
        keyword = self.quick_search.text().strip().lower()
        if len(keyword) < 2:  # Tối thiểu 2 ký tự
            self.apply_filters()
            return

        # Lọc các dòng trong bảng hiện tại
        for row in range(self.q_table.rowCount()):
            content_item = self.q_table.item(row, 2)  # Cột nội dung
            tags_item = self.q_table.item(row, 7)  # Cột tags

            show_row = False
            if content_item and keyword in content_item.text().lower():
                show_row = True
            elif tags_item and keyword in tags_item.text().lower():
                show_row = True

            self.q_table.setRowHidden(row, not show_row)

        self.update_stats_label()

    def apply_filters(self):
        """Áp dụng các bộ lọc"""
        content_filter = self.content_type_filter.currentText()
        difficulty_filter = self.difficulty_filter.currentText()

        visible_count = 0
        for row in range(self.q_table.rowCount()):
            show_row = True

            # Lọc theo loại content
            if content_filter != "Tất cả loại":
                type_item = self.q_table.item(row, 3)
                if type_item:
                    type_text = type_item.text()
                    if content_filter == "📝 Text" and "📝" not in type_text:
                        show_row = False
                    elif content_filter == "🖼️ Image" and "🖼️" not in type_text:
                        show_row = False
                    elif content_filter == "📄 PDF" and "📄" not in type_text:
                        show_row = False
                    elif content_filter == "📘 Word" and "📘" not in type_text:
                        show_row = False

            # Lọc theo độ khó
            if difficulty_filter != "Tất cả độ khó":
                diff_item = self.q_table.item(row, 4)
                if diff_item:
                    diff_text = diff_item.text()
                    if difficulty_filter == "🟢 Dễ" and "🟢" not in diff_text:
                        show_row = False
                    elif difficulty_filter == "🟡 Trung bình" and "🟡" not in diff_text:
                        show_row = False
                    elif difficulty_filter == "🔴 Khó" and "🔴" not in diff_text:
                        show_row = False

            self.q_table.setRowHidden(row, not show_row)
            if show_row:
                visible_count += 1

        self.update_stats_label(visible_count)

    def update_stats_label(self, visible_count=None):
        """Cập nhật label thống kê"""
        total = self.q_table.rowCount()
        if visible_count is None:
            visible_count = sum(1 for row in range(total) if not self.q_table.isRowHidden(row))

        if visible_count == total:
            self.stats_label.setText(f"{total} câu hỏi")
        else:
            self.stats_label.setText(f"{visible_count}/{total} câu hỏi")

    def show_view_options(self):
        """Hiển thị menu tùy chọn hiển thị"""
        menu = QMenu(self)

        # Tùy chọn ẩn/hiện cột
        columns_menu = menu.addMenu("Ẩn/Hiện cột")
        headers = ["☑️", "ID", "📊 Loại", "🎯 Độ khó", "✅ Đáp án", "📁 Chủ đề", "🏷️ Tags", "📅 Ngày tạo"]

        for i, header in enumerate(headers):
            action = columns_menu.addAction(header)
            action.setCheckable(True)
            action.setChecked(not self.q_table.isColumnHidden(i))
            action.triggered.connect(lambda checked, col=i: self.toggle_column(col, checked))

        menu.addSeparator()

        # Tùy chọn kích thước
        compact_action = menu.addAction("🔸 Chế độ thu gọn")
        compact_action.triggered.connect(self.toggle_compact_mode)

        detail_action = menu.addAction("🔹 Chế độ chi tiết")
        detail_action.triggered.connect(self.toggle_detail_mode)

        menu.exec(self.view_options_btn.mapToGlobal(self.view_options_btn.rect().bottomLeft()))

    def toggle_column(self, column, visible):
        """Ẩn/hiện cột"""
        self.q_table.setColumnHidden(column, not visible)

    def toggle_compact_mode(self):
        """Chuyển sang chế độ thu gọn"""
        # Ẩn một số cột không cần thiết
        columns_to_hide = [3, 5, 6, 7, 8]
        for col in columns_to_hide:
            self.q_table.setColumnHidden(col, True)

    def toggle_detail_mode(self):
        """Chuyển sang chế độ chi tiết"""
        # Hiện tất cả cột
        for col in range(self.q_table.columnCount()):
            self.q_table.setColumnHidden(col, False)

    def on_question_select_enhanced(self):
        """Xử lý chọn câu hỏi với preview nâng cao"""
        items = self.q_table.selectedItems()
        if not items:
            self.clear_preview()
            return

        row = items[0].row()
        question_id = self.q_table.item(row, 1).text()

        try:
            qid = int(question_id)
            self.load_question_preview(qid)
            self.current_question_id = qid
        except (ValueError, IndexError):
            self.clear_preview()

    def load_question_preview(self, question_id):
        """Load preview câu hỏi đơn giản - không dùng widget phức tạp"""
        q = self.db.execute_query("SELECT * FROM question_bank WHERE id=?", (question_id,), fetch="one")
        if not q:
            self.clear_preview()
            return

        try:
            # Chuyển đổi sang dict
            q_dict = self.row_to_dict(q)

            # Cập nhật info labels
            self.preview_id_label.setText(f"#{question_id}")

            # Chủ đề
            tree_id = self.safe_get(q_dict, "tree_id", 0)
            tree_path = self.get_tree_path(tree_id) if tree_id else []
            topic_path = " > ".join([p["name"] for p in tree_path]) if tree_path else "Chưa phân loại"
            self.preview_topic_label.setText(topic_path)

            # Độ khó
            difficulty = self.safe_get(q_dict, "difficulty_level", "medium")
            self.preview_difficulty_label.setText(difficulty.title())

            # Ngày tạo
            created_date = self.safe_get(q_dict, "created_date", "")
            date_str = created_date[:16] if created_date else "Không xác định"
            self.preview_date_label.setText(date_str)

            # Nội dung preview đơn giản
            content_type = self.safe_get(q_dict, "content_type", "text")
            content_text = self.safe_get(q_dict, "content_text", "")
            content_data = self.safe_get(q_dict, "content_data", None)

            if content_type == "image" and content_data:
                # Hiển thị ảnh trong preview
                self.preview_container.setCurrentIndex(1)  # Image viewer
                self.preview_image.load_image_from_data(content_data)
            else:
                # Hiển thị text
                self.preview_container.setCurrentIndex(0)  # Text label
                if content_text:
                    preview_text = content_text[:200] + ("..." if len(content_text) > 200 else "")
                else:
                    preview_text = "❌ [Nội dung trống]"
                self.preview_text.setText(preview_text)

        except Exception as e:
            print(f"⚠️ Lỗi load preview: {e}")
            self.clear_preview()

    def clear_preview(self):
        """Clear preview đơn giản"""
        try:
            if hasattr(self, 'preview_id_label'):
                self.preview_id_label.setText("-")
            if hasattr(self, 'preview_topic_label'):
                self.preview_topic_label.setText("-")
            if hasattr(self, 'preview_difficulty_label'):
                self.preview_difficulty_label.setText("-")
            if hasattr(self, 'preview_date_label'):
                self.preview_date_label.setText("-")
            if hasattr(self, 'preview_content'):
                self.preview_content.setText("Chọn câu hỏi để xem preview...")

            # Reset variables
            self.original_pixmap = None
            self.current_zoom = 1.0
        except Exception as e:
            print(f"⚠️ Lỗi clear_preview: {e}")

    # Chuột phải vào câu hỏi trong bảng ể hiện thị các menu
    def show_enhanced_context_menu(self, position):
        """Hiển thị context menu nâng cao"""
        item = self.q_table.itemAt(position)
        if not item:
            return

        menu = QMenu(self)

        # Các hành động cơ bản
        edit_action = menu.addAction("✏️ Chỉnh sửa câu hỏi")
        edit_action.triggered.connect(self.edit_current_question)

        view_action = menu.addAction("👁️ Xem chi tiết")
        view_action.triggered.connect(self.view_question_detail)

        menu.addSeparator()

        # Quản lý
        move_action = menu.addAction("🗂️ Di chuyển đến...")
        move_action.triggered.connect(self.move_question_to_folder)

        tags_action = menu.addAction("🏷️ Quản lý tags")
        tags_action.triggered.connect(self.manage_question_tags)

        copy_action = menu.addAction("📋 Sao chép")
        copy_action.triggered.connect(self.copy_question)

        menu.addSeparator()

        # Xóa
        delete_action = menu.addAction("🗑️ Xóa câu hỏi")
        delete_action.triggered.connect(self.delete_question)

        menu.exec(self.q_table.mapToGlobal(position))

    # ========== NHÓM 4: QUẢN LÝ DANH SÁCH CÂU HỎI ========== #
    def _load_question_rows(self, rows):
        """Load danh sách câu hỏi với xử lý lỗi an toàn"""
        if not rows:
            self.q_table.setRowCount(0)
            if hasattr(self, 'update_stats_label'):
                self.update_stats_label()
            return

        self.q_table.setRowCount(0)

        # ========== XỬ LÝ TỪNG DÒNG DỮ LIỆU VỚI ERROR HANDLING ========== #
        processed_count = 0
        error_count = 0

        for i, r in enumerate(rows):
            try:
                # Xử lý dữ liệu an toàn
                if r is None:
                    continue

                # Chuyển đổi sang dict nếu cần
                if isinstance(r, dict):
                    row_dict = r
                else:
                    row_dict = self.row_to_dict(r)

                if not row_dict:
                    print(f"⚠️ Empty row data at index {i}")
                    continue

                # Tạo checkbox
                checkbox = QtWidgets.QCheckBox()
                checkbox.setChecked(False)

                # Thông tin cơ bản với validation
                row_id = self.safe_get(row_dict, "id", 0)
                if not row_id:
                    continue

                # Loại content với fallback
                content_type = self.safe_get(row_dict, "content_type", "text")
                type_display = {
                    "text": "📝 Text",
                    "image": "🖼️ Image",
                    "pdf": "📄 PDF",
                    "word": "📘 Word",
                    "mixed": "🔀 Mixed"
                }.get(content_type, "📝 Text")

                # Độ khó với màu sắc
                difficulty = self.safe_get(row_dict, "difficulty_level", "medium")
                difficulty_colors = {
                    "easy": ("#d4edda", "🟢 Dễ"),
                    "medium": ("#fff3cd", "🟡 Trung bình"),
                    "hard": ("#f8d7da", "🔴 Khó")
                }
                difficulty_color, difficulty_display = difficulty_colors.get(difficulty, ("#f8f9fa", "🟡 Trung bình"))

                # Đáp án
                answer_text = self.safe_get(row_dict, "answer_text", "") or self.safe_get(row_dict, "correct_answer",
                                                                                          "")  # Fallback
                answer_type = self.safe_get(row_dict, "answer_type", "text")

                if answer_type == "image":
                    answer_display = "🖼️ [Đáp án hình ảnh]"
                elif answer_type == "pdf":
                    answer_display = "📄 [Đáp án PDF]"
                else:
                    answer_display = answer_text[:30] + (
                        "..." if len(answer_text) > 30 else "") if answer_text else "Chưa có đáp án"

                # Trạng thái
                status = self.safe_get(row_dict, "status", "active")
                status_display = {
                    "active": "✅ Hoạt động",
                    "draft": "📝 Nháp",
                    "archived": "📁 Lưu trữ"
                }.get(status, "✅ Hoạt động")

                # Thống kê sử dụng
                usage_count = self.safe_get(row_dict, "usage_count", 0)
                usage_display = f"{usage_count} lần"

                # Thời gian
                created_date = self.safe_get(row_dict, "created_date", "")
                if created_date:
                    try:
                        from datetime import datetime
                        if 'T' in str(created_date):
                            dt = datetime.fromisoformat(str(created_date).replace('Z', '+00:00'))
                        else:
                            dt = datetime.strptime(str(created_date)[:19], '%Y-%m-%d %H:%M:%S')
                        date_display = dt.strftime("%d/%m/%Y")
                    except:
                        date_display = str(created_date)[:10]
                else:
                    date_display = ""

                # Chủ đề từ tree
                tree_id = self.safe_get(row_dict, "tree_id", 0)
                try:
                    path = self.get_tree_path(tree_id) if tree_id else []
                    topic_display = " > ".join([p.get("name", "") for p in path[-2:]]) if len(path) >= 2 else (
                        path[0].get("name", "") if path else "Chưa phân loại")
                except:
                    topic_display = "Lỗi load chủ đề"

                # Tags
                try:
                    tags = self.db.execute_query(
                        "SELECT tag_name FROM question_tags WHERE question_id=?",
                        (row_id,), fetch="all"
                    ) or []
                    tags_text = ", ".join(
                        [self.safe_get(dict(tag) if hasattr(tag, 'keys') else {}, "tag_name", "") for tag in
                         tags]) if tags else ""
                except:
                    tags_text = ""

                # Thêm dòng vào bảng (cập nhật số cột nếu cần)
                row_idx = self.q_table.rowCount()
                self.q_table.insertRow(row_idx)

                # Đặt dữ liệu vào các cột
                try:
                    self.q_table.setCellWidget(row_idx, 0, checkbox)
                    self.q_table.setItem(row_idx, 1, QtWidgets.QTableWidgetItem(str(row_id)))
                    self.q_table.setItem(row_idx, 2, QtWidgets.QTableWidgetItem(type_display))

                    # Cột độ khó với màu nền
                    difficulty_item = QtWidgets.QTableWidgetItem(difficulty_display)
                    difficulty_item.setBackground(QtGui.QColor(difficulty_color))
                    self.q_table.setItem(row_idx, 3, difficulty_item)

                    self.q_table.setItem(row_idx, 4, QtWidgets.QTableWidgetItem(answer_display))
                    self.q_table.setItem(row_idx, 5, QtWidgets.QTableWidgetItem(topic_display))

                    # Tags với màu nền
                    tags_item = QtWidgets.QTableWidgetItem(tags_text)
                    if tags_text:
                        tags_item.setBackground(QtGui.QColor("#e3f2fd"))
                    self.q_table.setItem(row_idx, 6, tags_item)

                    # Kiểm tra số cột có đủ không
                    if self.q_table.columnCount() > 7:
                        self.q_table.setItem(row_idx, 7, QtWidgets.QTableWidgetItem(usage_display))
                    if self.q_table.columnCount() > 8:
                        self.q_table.setItem(row_idx, 8, QtWidgets.QTableWidgetItem(date_display))

                    processed_count += 1

                except Exception as e:
                    print(f"⚠️ Lỗi set table item row {row_idx}: {e}")
                    error_count += 1

            except Exception as e:
                print(f"⚠️ Lỗi xử lý row {i}: {e}")
                error_count += 1
                continue

        print(f"✅ Đã load {processed_count} câu hỏi, {error_count} lỗi")

        # Cập nhật thống kê
        if hasattr(self, 'update_stats_label'):
            try:
                self.update_stats_label()
            except:
                pass

    def on_question_select(self):
        """Load câu hỏi được chọn với xử lý sqlite3.Row an toàn"""
        items = self.q_table.selectedItems()
        if not items:
            return

        row = items[0].row()
        item_text = self.q_table.item(row, 1).text()

        try:
            qid = int(item_text)
        except (ValueError, IndexError):
            return

        # ========== TẢI DỮ LIỆU CÂU HỎI VỚI XỬ LÝ AN TOÀN ========== #
        q = self.db.execute_query("SELECT * FROM question_bank WHERE id=?", (qid,), fetch="one")
        if not q:
            return

        # Chuyển đổi sang dict để sử dụng .get()
        q_dict = self.row_to_dict(q)

        # Cập nhật ID hiện tại
        self.current_question_id = qid

        # Hiển thị nội dung với truy cập an toàn
        content_text = self.safe_get(q_dict, "content_text", "")
        content_type = self.safe_get(q_dict, "content_type", "text")

        if hasattr(self, 'content_text'):
            self.content_text.setPlainText(content_text)

        # Hiển thị đáp án
        answer_text = self.safe_get(q_dict, "answer_text", "") or self.safe_get(q_dict, "correct_answer",
                                                                                "")  # Fallback
        answer_type = self.safe_get(q_dict, "answer_type", "text")
        answer_data = self.safe_get(q_dict, "answer_data", "")

        # Load preview nếu có method
        if hasattr(self, 'load_question_preview'):
            self.load_question_preview(qid)

        # Load lịch sử nếu có method
        if hasattr(self, '_load_question_history'):
            self._load_question_history(qid)

    # (Phương thức xử lý double-click vào ô trong bảng)
    def on_cell_double_clicked(self, row, column):
        """Xử lý double-click vào ô trong bảng câu hỏi"""
        try:
            # Kiểm tra nếu double-click vào cột "Đáp án" (cột 4)
            if column != 4:
                return

            # Lấy ID câu hỏi từ cột ID (cột 1)
            id_item = self.q_table.item(row, 1)
            if not id_item:
                return

            question_id = int(id_item.text())

            # Truy vấn thông tin câu hỏi từ database
            question = self.db.execute_query(
                "SELECT answer_type, answer_data, answer_text FROM question_bank WHERE id=?",
                (question_id,), fetch="one"
            )

            if not question:
                QtWidgets.QMessageBox.warning(self, "Lỗi", "Không tìm thấy câu hỏi!")
                return

            # Chuyển đổi sang dict để truy cập an toàn
            q_dict = dict(question) if hasattr(question, 'keys') else question

            answer_type = self.safe_get(q_dict, "answer_type", "text")
            answer_data = self.safe_get(q_dict, "answer_data")

            # ÁP DỤNG FALLBACK LOGIC GIỐNG CÁC PHƯƠNG THỨC KHÁC
            answer_text = self.safe_get(q_dict, "answer_text", "")

            # Kiểm tra đáp án ảnh thuần túy
            if answer_type == "image" and answer_data:
                self.show_answer_image_dialog(question_id, answer_data)
                return

            # Hiển thị đáp án text thông thường
            if answer_text:
                import re
                clean_text = re.sub(r'<[^>]+>', '', answer_text).strip()
                if clean_text:
                    QtWidgets.QMessageBox.information(
                        self, "Đáp án câu hỏi",
                        f"Đáp án: {clean_text}"
                    )
                else:
                    QtWidgets.QMessageBox.information(
                        self, "Đáp án câu hỏi",
                        "Đáp án chứa nội dung HTML phức tạp"
                    )
            else:
                QtWidgets.QMessageBox.information(
                    self, "Đáp án câu hỏi",
                    "Chưa có đáp án"
                )
        except Exception as e:
            QtWidgets.QMessageBox.warning(self, "Lỗi", f"Không thể hiển thị đáp án: {e}")

    # (Hiển thị dialog ảnh đáp án với giao diện đầy đủ giống show_html_answer_dialog)
    def show_answer_image_dialog(self, question_id, answer_data):
        """Hiển thị dialog chứa hình ảnh đáp án với giao diện chuyên nghiệp giống ImageViewerDialog"""
        try:
            # Tạo dialog với styling chuyên nghiệp
            dialog = QtWidgets.QDialog(self)
            dialog.setWindowTitle(f"🖼️ Đáp án hình ảnh - Câu #{question_id}")
            dialog.setModal(False)
            dialog.resize(1200, 900)

            # (Window flags đầy đủ)
            dialog.setWindowFlags(
                Qt.Window |
                Qt.WindowMinimizeButtonHint |
                Qt.WindowMaximizeButtonHint |
                Qt.WindowCloseButtonHint |
                Qt.WindowSystemMenuHint |
                Qt.WindowTitleHint |
                Qt.CustomizeWindowHint
            )

            dialog.setMinimumSize(400, 300)
            dialog.setSizeGripEnabled(True)

            layout = QtWidgets.QVBoxLayout(dialog)
            layout.setContentsMargins(0, 0, 0, 0)
            layout.setSpacing(0)

            # Load ảnh từ binary data trước
            original_pixmap = self.load_pixmap_from_answer_data(answer_data)
            current_zoom = 1.0

            # (Main toolbar với styling chuyên nghiệp - giống ImageViewerDialog)
            main_toolbar = QtWidgets.QWidget()
            main_toolbar.setFixedHeight(55)
            main_toolbar.setStyleSheet("""
                QWidget {
                    background: qlineargradient(x1: 0, y1: 0, x2: 0, y2: 1,
                                              stop: 0 #f8f9fa, stop: 1 #e9ecef);
                    border-bottom: 2px solid #adb5bd;
                }
                QPushButton {
                    background: white;
                    border: 1px solid #ced4da;
                    border-radius: 6px;
                    color: #495057;
                    font-weight: 500;
                    padding: 8px 12px;
                    margin: 2px;
                    min-width: 80px;
                }
                QPushButton:hover {
                    background: #e3f2fd;
                    border-color: #2196f3;
                    color: #1976d2;
                }
                QPushButton:pressed {
                    background: #bbdefb;
                }
                QPushButton:checked {
                    background: #4caf50;
                    color: white;
                    border-color: #388e3c;
                }
                QSlider::groove:horizontal {
                    border: 1px solid #bbb;
                    height: 8px;
                    background: white;
                    border-radius: 4px;
                }
                QSlider::handle:horizontal {
                    background: #2196f3;
                    border: 1px solid #1976d2;
                    width: 18px;
                    border-radius: 9px;
                    margin: -5px 0;
                }
                QSlider::handle:horizontal:hover {
                    background: #1976d2;
                }
                QLabel {
                    color: #495057;
                    font-weight: 500;
                    padding: 0 8px;
                }
            """)

            toolbar_layout = QtWidgets.QHBoxLayout(main_toolbar)
            toolbar_layout.setContentsMargins(15, 8, 15, 8)
            toolbar_layout.setSpacing(8)

            # (Nhóm điều khiển cửa sổ)
            window_group = QtWidgets.QWidget()
            window_layout = QtWidgets.QHBoxLayout(window_group)
            window_layout.setContentsMargins(0, 0, 0, 0)
            window_layout.setSpacing(4)

            minimize_btn = QtWidgets.QPushButton("🗕 Thu nhỏ")
            minimize_btn.setToolTip("Thu nhỏ cửa sổ (Alt+F9)")
            minimize_btn.clicked.connect(dialog.showMinimized)

            maximize_btn = QtWidgets.QPushButton("🗖 Phóng to")
            maximize_btn.setToolTip("Phóng to/Khôi phục (Alt+F10)")

            hide_btn = QtWidgets.QPushButton("👁️ Ẩn")
            hide_btn.setToolTip("Ẩn cửa sổ (Ctrl+H)")
            hide_btn.clicked.connect(dialog.hide)

            window_layout.addWidget(minimize_btn)
            window_layout.addWidget(maximize_btn)
            window_layout.addWidget(hide_btn)
            toolbar_layout.addWidget(window_group)

            # Separator
            separator1 = QtWidgets.QFrame()
            separator1.setFrameShape(QtWidgets.QFrame.VLine)
            separator1.setStyleSheet("color: #adb5bd;")
            toolbar_layout.addWidget(separator1)

            # (Nhóm zoom controls với slider)
            zoom_group = QtWidgets.QWidget()
            zoom_layout = QtWidgets.QHBoxLayout(zoom_group)
            zoom_layout.setContentsMargins(0, 0, 0, 0)
            zoom_layout.setSpacing(4)

            zoom_out_btn = QtWidgets.QPushButton("🔍−")
            zoom_out_btn.setToolTip("Thu nhỏ (Ctrl + -)")
            zoom_out_btn.setFixedWidth(40)

            # Zoom slider với styling đẹp
            zoom_slider = QtWidgets.QSlider(Qt.Horizontal)
            zoom_slider.setRange(10, 500)  # 10% đến 500%
            zoom_slider.setValue(int(current_zoom * 100))
            zoom_slider.setFixedWidth(120)
            zoom_slider.setToolTip("Kéo để điều chỉnh zoom")

            zoom_in_btn = QtWidgets.QPushButton("🔍+")
            zoom_in_btn.setToolTip("Phóng to (Ctrl + +)")
            zoom_in_btn.setFixedWidth(40)

            # Zoom percentage display
            zoom_label = QtWidgets.QLabel("100%")
            zoom_label.setFixedWidth(50)
            zoom_label.setAlignment(Qt.AlignCenter)

            zoom_layout.addWidget(zoom_out_btn)
            zoom_layout.addWidget(zoom_slider)
            zoom_layout.addWidget(zoom_in_btn)
            zoom_layout.addWidget(zoom_label)
            toolbar_layout.addWidget(zoom_group)

            # Separator
            separator2 = QtWidgets.QFrame()
            separator2.setFrameShape(QtWidgets.QFrame.VLine)
            separator2.setStyleSheet("color: #adb5bd;")
            toolbar_layout.addWidget(separator2)

            # (Nhóm fit controls)
            fit_group = QtWidgets.QWidget()
            fit_layout = QtWidgets.QHBoxLayout(fit_group)
            fit_layout.setContentsMargins(0, 0, 0, 0)
            fit_layout.setSpacing(4)

            fit_window_btn = QtWidgets.QPushButton("📐 Vừa cửa sổ")
            fit_window_btn.setToolTip("Vừa khít cửa sổ (Ctrl + 0)")

            actual_size_btn = QtWidgets.QPushButton("1:1 Gốc")
            actual_size_btn.setToolTip("Hiển thị kích thước gốc (Ctrl + 1)")

            fit_layout.addWidget(fit_window_btn)
            fit_layout.addWidget(actual_size_btn)
            toolbar_layout.addWidget(fit_group)

            # Separator
            separator3 = QtWidgets.QFrame()
            separator3.setFrameShape(QtWidgets.QFrame.VLine)
            separator3.setStyleSheet("color: #adb5bd;")
            toolbar_layout.addWidget(separator3)

            # (Nhóm tiện ích)
            utility_group = QtWidgets.QWidget()
            utility_layout = QtWidgets.QHBoxLayout(utility_group)
            utility_layout.setContentsMargins(0, 0, 0, 0)
            utility_layout.setSpacing(4)

            always_top_btn = QtWidgets.QPushButton("📌 Luôn trên")
            always_top_btn.setCheckable(True)
            always_top_btn.setToolTip("Giữ cửa sổ luôn ở trên (Ctrl+T)")

            opacity_btn = QtWidgets.QPushButton("🔅 Độ mờ")
            opacity_btn.setToolTip("Điều chỉnh độ trong suốt")

            utility_layout.addWidget(always_top_btn)
            utility_layout.addWidget(opacity_btn)
            toolbar_layout.addWidget(utility_group)

            # Spacer để đẩy close button sang phải
            toolbar_layout.addStretch()

            # (Nút đóng nổi bật)
            close_btn = QtWidgets.QPushButton("✖ Đóng")
            close_btn.setToolTip("Đóng cửa sổ (ESC)")
            close_btn.clicked.connect(dialog.close)
            close_btn.setStyleSheet("""
                QPushButton {
                    background: #dc3545;
                    color: white;
                    border: 1px solid #c82333;
                    font-weight: bold;
                    min-width: 60px;
                }
                QPushButton:hover {
                    background: #c82333;
                    border-color: #bd2130;
                }
                QPushButton:pressed {
                    background: #bd2130;
                }
            """)
            toolbar_layout.addWidget(close_btn)

            layout.addWidget(main_toolbar)

            # (Container chính với background tối chuyên nghiệp - giống ImageViewerDialog)
            main_container = QtWidgets.QWidget()
            main_container.setStyleSheet("""
                QWidget {
                    background: qlineargradient(x1: 0, y1: 0, x2: 0, y2: 1,
                                              stop: 0 #1a1a1a, stop: 1 #000000);
                }
            """)
            main_layout = QtWidgets.QVBoxLayout(main_container)
            main_layout.setContentsMargins(8, 8, 8, 8)

            # Scroll area với styling cao cấp
            scroll_area = QtWidgets.QScrollArea()
            scroll_area.setAlignment(Qt.AlignCenter)
            scroll_area.setStyleSheet("""
                QScrollArea { 
                    border: 2px solid #333;
                    border-radius: 8px;
                    background: #000000;
                }
                QScrollArea QScrollBar:vertical {
                    background: #2b2b2b;
                    width: 12px;
                    border-radius: 6px;
                }
                QScrollArea QScrollBar::handle:vertical {
                    background: #555;
                    border-radius: 6px;
                    min-height: 30px;
                }
                QScrollArea QScrollBar::handle:vertical:hover {
                    background: #777;
                }
                QScrollArea QScrollBar:horizontal {
                    background: #2b2b2b;
                    height: 12px;
                    border-radius: 6px;
                }
                QScrollArea QScrollBar::handle:horizontal {
                    background: #555;
                    border-radius: 6px;
                    min-width: 30px;
                }
                QScrollArea QScrollBar::handle:horizontal:hover {
                    background: #777;
                }
            """)

            # (Image label với effects và interactions - giống ImageViewerDialog)
            image_label = QtWidgets.QLabel()
            image_label.setAlignment(Qt.AlignCenter)
            image_label.setStyleSheet("""
                QLabel { 
                    background: transparent;
                    border: 1px solid #444;
                    border-radius: 4px;
                    padding: 5px;
                }
            """)
            image_label.setMinimumSize(200, 150)

            if original_pixmap and not original_pixmap.isNull():
                image_label.setPixmap(original_pixmap)
                image_label.setCursor(Qt.OpenHandCursor)
            else:
                image_label.setText("⚠️ Không thể hiển thị ảnh")
                image_label.setStyleSheet(
                    "color: #dc3545; font-size: 18px; background: #2b2b2b; border: 2px dashed #dc3545;")

            scroll_area.setWidget(image_label)
            main_layout.addWidget(scroll_area)
            layout.addWidget(main_container)

            # (Status bar nâng cao với thông tin chi tiết - giống ImageViewerDialog)
            status_bar = QtWidgets.QWidget()
            status_bar.setFixedHeight(35)
            status_bar.setStyleSheet("""
                QWidget {
                    background: qlineargradient(x1: 0, y1: 0, x2: 0, y2: 1,
                                              stop: 0 #2b2b2b, stop: 1 #1a1a1a);
                    border-top: 1px solid #444;
                }
                QLabel {
                    color: #ffffff;
                    padding: 0 12px;
                    font-size: 11px;
                }
                QLabel#main_info {
                    font-weight: bold;
                    color: #4fc3f7;
                }
                QLabel#tips {
                    color: #81c784;
                    font-style: italic;
                }
            """)

            status_layout = QtWidgets.QHBoxLayout(status_bar)
            status_layout.setContentsMargins(8, 0, 8, 0)

            # Main image info
            main_status_label = QtWidgets.QLabel()
            main_status_label.setObjectName("main_info")
            status_layout.addWidget(main_status_label)

            # Spacer
            status_layout.addStretch()

            # Tips label
            tips_label = QtWidgets.QLabel("💡 Kéo chuột để di chuyển • Ctrl+Scroll để zoom • Double-click để fit")
            tips_label.setObjectName("tips")
            status_layout.addWidget(tips_label)

            layout.addWidget(status_bar)

            # (Logic zoom functions - giữ nguyên logic, chỉ cập nhật UI)
            def update_zoom_display():
                if original_pixmap and not original_pixmap.isNull():
                    new_size = original_pixmap.size() * current_zoom
                    scaled_pixmap = original_pixmap.scaled(
                        new_size, Qt.KeepAspectRatio, Qt.SmoothTransformation
                    )
                    image_label.setPixmap(scaled_pixmap)
                    image_label.resize(new_size)

                    # Update UI elements
                    zoom_percent = int(current_zoom * 100)
                    zoom_label.setText(f"{zoom_percent}%")
                    zoom_slider.blockSignals(True)
                    zoom_slider.setValue(zoom_percent)
                    zoom_slider.blockSignals(False)

                    # Update status bar
                    original_size = original_pixmap.size()
                    main_status_label.setText(
                        f"🖼️ Gốc: {original_size.width()}×{original_size.height()} | "
                        f"Hiển thị: {new_size.width()}×{new_size.height()} | "
                        f"Zoom: {zoom_percent}%"
                    )

            def zoom_in():
                nonlocal current_zoom
                if current_zoom < 5.0:
                    current_zoom *= 1.25
                    update_zoom_display()

            def zoom_out():
                nonlocal current_zoom
                if current_zoom > 0.1:
                    current_zoom *= 0.8
                    update_zoom_display()

            def fit_to_window():
                nonlocal current_zoom
                if original_pixmap and not original_pixmap.isNull():
                    available_size = scroll_area.viewport().size()
                    scale_w = available_size.width() / original_pixmap.width()
                    scale_h = available_size.height() / original_pixmap.height()
                    current_zoom = min(scale_w, scale_h) * 0.9
                    update_zoom_display()

            def actual_size():
                nonlocal current_zoom
                current_zoom = 1.0
                update_zoom_display()

            def zoom_slider_changed(value):
                nonlocal current_zoom
                current_zoom = value / 100.0
                update_zoom_display()

            # Connect zoom controls
            zoom_in_btn.clicked.connect(zoom_in)
            zoom_out_btn.clicked.connect(zoom_out)
            fit_window_btn.clicked.connect(fit_to_window)
            actual_size_btn.clicked.connect(actual_size)
            zoom_slider.valueChanged.connect(zoom_slider_changed)

            # (Logic window controls)
            def toggle_maximize():
                if dialog.isMaximized():
                    dialog.showNormal()
                    maximize_btn.setText("🗖 Phóng to")
                else:
                    dialog.showMaximized()
                    maximize_btn.setText("🗗 Khôi phục")

            maximize_btn.clicked.connect(toggle_maximize)

            def toggle_always_on_top(checked):
                if checked:
                    dialog.setWindowFlags(dialog.windowFlags() | Qt.WindowStaysOnTopHint)
                    dialog.show()
                    always_top_btn.setText("📌 Luôn trên ✓")
                else:
                    dialog.setWindowFlags(dialog.windowFlags() & ~Qt.WindowStaysOnTopHint)
                    dialog.show()
                    always_top_btn.setText("📌 Luôn trên")

            always_top_btn.toggled.connect(toggle_always_on_top)

            def show_opacity_dialog():
                opacity_dialog = QtWidgets.QDialog(dialog)
                opacity_dialog.setWindowTitle("Điều chỉnh độ mờ")
                opacity_dialog.resize(320, 140)

                dlg_layout = QtWidgets.QVBoxLayout(opacity_dialog)
                dlg_layout.setSpacing(15)

                title = QtWidgets.QLabel("🔅 Điều chỉnh độ trong suốt")
                title.setStyleSheet("font-weight: bold; font-size: 14px; color: #2c3e50;")
                title.setAlignment(Qt.AlignCenter)
                dlg_layout.addWidget(title)

                slider = QtWidgets.QSlider(Qt.Horizontal)
                slider.setRange(20, 100)
                slider.setValue(int(dialog.windowOpacity() * 100))

                label = QtWidgets.QLabel(f"Độ mờ: {slider.value()}%")
                label.setAlignment(Qt.AlignCenter)

                def update_opacity(value):
                    dialog.setWindowOpacity(value / 100.0)
                    label.setText(f"Độ mờ: {value}%")

                slider.valueChanged.connect(update_opacity)

                dlg_layout.addWidget(label)
                dlg_layout.addWidget(slider)

                btn_layout = QtWidgets.QHBoxLayout()
                reset_btn = QtWidgets.QPushButton("🔄 Reset")
                reset_btn.clicked.connect(lambda: slider.setValue(100))
                close_dlg_btn = QtWidgets.QPushButton("✓ Đóng")
                close_dlg_btn.clicked.connect(opacity_dialog.accept)

                btn_layout.addWidget(reset_btn)
                btn_layout.addWidget(close_dlg_btn)
                dlg_layout.addLayout(btn_layout)

                opacity_dialog.exec()

            opacity_btn.clicked.connect(show_opacity_dialog)

            # (Mouse interactions - giống ImageViewerDialog)
            is_panning = False
            last_pan_point = QtCore.QPoint()

            def mouse_press_event(event):
                nonlocal is_panning, last_pan_point
                if event.button() == Qt.LeftButton:
                    is_panning = True
                    last_pan_point = event.pos()
                    image_label.setCursor(Qt.ClosedHandCursor)

            def mouse_move_event(event):
                nonlocal is_panning, last_pan_point
                if is_panning and (event.buttons() & Qt.LeftButton):
                    delta = event.pos() - last_pan_point
                    h_scroll = scroll_area.horizontalScrollBar()
                    v_scroll = scroll_area.verticalScrollBar()
                    h_scroll.setValue(h_scroll.value() - delta.x())
                    v_scroll.setValue(v_scroll.value() - delta.y())
                    last_pan_point = event.pos()

            def mouse_release_event(event):
                nonlocal is_panning
                if event.button() == Qt.LeftButton:
                    is_panning = False
                    image_label.setCursor(Qt.OpenHandCursor)

            def mouse_double_click_event(event):
                if event.button() == Qt.LeftButton:
                    fit_to_window()

            # Set mouse events
            if original_pixmap and not original_pixmap.isNull():
                image_label.mousePressEvent = mouse_press_event
                image_label.mouseMoveEvent = mouse_move_event
                image_label.mouseReleaseEvent = mouse_release_event
                image_label.mouseDoubleClickEvent = mouse_double_click_event
                image_label.setMouseTracking(True)

            # (Mouse wheel zoom)
            def wheelEvent(event):
                modifiers = QtWidgets.QApplication.keyboardModifiers()
                if modifiers == Qt.ControlModifier:
                    if event.angleDelta().y() > 0:
                        zoom_in()
                    else:
                        zoom_out()
                    event.accept()
                else:
                    QtWidgets.QScrollArea.wheelEvent(scroll_area, event)

            scroll_area.wheelEvent = wheelEvent

            # (Context menu - giống ImageViewerDialog)
            def show_context_menu(pos):
                menu = QtWidgets.QMenu(dialog)
                menu.setStyleSheet("""
                    QMenu {
                        background: #2b2b2b;
                        color: white;
                        border: 1px solid #555;
                        border-radius: 4px;
                        padding: 4px;
                    }
                    QMenu::item {
                        padding: 8px 20px;
                        border-radius: 3px;
                    }
                    QMenu::item:selected {
                        background: #4fc3f7;
                    }
                    QMenu::separator {
                        height: 1px;
                        background: #555;
                        margin: 4px 0;
                    }
                """)

                # Window controls
                menu.addAction("🗕 Thu nhỏ", dialog.showMinimized)
                if dialog.isMaximized():
                    menu.addAction("🗗 Khôi phục", dialog.showNormal)
                else:
                    menu.addAction("🗖 Phóng to", dialog.showMaximized)
                menu.addAction("👁️ Ẩn", dialog.hide)
                menu.addSeparator()

                # Zoom controls
                if original_pixmap and not original_pixmap.isNull():
                    menu.addAction("🔍+ Phóng to ảnh", zoom_in)
                    menu.addAction("🔍− Thu nhỏ ảnh", zoom_out)
                    menu.addAction("📐 Vừa cửa sổ", fit_to_window)
                    menu.addAction("1:1 Kích thước gốc", actual_size)
                    menu.addSeparator()
                    menu.addAction("💾 Lưu ảnh...", lambda: self._save_image_to_file(original_pixmap))
                    menu.addAction("📋 Copy ảnh", lambda: QtWidgets.QApplication.clipboard().setPixmap(original_pixmap))
                    menu.addSeparator()

                # Window utilities
                always_top_text = "📌 Luôn trên ✓" if always_top_btn.isChecked() else "📌 Luôn trên"
                menu.addAction(always_top_text, lambda: always_top_btn.toggle())
                menu.addAction("🔅 Độ mờ...", show_opacity_dialog)
                menu.addSeparator()
                menu.addAction("✖ Đóng", dialog.close)

                menu.exec(dialog.mapToGlobal(pos))

            dialog.setContextMenuPolicy(Qt.CustomContextMenu)
            dialog.customContextMenuRequested.connect(show_context_menu)

            # (Keyboard shortcuts)
            QtGui.QShortcut(QtGui.QKeySequence("Alt+F9"), dialog).activated.connect(dialog.showMinimized)
            QtGui.QShortcut(QtGui.QKeySequence("Alt+F10"), dialog).activated.connect(toggle_maximize)
            QtGui.QShortcut(QtGui.QKeySequence("Ctrl+H"), dialog).activated.connect(dialog.hide)
            QtGui.QShortcut(QtGui.QKeySequence("Ctrl+T"), dialog).activated.connect(lambda: always_top_btn.toggle())
            QtGui.QShortcut(QtGui.QKeySequence("Escape"), dialog).activated.connect(dialog.close)
            QtGui.QShortcut(QtGui.QKeySequence("Ctrl++"), dialog).activated.connect(zoom_in)
            QtGui.QShortcut(QtGui.QKeySequence("Ctrl+-"), dialog).activated.connect(zoom_out)
            QtGui.QShortcut(QtGui.QKeySequence("Ctrl+0"), dialog).activated.connect(fit_to_window)
            QtGui.QShortcut(QtGui.QKeySequence("Ctrl+1"), dialog).activated.connect(actual_size)

            # Initial setup
            if original_pixmap and not original_pixmap.isNull():
                QtCore.QTimer.singleShot(100, fit_to_window)
            else:
                main_status_label.setText("❌ Không thể tải ảnh đáp án")

            # Show dialog
            dialog.show()
            dialog.raise_()
            dialog.activateWindow()

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể hiển thị ảnh đáp án: {e}")
    # (Helper method save ảnh ra file)
    def _save_image_to_file(self, pixmap):
        """Lưu ảnh ra file"""
        try:
            file_path, _ = QtWidgets.QFileDialog.getSaveFileName(
                self, "Lưu ảnh đáp án",
                f"dap_an_anh_{QtCore.QDateTime.currentDateTime().toString('yyyyMMdd_hhmmss')}.png",
                "PNG files (*.png);;JPG files (*.jpg);;All files (*.*)"
            )

            if file_path:
                if pixmap.save(file_path):
                    QtWidgets.QMessageBox.information(self, "Thành công", f"Đã lưu ảnh: {file_path}")
                else:
                    QtWidgets.QMessageBox.warning(self, "Lỗi", "Không thể lưu ảnh")

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Lỗi lưu ảnh: {e}")
    # #(Phương thức hiển thị dialog đáp án HTML chứa ảnh)
    # #(Helper method hiển thị mã nguồn HTML)
    def show_html_source_dialog(self, html_content):
        """Hiển thị mã nguồn HTML trong dialog riêng"""
        try:
            dialog = QtWidgets.QDialog(self)
            dialog.setWindowTitle("Mã nguồn HTML đáp án")
            dialog.setModal(True)
            dialog.resize(800, 600)

            layout = QtWidgets.QVBoxLayout(dialog)

            # Text editor để hiển thị HTML source
            source_editor = QtWidgets.QTextEdit()
            source_editor.setPlainText(html_content)
            source_editor.setReadOnly(True)
            source_editor.setFont(QtGui.QFont("Courier", 10))

            layout.addWidget(source_editor)

            # Nút đóng
            button_layout = QtWidgets.QHBoxLayout()
            close_btn = QtWidgets.QPushButton("Đóng")
            close_btn.clicked.connect(dialog.accept)
            button_layout.addStretch()
            button_layout.addWidget(close_btn)
            layout.addLayout(button_layout)

            dialog.exec()

        except Exception as e:
            QtWidgets.QMessageBox.warning(self, "Lỗi", f"Không thể hiển thị mã nguồn: {e}")

    def show_table_context_menu(self, position):
        """Hiển thị context menu cho bảng"""
        if not self.q_table.itemAt(position):
            return

        menu = QtWidgets.QMenu(self)

        edit_action = menu.addAction("✏️ Chỉnh sửa")
        edit_action.triggered.connect(self.on_question_select)

        menu.addSeparator()

        export_menu = menu.addMenu("📤 Xuất")
        export_menu.addAction("Xuất ra Word").triggered.connect(self.export_to_word)
        export_menu.addAction("Xuất ra PDF").triggered.connect(self.export_to_pdf)

        menu.addSeparator()

        delete_action = menu.addAction("🗑️ Xóa")
        delete_action.triggered.connect(self.delete_question)

        menu.exec(self.q_table.mapToGlobal(position))

    def get_tree_path(self, tree_id):
        """Lấy đường dẫn cây thư mục với xử lý an toàn"""
        if not tree_id:
            return []

        path = []
        current_id = tree_id

        # ========== TRAVERSE TREE VỚI XỬ LÝ AN TOÀN ========== #
        while current_id and len(path) < 10:  # Giới hạn độ sâu để tránh vòng lặp vô tận
            try:
                row = self.db.execute_query(
                    "SELECT id, parent_id, name, level FROM exercise_tree WHERE id=?",
                    (current_id,), fetch="one"
                )

                if row:
                    row_dict = self.row_to_dict(row)
                    path.insert(0, row_dict)
                    current_id = self.safe_get(row_dict, "parent_id")
                else:
                    break

            except Exception as e:
                print(f"⚠️ Lỗi get_tree_path: {e}")
                break

        return path

    # ========== NHÓM 5: LƯU/CẬP NHẬT/XÓA CÂU HỎI ========== #
    def save_question(self):
        """Lưu câu hỏi với cấu trúc mới"""
        tree_id = self._current_tree_id()
        if not tree_id:
            QtWidgets.QMessageBox.warning(self, "Chưa chọn thư mục", "Vui lòng chọn vị trí lưu trong cây.")
            return

        content = self.content_text.toPlainText().strip()
        answer = self.answer_text.toPlainText().strip()

        if not content:
            QtWidgets.QMessageBox.warning(self, "Lỗi", "Nội dung câu hỏi không được để trống")
            return

        try:
            # Chuẩn bị dữ liệu câu hỏi
            question_data = {
                'content_text': content,
                'content_type': getattr(self, 'content_type', 'text'),
                'answer_type': 'text',
                'correct_answer': answer,
                'tree_id': tree_id,
                'difficulty_level': getattr(self, 'current_difficulty', 'medium'),
                'question_type': 'knowledge',
                'status': 'active',
                'created_by': 'user'
            }

            if self.current_question_id:
                # Cập nhật câu hỏi hiện có
                success = self.db.update_question(
                    self.current_question_id,
                    question_data,
                    changed_fields=['content_text', 'correct_answer', 'tree_id']
                )

                if success:
                    QtWidgets.QMessageBox.information(self, "Thành công", "Đã cập nhật câu hỏi.")
                else:
                    QtWidgets.QMessageBox.critical(self, "Lỗi", "Không thể cập nhật câu hỏi.")
            else:
                # Tạo câu hỏi mới
                new_id = self.db.create_question(question_data)

                if new_id:
                    self.current_question_id = new_id
                    QtWidgets.QMessageBox.information(self, "Thành công", "Đã tạo câu hỏi mới.")
                else:
                    QtWidgets.QMessageBox.critical(self, "Lỗi", "Không thể tạo câu hỏi.")

            # Lưu tags nếu có
            self._save_question_tags()

            # Refresh danh sách
            self.on_tree_select()

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể lưu: {e}")

    def delete_question(self):
        """Xóa câu hỏi"""
        if not self.current_question_id:
            QtWidgets.QMessageBox.warning(self, "Chưa chọn", "Vui lòng chọn câu hỏi để xóa.")
            return

        if QtWidgets.QMessageBox.question(self, "Xác nhận",
                                          "Bạn có chắc muốn xóa câu hỏi này?") != QtWidgets.QMessageBox.Yes:
            return

        try:
            self.db.execute_query("DELETE FROM question_bank WHERE id=?", (self.current_question_id,))
            self.clear_question_form()

            tree_id = self._current_tree_id()
            if tree_id:
                rows = self.db.execute_query("SELECT * FROM question_bank WHERE tree_id=?", (tree_id,),
                                             fetch="all") or []
                self._load_question_rows(rows)

            QtWidgets.QMessageBox.information(self, "Thành công", "Đã xóa câu hỏi.")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể xóa: {e}")

    def new_question(self):
        """Tạo câu hỏi mới"""
        self.clear_question_form()
        self.content_text.setFocus()

    def clear_question_form(self):
        """Xóa form"""
        self.current_question_id = None
        if hasattr(self, 'content_text'):
            self.content_text.clear()
        if hasattr(self, 'answer_text'):
            self.answer_text.clear()
        if hasattr(self, 'tags_edit'):
            self.tags_edit.clear()

    def _current_tree_id(self) -> int | None:
        """Lấy tree_id hiện tại"""
        items = self.tree.selectedItems()
        if not items:
            return None
        return items[0].data(0, Qt.UserRole)

    def _save_question_tags(self):
        """Lưu tags cho câu hỏi"""
        if not self.current_question_id or not hasattr(self, 'tags_edit'):
            return

        tags_text = self.tags_edit.text().strip()
        if not tags_text:
            return

        # Xóa tags cũ
        self.db.execute_query("DELETE FROM question_tags WHERE question_id=?", (self.current_question_id,))

        # Thêm tags mới
        tag_names = [tag.strip() for tag in tags_text.split(',') if tag.strip()]
        for tag_name in tag_names:
            try:
                self.db.execute_query(
                    "INSERT INTO question_tags(question_id, tag_name) VALUES (?,?)",
                    (self.current_question_id, tag_name)
                )
            except:
                pass

    def _save_question_history(self, question_id, action_type, old_content, new_content):
        """Lưu lịch sử thay đổi"""
        try:
            self.db.execute_query(
                "INSERT INTO question_history(question_id, action_type, old_content, new_content) VALUES (?,?,?,?)",
                (question_id, action_type, old_content, new_content)
            )
        except:
            pass

    def _load_question_history(self, question_id):
        """Load lịch sử thay đổi"""
        if not hasattr(self, 'history_table'):
            return

        history = self.db.execute_query(
            "SELECT * FROM question_history WHERE question_id=? ORDER BY changed_date DESC LIMIT 50",
            (question_id,), fetch="all"
        ) or []

        self.history_table.setRowCount(0)

        for h in history:
            row_idx = self.history_table.rowCount()
            self.history_table.insertRow(row_idx)

            time_str = h.get("changed_date", "")

            self.history_table.setItem(row_idx, 0, QtWidgets.QTableWidgetItem(time_str))
            self.history_table.setItem(row_idx, 1, QtWidgets.QTableWidgetItem(h.get("action_type", "")))

            old_content = (h.get("old_content", "") or "")[:100]
            new_content = (h.get("new_content", "") or "")[:100]

            self.history_table.setItem(row_idx, 2, QtWidgets.QTableWidgetItem(old_content))
            self.history_table.setItem(row_idx, 3, QtWidgets.QTableWidgetItem(new_content))

    # CHỉnh sửa ca
    def edit_current_question(self):
        """Mở dialog edit câu hỏi hiện tại"""
        if not self.current_question_id:
            QtWidgets.QMessageBox.warning(self, "Chưa chọn câu hỏi",
                                          "Vui lòng chọn câu hỏi để chỉnh sửa")
            return

        tree_id = self._current_tree_id()
        dialog = QuestionEditDialog(self.db, tree_id=tree_id,
                                    question_id=self.current_question_id, parent=self)
        if dialog.exec() == QtWidgets.QDialog.Accepted:
            self.on_tree_select()

    # ========== NHÓM 6: TÌM KIẾM VÀ LỌC ========== #
    def search_questions(self):
        """Tìm kiếm câu hỏi"""
        keyword = (self.search_edit.text() or "").strip().lower()
        if not keyword:
            self.on_tree_select()
            return

        items = self.tree.selectedItems()
        if not items:
            QtWidgets.QMessageBox.warning(self, "Chưa chọn", "Hãy chọn thư mục để tìm.")
            return

        root_id = items[0].data(0, Qt.UserRole)
        all_ids = self.get_all_subtree_ids(root_id)
        if not all_ids:
            return

        placeholders = ",".join(["?"] * len(all_ids))
        query = f"SELECT * FROM question_bank WHERE tree_id IN ({placeholders})"
        rows = self.db.execute_query(query, tuple(all_ids), fetch="all") or []

        # Filter theo keyword
        rows = [r for r in rows if keyword in (r.get("content_text", "") or "").lower()]
        self._load_question_rows(rows)

    def get_all_subtree_ids(self, root_id: int) -> List[int]:
        """Lấy tất cả ID con"""
        ids = [root_id]
        children = self.db.execute_query("SELECT id FROM exercise_tree WHERE parent_id=?", (root_id,),
                                         fetch="all") or []
        for c in children:
            ids.extend(self.get_all_subtree_ids(c["id"]))
        return ids

    def focus_search(self):
        """Focus vào ô tìm kiếm"""
        self.search_edit.setFocus()
        self.search_edit.selectAll()

    def filter_by_combobox(self):
        """Lọc theo combobox"""
        # Simplified filtering logic
        self.on_tree_select()

    def clear_filters(self):
        """Xóa bộ lọc"""
        self.subject_cb.setCurrentIndex(0)
        self.grade_cb.setCurrentIndex(0)
        self.topic_cb.setCurrentIndex(0)
        self.type_cb.setCurrentIndex(0)
        self.level_cb.setCurrentIndex(0)
        self.on_tree_select()

    # ========== NHÓM 7: LOAD DỮ LIỆU COMBOBOX ========== #
    def load_available_subjects(self):
        """Load danh sách môn"""
        rows = self.db.execute_query(
            "SELECT DISTINCT name FROM exercise_tree WHERE level='Môn' ORDER BY name ASC",
            fetch="all"
        ) or []
        self.subject_cb.clear()
        self.subject_cb.addItem("")
        for r in rows:
            self.subject_cb.addItem(r["name"])

    def load_available_grades(self):
        """Load danh sách lớp"""
        rows = self.db.execute_query(
            "SELECT DISTINCT name FROM exercise_tree WHERE level='Lớp' ORDER BY name ASC",
            fetch="all"
        ) or []
        self.grade_cb.clear()
        self.grade_cb.addItem("")
        for r in rows:
            self.grade_cb.addItem(r["name"])

    def load_available_topics(self):
        """Load danh sách chủ đề"""
        subject = self.subject_cb.currentText().strip()
        grade = self.grade_cb.currentText().strip()

        if not subject or not grade:
            self.topic_cb.clear()
            self.type_cb.clear()
            return

        rows = self.db.execute_query("""
            SELECT name FROM exercise_tree 
            WHERE level='Chủ đề' AND parent_id IN (
                SELECT id FROM exercise_tree 
                WHERE name=? AND level='Lớp' AND parent_id IN (
                    SELECT id FROM exercise_tree WHERE name=? AND level='Môn'
                )
            )
        """, (grade, subject), fetch="all") or []

        self.topic_cb.clear()
        self.topic_cb.addItem("")
        for r in rows:
            self.topic_cb.addItem(r["name"])

    def load_available_types(self):
        """Load danh sách dạng"""
        topic = self.topic_cb.currentText().strip()
        if not topic:
            self.type_cb.clear()
            return

        rows = self.db.execute_query("""
            SELECT name FROM exercise_tree
            WHERE level='Dạng' AND parent_id IN (
                SELECT id FROM exercise_tree WHERE level='Chủ đề' AND name=?
            )
        """, (topic,), fetch="all") or []

        self.type_cb.clear()
        self.type_cb.addItem("")
        for r in rows:
            self.type_cb.addItem(r["name"])

    # ========== NHÓM 8: IMPORT/EXPORT ========== #
    def import_from_word(self):
        """Import từ Word"""
        try:
            from docx import Document
        except ImportError:
            QtWidgets.QMessageBox.critical(self, "Lỗi", "Cần cài đặt python-docx: pip install python-docx")
            return

        file_path, _ = QtWidgets.QFileDialog.getOpenFileName(
            self, "Chọn file Word", "", "Word files (*.docx)")
        if not file_path:
            return

        tree_id = self._current_tree_id()
        if not tree_id:
            QtWidgets.QMessageBox.warning(self, "Thiếu thư mục", "Vui lòng chọn nơi lưu câu hỏi.")
            return

        try:
            doc = Document(file_path)
            count = 0

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Simple import - just save as content
                    self.db.execute_query(
                        "INSERT INTO question_bank(content_text, tree_id) VALUES (?,?)",
                        (text, tree_id)
                    )
                    count += 1

            # Reload
            rows = self.db.execute_query("SELECT * FROM question_bank WHERE tree_id=?", (tree_id,), fetch="all") or []
            self._load_question_rows(rows)

            QtWidgets.QMessageBox.information(self, "Thành công", f"Đã import {count} câu hỏi.")

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể import: {e}")

    def export_to_word(self):
        """Export ra Word"""
        tree_id = self._current_tree_id()
        if not tree_id:
            QtWidgets.QMessageBox.warning(self, "Chưa chọn thư mục", "Vui lòng chọn thư mục để xuất.")
            return

        file_path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Lưu file Word", "", "Word files (*.docx)")
        if not file_path:
            return

        try:
            from docx import Document

            doc = Document()
            doc.add_heading('NGÂN HÀNG CÂU HỎI', 0)

            # Thông tin đường dẫn
            path_info = self.get_tree_path(tree_id)
            if path_info:
                path_text = " > ".join([p["name"] for p in path_info])
                doc.add_paragraph(f"Đường dẫn: {path_text}")

            # Lấy câu hỏi
            rows = self.db.execute_query("SELECT * FROM question_bank WHERE tree_id=?", (tree_id,), fetch="all") or []

            for i, row in enumerate(rows, 1):
                doc.add_paragraph(f"Câu {i}: {row.get('content_text', '')}", style='Heading 3')

                if row.get('correct'):
                    doc.add_paragraph(f"Đáp án: {row['correct']}")

                doc.add_paragraph("")  # Dòng trống

            doc.save(file_path)
            QtWidgets.QMessageBox.information(self, "Thành công", f"Đã xuất {len(rows)} câu hỏi ra file Word.")

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể xuất file: {e}")

    def export_to_pdf(self):
        """Export ra PDF"""
        tree_id = self._current_tree_id()
        if not tree_id:
            QtWidgets.QMessageBox.warning(self, "Chưa chọn thư mục", "Vui lòng chọn thư mục để xuất.")
            return

        file_path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Lưu file PDF", "", "PDF files (*.pdf)")
        if not file_path:
            return

        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch

            doc = SimpleDocTemplate(file_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []

            # Tiêu đề
            title = Paragraph("NGÂN HÀNG CÂU HỎI", styles['Title'])
            story.append(title)
            story.append(Spacer(1, 0.2 * inch))

            # Lấy câu hỏi
            rows = self.db.execute_query("SELECT * FROM question_bank WHERE tree_id=?", (tree_id,), fetch="all") or []

            for i, row in enumerate(rows, 1):
                question_para = Paragraph(f"<b>Câu {i}:</b> {row.get('content_text', '')}", styles['Normal'])
                story.append(question_para)

                if row.get('correct'):
                    answer_para = Paragraph(f"<b>Đáp án:</b> {row['correct']}", styles['Normal'])
                    story.append(answer_para)

                story.append(Spacer(1, 0.2 * inch))

            doc.build(story)
            QtWidgets.QMessageBox.information(self, "Thành công", f"Đã xuất {len(rows)} câu hỏi ra file PDF.")

        except ImportError:
            QtWidgets.QMessageBox.critical(self, "Lỗi", "Cần cài đặt reportlab: pip install reportlab")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể xuất PDF: {e}")

    # ========== NHÓM 9: TEMPLATE ========== #
    def show_template_dialog(self):
        """Hiển thị dialog template"""
        QtWidgets.QMessageBox.information(self, "Thông tin", "Chức năng template đang phát triển.")

    # ========== NHÓM 10: TAGS ========== #
    def add_new_tag(self):
        """Thêm tag mới"""
        if not hasattr(self, 'tags_edit') or not self.current_question_id:
            return

        tags_text = self.tags_edit.text().strip()
        if not tags_text:
            return

        tag_names = [tag.strip() for tag in tags_text.split(',') if tag.strip()]

        added_count = 0
        for tag_name in tag_names:
            try:
                self.db.execute_query(
                    "INSERT INTO question_tags(question_id, tag_name) VALUES (?,?)",
                    (self.current_question_id, tag_name)
                )
                added_count += 1
            except:
                pass

        if added_count > 0:
            self.tags_edit.clear()
            tree_id = self._current_tree_id()
            if tree_id:
                rows = self.db.execute_query("SELECT * FROM question_bank WHERE tree_id=?", (tree_id,),
                                             fetch="all") or []
                self._load_question_rows(rows)

    # ========== NHÓM 11: PREVIEW VÀ THỐNG KÊ ========== #
    def update_preview(self):
        """Cập nhật preview"""
        if not hasattr(self, 'preview_widget'):
            return

        content = self.content_text.toPlainText() if hasattr(self, 'content_text') else ""
        answer = self.answer_text.toPlainText() if hasattr(self, 'answer_text') else ""

        html = f"""
        <div style="font-family: Arial, sans-serif;">
            <h3 style="color: #2c3e50;">📝 Câu hỏi</h3>
            <p style="background: #f8f9fa; padding: 15px; border-left: 4px solid #007bff;">
                {content or '<em>Chưa có nội dung...</em>'}
            </p>
            <h4 style="color: #2c3e50;">✅ Đáp án</h4>
            <p style="background: #e8f5e9; padding: 15px;">
                {answer or '<em>Chưa có đáp án...</em>'}
            </p>
        </div>
        """
        self.preview_widget.setHtml(html)

    def update_statistics(self):
        """Cập nhật thống kê"""
        if not hasattr(self, 'stats_widget'):
            return

        try:
            # Tổng số câu hỏi
            total = self.db.execute_query("SELECT COUNT(*) as count FROM question_bank", fetch="one")["count"]

            # Thống kê theo mức độ
            level_stats = self.db.execute_query("""
                SELECT e.name, COUNT(q.id) as count 
                FROM question_bank q 
                JOIN exercise_tree e ON e.id = q.tree_id 
                WHERE e.level = 'Mức độ'
                GROUP BY e.name
                ORDER BY count DESC
            """, fetch="all") or []

            # Tạo HTML thống kê
            stats_html = f"""
            <div style="font-family: Arial, sans-serif;">
                <h2 style="color: #2c3e50;">📊 Thống kê</h2>
                <p><strong>Tổng số câu hỏi:</strong> {total}</p>

                <h3>Phân bố theo mức độ:</h3>
                <ul>
            """

            for stat in level_stats:
                percentage = (stat["count"] / total * 100) if total > 0 else 0
                stats_html += f"<li>{stat['name']}: {stat['count']} ({percentage:.1f}%)</li>"

            stats_html += "</ul></div>"

            self.stats_widget.setHtml(stats_html)

        except Exception as e:
            self.stats_widget.setHtml(f"<p style='color: red;'>Lỗi: {e}</p>")

    def refresh_all(self):
        """Làm mới toàn bộ"""
        self.refresh_tree()
        self.load_available_subjects()
        self.load_available_grades()
        self.on_tree_select()
        self.update_statistics()

    # ========== NHÓM HELPER CHO TÍNH NĂNG MỚI ========== #

    def view_question_detail(self):
        """Xem chi tiết câu hỏi trong dialog riêng"""
        if not self.current_question_id:
            return

        dialog = QuestionDetailDialog(self.db, self.current_question_id, parent=self)
        dialog.exec()

    def move_question_to_folder(self):
        """Di chuyển câu hỏi đến thư mục khác"""
        if not self.current_question_id:
            return

        # Tạo dialog chọn thư mục đích
        dialog = FolderSelectDialog(self.db, parent=self)
        if dialog.exec() == QtWidgets.QDialog.Accepted:
            target_tree_id = dialog.selected_tree_id
            if target_tree_id:
                try:
                    self.db.execute_query(
                        "UPDATE question_bank SET tree_id=?, modified_date=CURRENT_TIMESTAMP WHERE id=?",
                        (target_tree_id, self.current_question_id)
                    )
                    QtWidgets.QMessageBox.information(self, "Thành công", "Đã di chuyển câu hỏi.")
                    self.on_tree_select()  # Refresh danh sách
                except Exception as e:
                    QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể di chuyển: {e}")

    def manage_question_tags(self):
        """Quản lý tags của câu hỏi"""
        if not self.current_question_id:
            return

        dialog = TagsManagerDialog(self.db, self.current_question_id, parent=self)
        if dialog.exec() == QtWidgets.QDialog.Accepted:
            self.on_tree_select()  # Refresh để cập nhật tags

    def copy_question(self):
        """Sao chép câu hỏi"""
        if not self.current_question_id:
            return

        # Lấy dữ liệu câu hỏi
        q = self.db.execute_query("SELECT * FROM question_bank WHERE id=?", (self.current_question_id,), fetch="one")
        if not q:
            return

        # Copy vào clipboard (định dạng JSON)
        import json
        question_data = {
            "content_text": self.safe_get("content_text", ""),
            "options": self.safe_get("options", ""),
            "correct": self.safe_get("correct", ""),
            "content_type": self.safe_get("content_type", "text"),
            "difficulty_level": self.safe_get("difficulty_level", "")
        }

        from PySide6.QtGui import QClipboard
        clipboard = QtWidgets.QApplication.clipboard()
        clipboard.setText(json.dumps(question_data, ensure_ascii=False, indent=2))

        QtWidgets.QMessageBox.information(self, "Thành công", "Đã sao chép câu hỏi vào clipboard.")

    def open_fullscreen_preview(self):
        """Mở ảnh preview fullscreen"""
        # #(Phương thức xem ảnh preview ở chế độ toàn màn hình với kiểm tra an toàn)
        if not hasattr(self, 'preview_image'):
            QtWidgets.QMessageBox.information(self, "Thông báo", "Không có preview image để hiển thị")
            return

        if not hasattr(self.preview_image, 'current_pixmap') or not self.preview_image.current_pixmap:
            QtWidgets.QMessageBox.information(self, "Thông báo", "Không có ảnh nào để xem fullscreen")
            return

        if self.preview_image.current_pixmap.isNull():
            QtWidgets.QMessageBox.information(self, "Thông báo", "Ảnh preview không hợp lệ")
            return

        try:
            dialog = ImageViewerDialog(self.preview_image.current_pixmap, self)
            dialog.exec()
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể mở ảnh fullscreen: {e}")

    def setup_preview_interactions(self):
        """Thiết lập tương tác cho preview"""
        # #(Thiết lập sự kiện double-click và context menu cho preview image với cursor phù hợp)
        if hasattr(self, 'preview_image'):
            # Double-click để xem fullscreen
            def handle_double_click(event):
                # Kiểm tra có ảnh không trước khi mở fullscreen
                if (hasattr(self.preview_image, 'current_pixmap') and
                        self.preview_image.current_pixmap and
                        not self.preview_image.current_pixmap.isNull()):
                    self.open_fullscreen_preview()

            self.preview_image.mouseDoubleClickEvent = handle_double_click

            # Right-click menu
            self.preview_image.setContextMenuPolicy(Qt.CustomContextMenu)
            self.preview_image.customContextMenuRequested.connect(self.show_preview_context_menu)

            # Cập nhật tooltip và cursor
            self.preview_image.setToolTip(
                "🖱️ Double-click để xem fullscreen\n"
                "🎮 Dùng nút +/- để zoom trong preview\n"
                "📱 Chuột phải để xem menu"
            )

            # Set cursor để người dùng biết có thể click
            if hasattr(self.preview_image, 'image_label'):
                self.preview_image.image_label.setCursor(Qt.PointingHandCursor)

    def show_preview_context_menu(self, position):
        """Context menu cho ảnh preview"""
        # #(Menu chuột phải cho preview với kiểm tra an toàn)
        if not hasattr(self, 'preview_image'):
            return

        menu = QtWidgets.QMenu(self)

        fullscreen_action = menu.addAction("🔍 Xem fullscreen")
        fullscreen_action.triggered.connect(self.open_fullscreen_preview)

        # Chỉ thêm zoom menu nếu có ảnh
        if hasattr(self.preview_image, 'current_pixmap') and self.preview_image.current_pixmap:
            zoom_in_action = menu.addAction("🔍+ Phóng to")
            zoom_in_action.triggered.connect(self.preview_image._zoom_in)

            zoom_out_action = menu.addAction("🔍- Thu nhỏ")
            zoom_out_action.triggered.connect(self.preview_image._zoom_out)

        menu.exec(self.preview_image.mapToGlobal(position))  # ========== XỬ LÝ ẢNH TRONG PREVIEW ========== #

    def load_image_from_data(self, content_data, content_metadata=None):
        """Load ảnh từ content_data với nhiều format"""
        if not content_data:
            return None

        try:
            pixmap = None

            # Nếu là bytes/BLOB data
            if isinstance(content_data, (bytes, bytearray)):
                # Thử load trực tiếp từ bytes
                pixmap = QtGui.QPixmap()
                if pixmap.loadFromData(content_data):
                    return pixmap

                # Thử decode base64
                try:
                    import base64
                    decoded_data = base64.b64decode(content_data)
                    pixmap = QtGui.QPixmap()
                    if pixmap.loadFromData(decoded_data):
                        return pixmap
                except:
                    pass

            # Nếu là string (có thể là base64 hoặc file path)
            elif isinstance(content_data, str):
                # Thử như file path
                if os.path.exists(content_data):
                    pixmap = QtGui.QPixmap(content_data)
                    if not pixmap.isNull():
                        return pixmap

                # Thử decode base64
                try:
                    import base64
                    if content_data.startswith('data:image'):
                        # Data URL format
                        header, data = content_data.split(',', 1)
                        decoded_data = base64.b64decode(data)
                    else:
                        # Pure base64
                        decoded_data = base64.b64decode(content_data)

                    pixmap = QtGui.QPixmap()
                    if pixmap.loadFromData(decoded_data):
                        return pixmap
                except:
                    pass

            return None

        except Exception as e:
            print(f"⚠️ Lỗi load image: {e}")
            return None

    def display_image_preview(self, pixmap):
        """Hiển thị thông tin ảnh đơn giản - không render ảnh thực tế"""
        try:
            if pixmap and not pixmap.isNull():
                image_info = (
                    f"🖼️ [Câu hỏi có hình ảnh]\n\n"
                    f"📏 Kích thước: {pixmap.width()}x{pixmap.height()}\n"
                    f"🎨 Format: {pixmap.hasAlpha()}"
                )
                self.show_text_preview(image_info)
            else:
                self.show_text_preview("❌ Không thể load ảnh")
        except Exception as e:
            print(f"⚠️ Lỗi display image: {e}")
            self.show_text_preview("🖼️ [Câu hỏi có hình ảnh]")

    def show_text_preview(self, text):
        """Hiển thị text preview đơn giản - không dùng widget phức tạp"""
        try:
            if hasattr(self, 'preview_content'):
                self.preview_content.setText(str(text))
        except Exception as e:
            print(f"⚠️ Lỗi show_text_preview: {e}")

    # ========== VALIDATION HELPER ========== #
    def validate_question_data(self, question_data):
        """Validate dữ liệu câu hỏi trước khi xử lý"""
        if not question_data:
            return False, "Dữ liệu câu hỏi trống"

        required_fields = ["content_text", "tree_id"]
        for field in required_fields:
            if not self.safe_get(question_data, field):
                return False, f"Thiếu trường bắt buộc: {field}"

        return True, ""

    # ========== HELPER KIỂM TRA WIDGET HỢP LỆ ========== #
    # ========== WRAPPER AN TOÀN CHO TRUY VẤN DATABASE ========== #
    def safe_db_query(self, query, params=(), fetch=None):
        """Wrapper an toàn cho database query"""
        try:
            result = self.db.execute_query(query, params, fetch)

            # Chuyển đổi kết quả nếu cần
            if fetch == "one" and result:
                return self.row_to_dict(result)
            elif fetch == "all" and result:
                return [self.row_to_dict(row) for row in result]

            return result
        except Exception as e:
            print(f"⚠️ Lỗi database query: {e}")
            return None if fetch == "one" else []


class ImageViewer(QtWidgets.QWidget):
    """Widget hiển thị và xử lý ảnh"""

    def __init__(self):
        super().__init__()
        layout = QtWidgets.QVBoxLayout(self)

        # Scroll area để hiển thị ảnh lớn
        scroll = QtWidgets.QScrollArea()
        self.image_label = QtWidgets.QLabel()
        self.image_label.setAlignment(Qt.AlignCenter)
        self.image_label.setScaledContents(False)  # Giữ tỷ lệ ảnh
        scroll.setWidget(self.image_label)
        scroll.setWidgetResizable(True)
        layout.addWidget(scroll)

        # Toolbar zoom
        toolbar = QtWidgets.QToolBar()
        toolbar.addAction("🔍+").triggered.connect(lambda: self.zoom(1.2))
        toolbar.addAction("🔍-").triggered.connect(lambda: self.zoom(0.8))
        toolbar.addAction("💯").triggered.connect(self.fit_to_window)
        toolbar.addAction("📋 Paste").triggered.connect(self.paste_image)
        layout.addWidget(toolbar)

        self.current_pixmap = None
        self.scale_factor = 1.0

        # Enable drop
        self.setAcceptDrops(True)

    def set_image(self, image):
        """Set QImage từ clipboard"""
        if isinstance(image, QtGui.QImage):
            self.current_pixmap = QtGui.QPixmap.fromImage(image)
            self.display_image()
            return True
        return False

    def set_pixmap(self, pixmap):
        """Set QPixmap"""
        if isinstance(pixmap, QtGui.QPixmap):
            self.current_pixmap = pixmap
            self.display_image()
            return True
        return False

    def display_image(self):
        if self.current_pixmap and not self.current_pixmap.isNull():
            # Scale ảnh nhưng giữ tỷ lệ
            scaled = self.current_pixmap.scaled(
                self.current_pixmap.size() * self.scale_factor,
                Qt.KeepAspectRatio,
                Qt.SmoothTransformation
            )
            self.image_label.setPixmap(scaled)

    def paste_image(self):
        """Paste ảnh trực tiếp trong viewer"""
        clipboard = QtWidgets.QApplication.clipboard()
        mime_data = clipboard.mimeData()

        if mime_data.hasImage():
            image = clipboard.image()
            if not image.isNull():
                self.set_image(image)
                print("✅ Đã paste ảnh vào viewer")
        else:
            print("❌ Không có ảnh trong clipboard")

    def dragEnterEvent(self, event):
        """Xử lý kéo thả file"""
        if event.mimeData().hasImage() or event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dropEvent(self, event):
        """Xử lý thả file"""
        mime_data = event.mimeData()

        if mime_data.hasImage():
            self.set_image(mime_data.imageData())
        elif mime_data.hasUrls():
            for url in mime_data.urls():
                file_path = url.toLocalFile()
                if file_path.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                    pixmap = QtGui.QPixmap(file_path)
                    if not pixmap.isNull():
                        self.set_pixmap(pixmap)
                        break

    # Thêm vào class ImageViewer hiện tại của bạn:

    def fit_to_window(self):
        """Reset zoom về 100%"""
        self.scale_factor = 1.0
        self.display_image()

    def zoom(self, factor):
        """Zoom ảnh"""
        if not self.current_pixmap:
            return

        self.scale_factor *= factor
        # Giới hạn zoom từ 10% đến 500%
        self.scale_factor = max(0.1, min(5.0, self.scale_factor))
        self.display_image()

    def paste_image(self):
        """Paste ảnh từ clipboard"""
        clipboard = QtWidgets.QApplication.clipboard()
        if clipboard.mimeData().hasImage():
            image = clipboard.image()
            if not image.isNull():
                self.set_image(image)

    def clear_image(self):
        """Xóa ảnh"""
        self.current_pixmap = None
        self.image_label.clear()




class LaTeXInputDialog(QtWidgets.QDialog):
    """Dialog nhập công thức LaTeX"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("∑ Nhập công thức LaTeX")
        self.resize(500, 300)

        layout = QtWidgets.QVBoxLayout(self)

        layout.addWidget(QtWidgets.QLabel("Nhập công thức LaTeX:"))

        self.latex_edit = QtWidgets.QTextEdit()
        self.latex_edit.setPlaceholderText(r"\frac{a}{b} hoặc \int_{0}^{1} x^2 dx")
        layout.addWidget(self.latex_edit)

        # Preview (nếu có thư viện render LaTeX)
        self.preview_label = QtWidgets.QLabel("Preview sẽ hiển thị ở đây")
        self.preview_label.setAlignment(Qt.AlignCenter)
        self.preview_label.setStyleSheet("border: 1px solid #ccc; padding: 10px;")
        layout.addWidget(self.preview_label)

        # Buttons
        button_layout = QtWidgets.QHBoxLayout()
        ok_btn = QtWidgets.QPushButton("OK")
        ok_btn.clicked.connect(self.accept)
        cancel_btn = QtWidgets.QPushButton("Hủy")
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(cancel_btn)
        button_layout.addWidget(ok_btn)
        layout.addLayout(button_layout)

    def get_latex(self):
        return self.latex_edit.toPlainText()




# ========== DIALOG XEM CHI TIẾT CÂU HỎI ========== #


# ========== CLASS MỚI - THÊM VÀO ========== #
class AdaptiveImageViewer(QtWidgets.QWidget):
    """Widget hiển thị ảnh tự động điều chỉnh kích thước theo ảnh"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.current_pixmap = None
        self.max_width = 600  # Chiều rộng tối đa
        self.max_height = 400  # Chiều cao tối đa
        self.min_height = 100  # Chiều cao tối thiểu

        self._setup_ui()

    # #(Xử lý tự động điều chỉnh khi resize widget)
    def resizeEvent(self, event):
        """Tự động điều chỉnh ảnh khi widget thay đổi kích thước"""
        super().resizeEvent(event)
        if hasattr(self, 'current_pixmap') and self.current_pixmap:
            QtCore.QTimer.singleShot(50, self.fit_to_container)  # Delay nhỏ để đảm bảo UI đã update

    def _setup_ui(self):
        """Thiết lập giao diện"""
        layout = QtWidgets.QVBoxLayout(self)
        layout.setContentsMargins(5, 5, 5, 5)

        # Label hiển thị ảnh
        self.image_label = QtWidgets.QLabel()
        self.image_label.setAlignment(Qt.AlignCenter)
        self.image_label.setStyleSheet("""
            QLabel {
                border: 1px solid #ddd;
                border-radius: 4px;
                background: white;
                padding: 5px;
            }
        """)

        # Scroll area cho ảnh lớn
        self.scroll_area = QtWidgets.QScrollArea()
        self.scroll_area.setWidget(self.image_label)
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)

        layout.addWidget(self.scroll_area)

        # Info label hiển thị kích thước ảnh
        self.info_label = QtWidgets.QLabel("No image loaded")
        self.info_label.setStyleSheet("color: #666; font-size: 11px;")
        layout.addWidget(self.info_label)

    def load_image_from_data(self, image_data):
        """Load ảnh từ binary data và tự động resize widget"""
        if not image_data:
            self.clear_image()
            return

        try:
            # Load pixmap từ data
            pixmap = QtGui.QPixmap()
            if isinstance(image_data, bytes):
                pixmap.loadFromData(image_data)
            elif isinstance(image_data, str):
                # Thử decode base64 nếu là string
                import base64
                try:
                    if image_data.startswith('data:image'):
                        header, data = image_data.split(',', 1)
                        decoded_data = base64.b64decode(data)
                    else:
                        decoded_data = base64.b64decode(image_data)
                    pixmap.loadFromData(decoded_data)
                except:
                    pixmap = QtGui.QPixmap(image_data)  # File path

            if pixmap.isNull():
                self.clear_image()
                return

            self.current_pixmap = pixmap
            QtCore.QTimer.singleShot(100, self.fit_to_container)

        except Exception as e:
            print(f"⚠️ Lỗi load image: {e}")
            self.clear_image()

    def _display_adaptive_image(self):
        """Hiển thị ảnh với kích thước tự động điều chỉnh"""
        if not self.current_pixmap or self.current_pixmap.isNull():
            return

        # Kích thước gốc của ảnh
        original_width = self.current_pixmap.width()
        original_height = self.current_pixmap.height()

        # Tính toán kích thước hiển thị
        display_width = min(original_width, self.max_width)
        display_height = min(original_height, self.max_height)

        # Giữ tỷ lệ khung hình
        if original_width > 0 and original_height > 0:
            aspect_ratio = original_width / original_height

            # Tính scale factor để vừa với giới hạn max_width và max_height
            scale_width = self.max_width / original_width
            scale_height = self.max_height / original_height
            scale_factor = min(scale_width, scale_height, 1.0)  # Không phóng to quá kích thước gốc

            # Áp dụng scale factor
            display_width = int(original_width * scale_factor)
            display_height = int(original_height * scale_factor)
        else:
            display_width = self.max_width
            display_height = self.max_height

        # Scale ảnh theo kích thước tính toán
        scaled_pixmap = self.current_pixmap.scaled(
            display_width, display_height,
            Qt.KeepAspectRatio, Qt.SmoothTransformation
        )

        # Hiển thị ảnh
        self.image_label.setPixmap(scaled_pixmap)
        self.image_label.setMinimumSize(display_width, display_height)

        # Điều chỉnh kích thước widget chứa
        widget_height = display_height + 80  # +80 cho info label và padding
        self.setMinimumHeight(widget_height)
        self.setMaximumHeight(max(widget_height, 400))  # Tối thiểu 400px

        # Cập nhật thông tin
        scale_percent = int((display_width / original_width) * 100) if original_width > 0 else 100
        self.info_label.setText(
            f"📷 Gốc: {original_width}×{original_height} | "
            f"Hiển thị: {display_width}×{display_height} | "
            f"Tỷ lệ: {scale_percent}%"
        )

        # #(Cập nhật layout parent để hiển thị thay đổi)
        if self.parent():
            self.parent().updateGeometry()

    def clear_image(self):
        """Xóa ảnh và reset kích thước"""
        self.current_pixmap = None
        self.image_label.clear()
        self.image_label.setText("No image")
        self.setMinimumHeight(self.min_height)
        self.setMaximumHeight(self.min_height)
        self.info_label.setText("No image loaded")

    # ========== PHƯƠNG THỨC CONFIGURATION - THÊM VÀO AdaptiveImageViewer ========== #
    def set_size_limits(self, max_width=600, max_height=400, min_height=100):
        """Thiết lập giới hạn kích thước hiển thị"""
        self.max_width = max_width
        self.max_height = max_height
        self.min_height = min_height

        # Refresh hiển thị nếu đã có ảnh
        if self.current_pixmap:
            self._display_adaptive_image()

    def enable_zoom_controls(self):
        """Thêm nút zoom cho ảnh lớn"""
        zoom_layout = QtWidgets.QHBoxLayout()

        zoom_out_btn = QtWidgets.QPushButton("🔍-")
        zoom_out_btn.setFixedSize(30, 25)
        zoom_out_btn.clicked.connect(self._zoom_out)

        zoom_in_btn = QtWidgets.QPushButton("🔍+")
        zoom_in_btn.setFixedSize(30, 25)
        zoom_in_btn.clicked.connect(self._zoom_in)

        zoom_layout.addWidget(zoom_out_btn)
        zoom_layout.addWidget(zoom_in_btn)
        zoom_layout.addStretch()

        # Thêm vào layout chính (sau info_label)
        self.layout().addLayout(zoom_layout)

    def _zoom_in(self):
        """Phóng to ảnh"""
        # #(Tăng kích thước giới hạn hiển thị để tạo hiệu ứng zoom)
        if self.max_width < 1200 and self.max_height < 1000:
            self.max_width = min(self.max_width + 100, 1200)
            self.max_height = min(self.max_height + 80, 1000)
            self._display_adaptive_image()
            print(f"🔍+ Zoom in: {self.max_width}×{self.max_height}")

    def _zoom_out(self):
        """Thu nhỏ ảnh"""
        # #(Giảm kích thước giới hạn hiển thị để tạo hiệu ứng zoom out)
        if self.max_width > 200 and self.max_height > 150:
            self.max_width = max(self.max_width - 100, 200)
            self.max_height = max(self.max_height - 80, 150)
            self._display_adaptive_image()
            print(f"🔍- Zoom out: {self.max_width}×{self.max_height}")

    def fit_to_container(self):
        """Điều chỉnh kích thước tối đa theo container thực tế"""
        if not self.current_pixmap or self.current_pixmap.isNull():
            return

        try:
            # Lấy kích thước container thực tế
            container_width = self.width() - 20  # Trừ margin
            container_height = self.height() - 80  # Trừ info label và padding

            # Đảm bảo kích thước tối thiểu
            container_width = max(container_width, 200)
            container_height = max(container_height, self.min_height)

            # Cập nhật giới hạn
            self.max_width = container_width
            self.max_height = container_height

            # Refresh hiển thị
            self._display_adaptive_image()
        except Exception as e:
            print(f"⚠️ Lỗi fit_to_container: {e}")  # ========== DIALOG XEM CHI TIẾT CÂU HỎI ========== #


# ========== DIALOG CHỌN THƯ MỤC ========== #
class FolderSelectDialog(QtWidgets.QDialog):
    """Dialog chọn thư mục đích để di chuyển câu hỏi"""

    def __init__(self, db_manager, parent=None):
        super().__init__(parent)
        self.db = db_manager
        self.selected_tree_id = None

        self.setWindowTitle("🗂️ Chọn thư mục đích")
        self.setModal(True)
        self.resize(400, 500)

        self._setup_ui()
        self._load_tree_data()

    def _setup_ui(self):
        """Thiết lập giao diện dialog"""
        layout = QtWidgets.QVBoxLayout(self)

        # Tiêu đề
        title = QtWidgets.QLabel("Chọn thư mục để di chuyển câu hỏi:")
        title.setStyleSheet("font-weight: bold; font-size: 14px; margin-bottom: 10px;")
        layout.addWidget(title)

        # Tree widget để hiển thị cây thư mục
        self.tree_widget = QtWidgets.QTreeWidget()
        self.tree_widget.setHeaderHidden(True)
        self.tree_widget.itemClicked.connect(self._on_tree_select)
        self.tree_widget.itemDoubleClicked.connect(self._on_tree_double_click)

        # Style cho tree
        self.tree_widget.setStyleSheet("""
            QTreeWidget {
                border: 1px solid #ddd;
                border-radius: 4px;
                background: white;
            }
            QTreeWidget::item {
                padding: 8px;
                border-bottom: 1px solid #f0f0f0;
            }
            QTreeWidget::item:selected {
                background-color: #e3f2fd;
                color: #1976d2;
            }
            QTreeWidget::item:hover {
                background-color: #f5f5f5;
            }
        """)

        layout.addWidget(self.tree_widget)

        # Label hiển thị lựa chọn hiện tại
        self.selection_label = QtWidgets.QLabel("Chưa chọn thư mục nào")
        self.selection_label.setStyleSheet("""
            QLabel {
                padding: 8px;
                background: #f8f9fa;
                border: 1px solid #dee2e6;
                border-radius: 4px;
                color: #495057;
            }
        """)
        layout.addWidget(self.selection_label)

        # Buttons
        button_layout = QtWidgets.QHBoxLayout()

        self.ok_button = QtWidgets.QPushButton("✅ Xác nhận")
        self.ok_button.setEnabled(False)
        self.ok_button.clicked.connect(self.accept)
        self.ok_button.setStyleSheet("""
            QPushButton {
                background-color: #28a745;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #218838;
            }
            QPushButton:disabled {
                background-color: #6c757d;
            }
        """)

        cancel_button = QtWidgets.QPushButton("❌ Hủy")
        cancel_button.clicked.connect(self.reject)
        cancel_button.setStyleSheet("""
            QPushButton {
                background-color: #6c757d;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #545b62;
            }
        """)

        button_layout.addStretch()
        button_layout.addWidget(cancel_button)
        button_layout.addWidget(self.ok_button)

        layout.addLayout(button_layout)

    def _load_tree_data(self):
        """Load dữ liệu cây thư mục"""
        try:
            # Load tất cả nodes
            nodes = self.db.execute_query(
                "SELECT id, parent_id, name, level FROM exercise_tree ORDER BY parent_id, name",
                fetch="all"
            ) or []

            # Chuyển đổi sang dict để dễ xử lý
            node_dict = {}
            root_nodes = []

            for node in nodes:
                node_data = dict(node) if hasattr(node, 'keys') else {}
                node_id = node_data.get('id')
                parent_id = node_data.get('parent_id')

                if node_id:
                    node_dict[node_id] = node_data
                    if not parent_id:
                        root_nodes.append(node_data)

            # Tạo tree structure
            self.tree_widget.clear()

            for root_node in root_nodes:
                root_item = self._create_tree_item(root_node)
                self.tree_widget.addTopLevelItem(root_item)
                self._add_child_items(root_item, root_node['id'], node_dict)

            # Expand tất cả các node
            self.tree_widget.expandAll()

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể load cây thư mục: {e}")

    def _create_tree_item(self, node_data):
        """Tạo QTreeWidgetItem từ node data"""
        name = node_data.get('name', 'Unknown')
        level = node_data.get('level', '')
        node_id = node_data.get('id')

        # Icon theo level
        level_icons = {
            'subject': '📚',
            'grade': '🎓',
            'topic': '📖',
            'subtopic': '📄',
            'difficulty': '🎯'
        }

        icon = level_icons.get(level.lower(), '📁')
        display_text = f"{icon} {name}"

        item = QtWidgets.QTreeWidgetItem([display_text])
        item.setData(0, Qt.UserRole, node_id)

        return item

    def _add_child_items(self, parent_item, parent_id, node_dict):
        """Thêm các item con vào parent item"""
        children = [node for node in node_dict.values() if node.get('parent_id') == parent_id]

        for child in children:
            child_item = self._create_tree_item(child)
            parent_item.addChild(child_item)

            # Recursively add grandchildren
            child_id = child.get('id')
            if child_id:
                self._add_child_items(child_item, child_id, node_dict)

    def _on_tree_select(self, item, column):
        """Xử lý khi chọn item trong tree"""
        if not item:
            return

        tree_id = item.data(0, Qt.UserRole)
        if tree_id:
            self.selected_tree_id = tree_id

            # Lấy đường dẫn đầy đủ
            path = self._get_item_path(item)
            self.selection_label.setText(f"📂 Đã chọn: {path}")
            self.selection_label.setStyleSheet("""
                QLabel {
                    padding: 8px;
                    background: #d4edda;
                    border: 1px solid #c3e6cb;
                    border-radius: 4px;
                    color: #155724;
                }
            """)

            self.ok_button.setEnabled(True)
        else:
            self.selected_tree_id = None
            self.selection_label.setText("Chưa chọn thư mục nào")
            self.ok_button.setEnabled(False)

    def _on_tree_double_click(self, item, column):
        """Xử lý double click - chọn và đóng dialog"""
        self._on_tree_select(item, column)
        if self.selected_tree_id:
            self.accept()

    def _get_item_path(self, item):
        """Lấy đường dẫn đầy đủ của item"""
        path_parts = []
        current_item = item

        while current_item:
            text = current_item.text(0)
            # Loại bỏ icon emoji khỏi text
            clean_text = ''.join(char for char in text if not self._is_emoji(char)).strip()
            path_parts.insert(0, clean_text)
            current_item = current_item.parent()

        return ' > '.join(path_parts)

    def _is_emoji(self, char):
        """Kiểm tra ký tự có phải emoji không"""
        return ord(char) > 0x1F000


# ========== DIALOG QUẢN LÝ TAGS CÂU HỎI ========== #
class TagsManagerDialog(QtWidgets.QDialog):
    """Dialog quản lý tags cho câu hỏi"""

    def __init__(self, db_manager, question_id, parent=None):
        super().__init__(parent)
        self.db = db_manager
        self.question_id = question_id

        self.setWindowTitle(f"🏷️ Quản lý Tags - Câu hỏi #{question_id}")
        self.setModal(True)
        self.resize(500, 400)

        self._setup_ui()
        self._load_current_tags()
        self._load_available_tags()

    def _setup_ui(self):
        """Thiết lập giao diện dialog"""
        layout = QtWidgets.QVBoxLayout(self)

        # Tiêu đề
        title = QtWidgets.QLabel("Quản lý Tags cho câu hỏi")
        title.setStyleSheet("font-weight: bold; font-size: 16px; margin-bottom: 15px;")
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)

        # Phần thêm tag mới
        add_group = QtWidgets.QGroupBox("➕ Thêm Tag Mới")
        add_layout = QtWidgets.QHBoxLayout(add_group)

        self.new_tag_input = QtWidgets.QLineEdit()
        self.new_tag_input.setPlaceholderText("Nhập tên tag mới...")
        self.new_tag_input.returnPressed.connect(self._add_new_tag)

        self.color_button = QtWidgets.QPushButton("🎨")
        self.color_button.setFixedSize(40, 30)
        self.color_button.clicked.connect(self._choose_color)
        self.color_button.setToolTip("Chọn màu tag")
        self.current_color = "#3498db"  # Màu mặc định
        self._update_color_button()

        add_button = QtWidgets.QPushButton("➕ Thêm")
        add_button.clicked.connect(self._add_new_tag)
        add_button.setStyleSheet("""
            QPushButton {
                background-color: #28a745;
                color: white;
                border: none;
                padding: 5px 15px;
                border-radius: 3px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #218838;
            }
        """)

        add_layout.addWidget(QtWidgets.QLabel("Tag:"))
        add_layout.addWidget(self.new_tag_input, 1)
        add_layout.addWidget(self.color_button)
        add_layout.addWidget(add_button)

        layout.addWidget(add_group)

        # Splitter cho tags hiện tại và có sẵn
        splitter = QtWidgets.QSplitter(Qt.Horizontal)

        # Phần tags hiện tại
        current_group = QtWidgets.QGroupBox("🏷️ Tags Hiện Tại")
        current_layout = QtWidgets.QVBoxLayout(current_group)

        self.current_tags_list = QtWidgets.QListWidget()
        self.current_tags_list.setMaximumHeight(150)
        current_layout.addWidget(self.current_tags_list)

        remove_button = QtWidgets.QPushButton("🗑️ Xóa Tag Đã Chọn")
        remove_button.clicked.connect(self._remove_selected_tag)
        remove_button.setStyleSheet("""
            QPushButton {
                background-color: #dc3545;
                color: white;
                border: none;
                padding: 5px 15px;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #c82333;
            }
        """)
        current_layout.addWidget(remove_button)

        splitter.addWidget(current_group)

        # Phần tags có sẵn
        available_group = QtWidgets.QGroupBox("📚 Tags Có Sẵn (Click để thêm)")
        available_layout = QtWidgets.QVBoxLayout(available_group)

        self.available_tags_list = QtWidgets.QListWidget()
        self.available_tags_list.setMaximumHeight(150)
        self.available_tags_list.itemClicked.connect(self._add_existing_tag)
        available_layout.addWidget(self.available_tags_list)

        splitter.addWidget(available_group)
        layout.addWidget(splitter)

        # Buttons
        button_layout = QtWidgets.QHBoxLayout()

        save_button = QtWidgets.QPushButton("💾 Lưu")
        save_button.clicked.connect(self.accept)
        save_button.setStyleSheet("""
            QPushButton {
                background-color: #007bff;
                color: white;
                border: none;
                padding: 8px 20px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #0056b3;
            }
        """)

        cancel_button = QtWidgets.QPushButton("❌ Hủy")
        cancel_button.clicked.connect(self.reject)
        cancel_button.setStyleSheet("""
            QPushButton {
                background-color: #6c757d;
                color: white;
                border: none;
                padding: 8px 20px;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #545b62;
            }
        """)

        button_layout.addStretch()
        button_layout.addWidget(cancel_button)
        button_layout.addWidget(save_button)

        layout.addLayout(button_layout)

    def _update_color_button(self):
        """Cập nhật màu nút chọn màu"""
        self.color_button.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.current_color};
                border: 2px solid #333;
                border-radius: 4px;
            }}
        """)

    def _choose_color(self):
        """Chọn màu cho tag"""
        color = QtWidgets.QColorDialog.getColor(QtGui.QColor(self.current_color), self)
        if color.isValid():
            self.current_color = color.name()
            self._update_color_button()

    def _load_current_tags(self):
        """Load tags hiện tại của câu hỏi"""
        try:
            tags = self.db.execute_query(
                "SELECT tag_name, color FROM question_tags WHERE question_id=?",
                (self.question_id,), fetch="all"
            ) or []

            self.current_tags_list.clear()

            for tag in tags:
                tag_dict = dict(tag) if hasattr(tag, 'keys') else {}
                tag_name = tag_dict.get('tag_name', '')
                tag_color = tag_dict.get('color', '#3498db')

                if tag_name:
                    item = QtWidgets.QListWidgetItem(tag_name)
                    item.setBackground(QtGui.QColor(tag_color))
                    item.setForeground(QtGui.QColor('white'))
                    item.setData(Qt.UserRole, tag_color)
                    self.current_tags_list.addItem(item)

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể load tags hiện tại: {e}")

    def _load_available_tags(self):
        """Load tất cả tags có sẵn trong hệ thống"""
        try:
            # Lấy tất cả tags unique, loại trừ tags đã có
            current_tag_names = [self.current_tags_list.item(i).text()
                                 for i in range(self.current_tags_list.count())]

            all_tags = self.db.execute_query(
                "SELECT DISTINCT tag_name, color FROM question_tags ORDER BY tag_name",
                fetch="all"
            ) or []

            self.available_tags_list.clear()

            for tag in all_tags:
                tag_dict = dict(tag) if hasattr(tag, 'keys') else {}
                tag_name = tag_dict.get('tag_name', '')
                tag_color = tag_dict.get('color', '#3498db')

                # Chỉ hiển thị tags chưa được thêm
                if tag_name and tag_name not in current_tag_names:
                    item = QtWidgets.QListWidgetItem(f"🏷️ {tag_name}")
                    item.setBackground(QtGui.QColor(tag_color))
                    item.setForeground(QtGui.QColor('white'))
                    item.setData(Qt.UserRole, {'name': tag_name, 'color': tag_color})
                    self.available_tags_list.addItem(item)

        except Exception as e:
            print(f"⚠️ Lỗi load available tags: {e}")

    def _add_new_tag(self):
        """Thêm tag mới"""
        tag_name = self.new_tag_input.text().strip()
        if not tag_name:
            QtWidgets.QMessageBox.warning(self, "Lỗi", "Vui lòng nhập tên tag")
            return

        # Kiểm tra tag đã tồn tại chưa
        existing_tags = [self.current_tags_list.item(i).text()
                         for i in range(self.current_tags_list.count())]

        if tag_name in existing_tags:
            QtWidgets.QMessageBox.warning(self, "Lỗi", "Tag này đã tồn tại")
            return

        try:
            # Thêm vào database
            self.db.execute_query(
                "INSERT OR IGNORE INTO question_tags (question_id, tag_name, color) VALUES (?, ?, ?)",
                (self.question_id, tag_name, self.current_color)
            )

            # Thêm vào danh sách hiện tại
            item = QtWidgets.QListWidgetItem(tag_name)
            item.setBackground(QtGui.QColor(self.current_color))
            item.setForeground(QtGui.QColor('white'))
            item.setData(Qt.UserRole, self.current_color)
            self.current_tags_list.addItem(item)

            # Clear input
            self.new_tag_input.clear()

            # Refresh available tags
            self._load_available_tags()

            QtWidgets.QMessageBox.information(self, "Thành công", f"Đã thêm tag '{tag_name}'")

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể thêm tag: {e}")

    def _add_existing_tag(self, item):
        """Thêm tag có sẵn vào câu hỏi"""
        if not item:
            return

        tag_data = item.data(Qt.UserRole)
        if not tag_data:
            return

        tag_name = tag_data['name']
        tag_color = tag_data['color']

        try:
            # Thêm vào database
            self.db.execute_query(
                "INSERT OR IGNORE INTO question_tags (question_id, tag_name, color) VALUES (?, ?, ?)",
                (self.question_id, tag_name, tag_color)
            )

            # Thêm vào danh sách hiện tại
            new_item = QtWidgets.QListWidgetItem(tag_name)
            new_item.setBackground(QtGui.QColor(tag_color))
            new_item.setForeground(QtGui.QColor('white'))
            new_item.setData(Qt.UserRole, tag_color)
            self.current_tags_list.addItem(new_item)

            # Xóa khỏi available tags
            row = self.available_tags_list.row(item)
            self.available_tags_list.takeItem(row)

            QtWidgets.QMessageBox.information(self, "Thành công", f"Đã thêm tag '{tag_name}'")

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể thêm tag: {e}")

    def _remove_selected_tag(self):
        """Xóa tag được chọn"""
        current_item = self.current_tags_list.currentItem()
        if not current_item:
            QtWidgets.QMessageBox.warning(self, "Lỗi", "Vui lòng chọn tag để xóa")
            return

        tag_name = current_item.text()

        if QtWidgets.QMessageBox.question(
                self, "Xác nhận", f"Bạn có muốn xóa tag '{tag_name}'?"
        ) != QtWidgets.QMessageBox.Yes:
            return

        try:
            # Xóa khỏi database
            self.db.execute_query(
                "DELETE FROM question_tags WHERE question_id=? AND tag_name=?",
                (self.question_id, tag_name)
            )

            # Xóa khỏi danh sách
            row = self.current_tags_list.row(current_item)
            self.current_tags_list.takeItem(row)

            # Refresh available tags
            self._load_available_tags()

            QtWidgets.QMessageBox.information(self, "Thành công", f"Đã xóa tag '{tag_name}'")

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể xóa tag: {e}")


class TreeNodeDialog(QtWidgets.QDialog):
    def __init__(self, db_manager, mode="add", node_id=None, parent=None):
        super().__init__(parent)
        self.db = db_manager
        self.mode = mode  # "add" hoặc "edit"
        self.node_id = node_id
        self.parent_id = None

        self._setup_dialog()
        self._build_ui()
        self._load_data()

    def _setup_dialog(self):
        """Thiết lập dialog"""
        if self.mode == "add":
            self.setWindowTitle("➕ Thêm nhánh mới")
        else:
            self.setWindowTitle("✏️ Sửa nhánh")

        self.setModal(True)
        self.resize(450, 400)

    def _build_ui(self):
        """Xây dựng giao diện"""
        layout = QtWidgets.QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)

        # Header
        header = QtWidgets.QLabel()
        if self.mode == "add":
            header.setText("➕ Thêm nhánh mới vào cây thư mục")
        else:
            header.setText("✏️ Chỉnh sửa thông tin nhánh")

        header.setStyleSheet("""
            QLabel {
                font-size: 16px;
                font-weight: bold;
                color: #2E86AB;
                padding: 15px;
                background-color: #f8f9fa;
                border-radius: 8px;
                border: 1px solid #dee2e6;
            }
        """)
        layout.addWidget(header)

        # Form container
        form_container = QtWidgets.QWidget()
        form_container.setStyleSheet("""
            QWidget {
                background-color: white;
                border-radius: 8px;
                border: 1px solid #e1e5e9;
            }
        """)

        form_layout = QtWidgets.QFormLayout(form_container)
        form_layout.setSpacing(12)
        form_layout.setContentsMargins(20, 20, 20, 20)

        # Parent selection (chỉ hiện khi thêm)
        if self.mode == "add":
            self.parent_combo = QtWidgets.QComboBox()
            self.parent_combo.addItem("(Không có parent - Cấp gốc)", None)
            self._load_parent_options()

            parent_label = QtWidgets.QLabel("📁 Nhánh cha:")
            parent_label.setStyleSheet("font-weight: 500; color: #495057;")
            form_layout.addRow(parent_label, self.parent_combo)

        # Tên nhánh
        self.name_edit = QtWidgets.QLineEdit()
        self.name_edit.setPlaceholderText("Nhập tên nhánh...")

        name_label = QtWidgets.QLabel("📝 Tên nhánh:")
        name_label.setStyleSheet("font-weight: 500; color: #495057;")
        form_layout.addRow(name_label, self.name_edit)

        # Cấp độ
        self.level_combo = QtWidgets.QComboBox()
        self.level_combo.addItems(["Môn", "Lớp", "Chủ đề", "Dạng", "Mức độ"])

        level_label = QtWidgets.QLabel("📊 Cấp độ:")
        level_label.setStyleSheet("font-weight: 500; color: #495057;")
        form_layout.addRow(level_label, self.level_combo)

        # Mô tả
        self.description_edit = QtWidgets.QTextEdit()
        self.description_edit.setMaximumHeight(100)
        self.description_edit.setPlaceholderText("Nhập mô tả chi tiết...")

        desc_label = QtWidgets.QLabel("📄 Mô tả:")
        desc_label.setStyleSheet("font-weight: 500; color: #495057;")
        form_layout.addRow(desc_label, self.description_edit)

        # Style cho form inputs
        input_style = """
            QLineEdit, QComboBox, QTextEdit {
                padding: 10px;
                border: 2px solid #e1e5e9;
                border-radius: 6px;
                font-size: 13px;
                background-color: white;
            }
            QLineEdit:focus, QComboBox:focus, QTextEdit:focus {
                border-color: #2E86AB;
                outline: none;
                background-color: #f8fbff;
            }
        """

        self.name_edit.setStyleSheet(input_style)
        self.level_combo.setStyleSheet(input_style)
        self.description_edit.setStyleSheet(input_style)

        if hasattr(self, 'parent_combo'):
            self.parent_combo.setStyleSheet(input_style)

        layout.addWidget(form_container)

        # Buttons
        button_layout = QtWidgets.QHBoxLayout()
        button_layout.setContentsMargins(0, 10, 0, 0)

        cancel_btn = QtWidgets.QPushButton("❌ Hủy")
        cancel_btn.setFixedSize(100, 40)
        cancel_btn.clicked.connect(self.reject)
        cancel_btn.setStyleSheet("""
            QPushButton {
                background-color: #6c757d;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 13px;
            }
            QPushButton:hover {
                background-color: #5a6268;
            }
        """)

        # Save button
        if self.mode == "add":
            save_btn = QtWidgets.QPushButton("➕ Thêm")
            save_btn.setStyleSheet("""
                QPushButton {
                    background-color: #28a745;
                    color: white;
                    border: none;
                    padding: 10px 20px;
                    border-radius: 6px;
                    font-weight: 600;
                    font-size: 13px;
                }
                QPushButton:hover {
                    background-color: #218838;
                }
            """)
        else:
            save_btn = QtWidgets.QPushButton("💾 Lưu")
            save_btn.setStyleSheet("""
                QPushButton {
                    background-color: #007bff;
                    color: white;
                    border: none;
                    padding: 10px 20px;
                    border-radius: 6px;
                    font-weight: 600;
                    font-size: 13px;
                }
                QPushButton:hover {
                    background-color: #0056b3;
                }
            """)

        save_btn.setFixedSize(100, 40)
        save_btn.clicked.connect(self.accept)
        save_btn.setDefault(True)

        button_layout.addStretch()
        button_layout.addWidget(cancel_btn)
        button_layout.addSpacing(10)
        button_layout.addWidget(save_btn)

        layout.addWidget(QtWidgets.QWidget())
        layout.layout().addLayout(button_layout)

        # Focus vào name edit
        self.name_edit.setFocus()

    def _load_parent_options(self):
        """Load danh sách parent có thể chọn"""
        if self.mode != "add":
            return

        try:
            rows = self.db.execute_query(
                "SELECT id, name, level FROM exercise_tree ORDER BY level, name",
                fetch="all"
            ) or []

            for row in rows:
                # Nếu đang edit, không cho chọn chính nó làm parent
                if self.mode == "edit" and row["id"] == self.node_id:
                    continue

                display_text = f"{row['name']} ({row['level']})"
                self.parent_combo.addItem(display_text, row["id"])

        except Exception as e:
            print(f"Lỗi load parent options: {e}")

    def set_parent_id(self, parent_id):
        """Đặt parent được chọn"""
        self.parent_id = parent_id

        if self.mode == "add" and hasattr(self, 'parent_combo'):
            # Tìm và chọn parent trong combo
            for i in range(self.parent_combo.count()):
                if self.parent_combo.itemData(i) == parent_id:
                    self.parent_combo.setCurrentIndex(i)
                    break

    def _load_data(self):
        """Load dữ liệu nếu đang edit"""
        if self.mode != "edit" or not self.node_id:
            return

        try:
            row = self.db.execute_query(
                "SELECT name, level, description FROM exercise_tree WHERE id = ?",
                (self.node_id,), fetch="one"
            )

            if row:
                self.name_edit.setText(row["name"] or "")

                # Set level
                level = row["level"] or "Môn"
                index = self.level_combo.findText(level)
                if index >= 0:
                    self.level_combo.setCurrentIndex(index)

                # Description
                description = row.get('description', '') or ''
                self.description_edit.setPlainText(description)

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể tải dữ liệu: {e}")

    def accept(self):
        """Xử lý khi người dùng nhấn Save/Add"""
        if not self._validate_input():
            return

        name = self.name_edit.text().strip()
        level = self.level_combo.currentText()
        description = self.description_edit.toPlainText().strip()

        try:
            if self.mode == "add":
                # Thêm node mới
                parent_id = None
                if hasattr(self, 'parent_combo'):
                    parent_id = self.parent_combo.currentData()
                elif self.parent_id:
                    parent_id = self.parent_id

                self.db.execute_query(
                    "INSERT INTO exercise_tree (parent_id, name, level, description) VALUES (?, ?, ?, ?)",
                    (parent_id, name, level, description)
                )

                QtWidgets.QMessageBox.information(
                    self, "Thành công",
                    f"Đã thêm nhánh '{name}' thành công!"
                )

            else:
                # Cập nhật node
                self.db.execute_query(
                    "UPDATE exercise_tree SET name = ?, level = ?, description = ? WHERE id = ?",
                    (name, level, description, self.node_id)
                )

                QtWidgets.QMessageBox.information(
                    self, "Thành công",
                    f"Đã cập nhật nhánh '{name}' thành công!"
                )

            super().accept()

        except Exception as e:
            QtWidgets.QMessageBox.critical(
                self, "Lỗi database",
                f"Không thể lưu dữ liệu:\n{str(e)}"
            )

    def _validate_input(self):
        """Validate dữ liệu đầu vào"""
        name = self.name_edit.text().strip()

        if not name:
            QtWidgets.QMessageBox.warning(
                self, "Lỗi",
                "Tên nhánh không được để trống!"
            )
            self.name_edit.setFocus()
            return False

        if len(name) > 100:
            QtWidgets.QMessageBox.warning(
                self, "Lỗi",
                "Tên nhánh không được quá 100 ký tự!"
            )
            self.name_edit.setFocus()
            return False

        return True


# ========== CLASS QUẢN LÝ CÂY NÂNG CAO ========== #
class TreeManagerDialog(QtWidgets.QDialog):
    def __init__(self, db_manager, parent=None):
        super().__init__(parent)
        self.db = db_manager
        self.setWindowTitle("⚙️ Quản lý cây thư mục")
        self.setModal(True)
        self.resize(800, 600)
        self._build_ui()
        self._load_tree_data()

    def _build_ui(self):
        """Xây dựng giao diện"""
        layout = QtWidgets.QVBoxLayout(self)

        # Toolbar
        toolbar = QtWidgets.QToolBar()
        toolbar.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)

        add_action = toolbar.addAction("➕ Thêm nhánh")
        add_action.triggered.connect(self._add_node)

        edit_action = toolbar.addAction("✏️ Sửa nhánh")
        edit_action.triggered.connect(self._edit_node)

        delete_action = toolbar.addAction("🗑️ Xóa nhánh")
        delete_action.triggered.connect(self._delete_node)

        toolbar.addSeparator()

        export_action = toolbar.addAction("📤 Xuất cấu trúc")
        export_action.triggered.connect(self._export_structure)

        layout.addWidget(toolbar)

        # Tree view
        self.tree_table = QtWidgets.QTreeWidget()
        self.tree_table.setHeaderLabels(["Tên", "Cấp độ", "Số câu hỏi", "Mô tả"])
        self.tree_table.setColumnWidth(0, 250)
        self.tree_table.setColumnWidth(1, 100)
        self.tree_table.setColumnWidth(2, 100)

        layout.addWidget(self.tree_table)

        # Buttons
        button_layout = QtWidgets.QHBoxLayout()

        close_btn = QtWidgets.QPushButton("Đóng")
        close_btn.clicked.connect(self.accept)

        button_layout.addStretch()
        button_layout.addWidget(close_btn)

        layout.addLayout(button_layout)

    def _load_tree_data(self):
        """Load dữ liệu cây"""
        # Implementation tương tự refresh_tree nhưng hiển thị trong table
        pass

    def _add_node(self):
        """Thêm node mới"""
        dialog = TreeNodeDialog(self.db, mode="add", parent=self)
        if dialog.exec() == QtWidgets.QDialog.Accepted:
            self._load_tree_data()

    def _edit_node(self):
        """Sửa node được chọn"""
        # Implementation
        pass

    def _delete_node(self):
        """Xóa node được chọn"""
        # Implementation
        pass

    def _export_structure(self):
        """Xuất cấu trúc cây"""
        # Implementation
        pass


# ========== DIALOG XEM ẢNH FULLSCREEN ========== #
# ========== DIALOG XEM ẢNH FULLSCREEN - HOÀN THIỆN ========== #
class ImageViewerDialog(QtWidgets.QDialog):
    """Dialog xem ảnh fullscreen với đầy đủ tính năng zoom, pan, keyboard shortcuts"""

    def __init__(self, pixmap, parent=None):
        super().__init__(parent)
        self.original_pixmap = pixmap
        self.current_zoom = 1.0
        self.min_zoom = 0.1
        self.max_zoom = 10.0
        self.zoom_step = 0.2
        self.pan_offset = QtCore.QPoint(0, 0)
        self.last_pan_point = QtCore.QPoint()
        self.is_panning = False

        self._setup_window()
        self._setup_ui()
        self._setup_shortcuts()
        self._fit_to_window()

    def _setup_window(self):
        """#(Thiết lập cửa sổ dialog maximized - phóng to nhưng vẫn có thanh tiêu đề và viền)"""
        self.setWindowTitle("🖼️ Xem ảnh - Nhấn ESC để thoát")

        # Sử dụng window flags bình thường để có thanh tiêu đề và viền
        self.setWindowFlags(
            Qt.Dialog |
            Qt.WindowMaximizeButtonHint |  # Có nút maximize
            Qt.WindowMinimizeButtonHint |  # Có nút minimize
            Qt.WindowCloseButtonHint  # Có nút close
        )

        # Màu nền đen cho viewer
        self.setStyleSheet("""
            QDialog {
                background-color: #000000;
            }
        """)
        self.showMaximized()
        QtCore.QTimer.singleShot(100, self._fit_to_window)
    # Giao diện xem ảnh câu hỏi khi nhấn đúp
    def _setup_ui(self):
        """Thiết lập giao diện với toolbar đầy đủ, controls và styling chuyên nghiệp"""
        layout = QtWidgets.QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # Window flags đầy đủ
        self.setWindowFlags(
            Qt.Window |
            Qt.WindowMinimizeButtonHint |
            Qt.WindowMaximizeButtonHint |
            Qt.WindowCloseButtonHint |
            Qt.WindowSystemMenuHint |
            Qt.WindowTitleHint |
            Qt.CustomizeWindowHint
        )

        # Thiết lập toolbar chính với styling đẹp
        self._setup_main_toolbar(layout)

        # (Container chính với background tối chuyên nghiệp)
        main_container = QtWidgets.QWidget()
        main_container.setStyleSheet("""
            QWidget {
                background: qlineargradient(x1: 0, y1: 0, x2: 0, y2: 1,
                                          stop: 0 #1a1a1a, stop: 1 #000000);
            }
        """)
        main_layout = QtWidgets.QVBoxLayout(main_container)
        main_layout.setContentsMargins(8, 8, 8, 8)

        # Scroll area cho ảnh với styling cao cấp
        self.scroll_area = QtWidgets.QScrollArea()
        self.scroll_area.setAlignment(Qt.AlignCenter)
        self.scroll_area.setStyleSheet("""
            QScrollArea { 
                border: 2px solid #333;
                border-radius: 8px;
                background: #000000;
            }
            QScrollArea QScrollBar:vertical {
                background: #2b2b2b;
                width: 12px;
                border-radius: 6px;
            }
            QScrollArea QScrollBar::handle:vertical {
                background: #555;
                border-radius: 6px;
                min-height: 30px;
            }
            QScrollArea QScrollBar::handle:vertical:hover {
                background: #777;
            }
            QScrollArea QScrollBar:horizontal {
                background: #2b2b2b;
                height: 12px;
                border-radius: 6px;
            }
            QScrollArea QScrollBar::handle:horizontal {
                background: #555;
                border-radius: 6px;
                min-width: 30px;
            }
            QScrollArea QScrollBar::handle:horizontal:hover {
                background: #777;
            }
        """)

        # (Image label với effects và interactions)
        self.image_label = QtWidgets.QLabel()
        self.image_label.setAlignment(Qt.AlignCenter)
        self.image_label.setStyleSheet("""
            QLabel { 
                background: transparent;
                border: 1px solid #444;
                border-radius: 4px;
                padding: 5px;
            }
        """)
        self.image_label.setPixmap(self.original_pixmap)
        self.image_label.setMinimumSize(200, 150)

        # Enable mouse interactions cho pan và zoom
        self.image_label.setMouseTracking(True)
        self.image_label.mousePressEvent = self._mouse_press_event
        self.image_label.mouseMoveEvent = self._mouse_move_event
        self.image_label.mouseReleaseEvent = self._mouse_release_event
        self.image_label.mouseDoubleClickEvent = self._mouse_double_click_event

        # Thiết lập cursor hints
        self.image_label.setCursor(Qt.OpenHandCursor)

        self.scroll_area.setWidget(self.image_label)
        main_layout.addWidget(self.scroll_area)
        layout.addWidget(main_container)

        # Status bar với thông tin chi tiết
        self._setup_enhanced_status_bar(layout)

        # Context menu cho toàn dialog
        self._setup_context_menu()

    def _setup_main_toolbar(self, layout):
        """#(Thiết lập toolbar chính với tất cả controls)"""
        toolbar = QtWidgets.QWidget()
        toolbar.setFixedHeight(55)
        toolbar.setStyleSheet("""
            QWidget {
                background: qlineargradient(x1: 0, y1: 0, x2: 0, y2: 1,
                                          stop: 0 #f8f9fa, stop: 1 #e9ecef);
                border-bottom: 2px solid #adb5bd;
            }
            QPushButton {
                background: white;
                border: 1px solid #ced4da;
                border-radius: 6px;
                color: #495057;
                font-weight: 500;
                padding: 8px 12px;
                margin: 2px;
                min-width: 80px;
            }
            QPushButton:hover {
                background: #e3f2fd;
                border-color: #2196f3;
                color: #1976d2;
            }
            QPushButton:pressed {
                background: #bbdefb;
            }
            QPushButton:checked {
                background: #4caf50;
                color: white;
                border-color: #388e3c;
            }
            QSlider::groove:horizontal {
                border: 1px solid #bbb;
                height: 8px;
                background: white;
                border-radius: 4px;
            }
            QSlider::handle:horizontal {
                background: #2196f3;
                border: 1px solid #1976d2;
                width: 18px;
                border-radius: 9px;
                margin: -5px 0;
            }
            QSlider::handle:horizontal:hover {
                background: #1976d2;
            }
            QLabel {
                color: #495057;
                font-weight: 500;
                padding: 0 8px;
            }
        """)

        toolbar_layout = QtWidgets.QHBoxLayout(toolbar)
        toolbar_layout.setContentsMargins(15, 8, 15, 8)
        toolbar_layout.setSpacing(8)

        # (Nhóm điều khiển cửa sổ)
        window_group = QtWidgets.QWidget()
        window_layout = QtWidgets.QHBoxLayout(window_group)
        window_layout.setContentsMargins(0, 0, 0, 0)
        window_layout.setSpacing(4)

        self.minimize_btn = QtWidgets.QPushButton("🗕 Thu nhỏ")
        self.minimize_btn.setToolTip("Thu nhỏ cửa sổ (Alt+F9)")
        self.minimize_btn.clicked.connect(self.showMinimized)

        self.maximize_btn = QtWidgets.QPushButton("🗖 Phóng to")
        self.maximize_btn.setToolTip("Phóng to/Khôi phục (Alt+F10)")

        self.hide_btn = QtWidgets.QPushButton("👁️ Ẩn")
        self.hide_btn.setToolTip("Ẩn cửa sổ (Ctrl+H)")
        self.hide_btn.clicked.connect(self.hide)

        window_layout.addWidget(self.minimize_btn)
        window_layout.addWidget(self.maximize_btn)
        window_layout.addWidget(self.hide_btn)
        toolbar_layout.addWidget(window_group)

        # Separator
        separator1 = QtWidgets.QFrame()
        separator1.setFrameShape(QtWidgets.QFrame.VLine)
        separator1.setStyleSheet("color: #adb5bd;")
        toolbar_layout.addWidget(separator1)

        # (Nhóm zoom controls)
        zoom_group = QtWidgets.QWidget()
        zoom_layout = QtWidgets.QHBoxLayout(zoom_group)
        zoom_layout.setContentsMargins(0, 0, 0, 0)
        zoom_layout.setSpacing(4)

        self.zoom_out_btn = QtWidgets.QPushButton("🔍−")
        self.zoom_out_btn.setToolTip("Thu nhỏ (Ctrl + -)")
        self.zoom_out_btn.clicked.connect(self._zoom_out)
        self.zoom_out_btn.setFixedWidth(40)

        # Zoom slider với styling đẹp
        self.zoom_slider = QtWidgets.QSlider(Qt.Horizontal)
        self.zoom_slider.setRange(int(self.min_zoom * 100), int(self.max_zoom * 100))
        self.zoom_slider.setValue(int(self.current_zoom * 100))
        self.zoom_slider.setFixedWidth(120)
        self.zoom_slider.setToolTip("Kéo để điều chỉnh zoom")
        self.zoom_slider.valueChanged.connect(self._zoom_slider_changed)

        self.zoom_in_btn = QtWidgets.QPushButton("🔍+")
        self.zoom_in_btn.setToolTip("Phóng to (Ctrl + +)")
        self.zoom_in_btn.clicked.connect(self._zoom_in)
        self.zoom_in_btn.setFixedWidth(40)

        # Zoom percentage display
        self.zoom_label = QtWidgets.QLabel("100%")
        self.zoom_label.setFixedWidth(50)
        self.zoom_label.setAlignment(Qt.AlignCenter)

        zoom_layout.addWidget(self.zoom_out_btn)
        zoom_layout.addWidget(self.zoom_slider)
        zoom_layout.addWidget(self.zoom_in_btn)
        zoom_layout.addWidget(self.zoom_label)
        toolbar_layout.addWidget(zoom_group)

        # Separator
        separator2 = QtWidgets.QFrame()
        separator2.setFrameShape(QtWidgets.QFrame.VLine)
        separator2.setStyleSheet("color: #adb5bd;")
        toolbar_layout.addWidget(separator2)

        # (Nhóm fit controls)
        fit_group = QtWidgets.QWidget()
        fit_layout = QtWidgets.QHBoxLayout(fit_group)
        fit_layout.setContentsMargins(0, 0, 0, 0)
        fit_layout.setSpacing(4)

        self.fit_window_btn = QtWidgets.QPushButton("📐 Vừa cửa sổ")
        self.fit_window_btn.setToolTip("Vừa khít cửa sổ (Ctrl + 0)")
        self.fit_window_btn.clicked.connect(self._fit_to_window)

        self.actual_size_btn = QtWidgets.QPushButton("1:1 Gốc")
        self.actual_size_btn.setToolTip("Hiển thị kích thước gốc (Ctrl + 1)")
        self.actual_size_btn.clicked.connect(self._actual_size)

        fit_layout.addWidget(self.fit_window_btn)
        fit_layout.addWidget(self.actual_size_btn)
        toolbar_layout.addWidget(fit_group)

        # Separator
        separator3 = QtWidgets.QFrame()
        separator3.setFrameShape(QtWidgets.QFrame.VLine)
        separator3.setStyleSheet("color: #adb5bd;")
        toolbar_layout.addWidget(separator3)

        # (Nhóm tiện ích)
        utility_group = QtWidgets.QWidget()
        utility_layout = QtWidgets.QHBoxLayout(utility_group)
        utility_layout.setContentsMargins(0, 0, 0, 0)
        utility_layout.setSpacing(4)

        self.always_top_btn = QtWidgets.QPushButton("📌 Luôn trên")
        self.always_top_btn.setCheckable(True)
        self.always_top_btn.setToolTip("Giữ cửa sổ luôn ở trên (Ctrl+T)")

        self.opacity_btn = QtWidgets.QPushButton("🔅 Độ mờ")
        self.opacity_btn.setToolTip("Điều chỉnh độ trong suốt")

        utility_layout.addWidget(self.always_top_btn)
        utility_layout.addWidget(self.opacity_btn)
        toolbar_layout.addWidget(utility_group)

        # Spacer để đẩy close button sang phải
        toolbar_layout.addStretch()

        # (Nút đóng nổi bật)
        self.close_btn = QtWidgets.QPushButton("✖ Đóng")
        self.close_btn.setToolTip("Đóng cửa sổ (ESC)")
        self.close_btn.clicked.connect(self.close)
        self.close_btn.setStyleSheet("""
            QPushButton {
                background: #dc3545;
                color: white;
                border: 1px solid #c82333;
                font-weight: bold;
                min-width: 60px;
            }
            QPushButton:hover {
                background: #c82333;
                border-color: #bd2130;
            }
            QPushButton:pressed {
                background: #bd2130;
            }
        """)
        toolbar_layout.addWidget(self.close_btn)

        layout.addWidget(toolbar)

        # Connect maximize toggle
        def toggle_maximize():
            if self.isMaximized():
                self.showNormal()
                self.maximize_btn.setText("🗖 Phóng to")
                self.maximize_btn.setToolTip("Phóng to cửa sổ")
            else:
                self.showMaximized()
                self.maximize_btn.setText("🗗 Khôi phục")
                self.maximize_btn.setToolTip("Khôi phục kích thước ban đầu")

        self.maximize_btn.clicked.connect(toggle_maximize)

        # Connect utility buttons
        self.always_top_btn.toggled.connect(self._toggle_always_on_top)
        self.opacity_btn.clicked.connect(self._show_opacity_dialog)

    def _setup_enhanced_status_bar(self, layout):
        """#(Thiết lập status bar nâng cao với nhiều thông tin)"""
        self.status_bar = QtWidgets.QWidget()
        self.status_bar.setFixedHeight(35)
        self.status_bar.setStyleSheet("""
            QWidget {
                background: qlineargradient(x1: 0, y1: 0, x2: 0, y2: 1,
                                          stop: 0 #2b2b2b, stop: 1 #1a1a1a);
                border-top: 1px solid #444;
            }
            QLabel {
                color: #ffffff;
                padding: 0 12px;
                font-size: 11px;
            }
            QLabel#main_info {
                font-weight: bold;
                color: #4fc3f7;
            }
            QLabel#tips {
                color: #81c784;
                font-style: italic;
            }
        """)

        status_layout = QtWidgets.QHBoxLayout(self.status_bar)
        status_layout.setContentsMargins(8, 0, 8, 0)

        # Main image info
        self.main_status_label = QtWidgets.QLabel()
        self.main_status_label.setObjectName("main_info")
        status_layout.addWidget(self.main_status_label)

        # Spacer
        status_layout.addStretch()

        # Tips label
        self.tips_label = QtWidgets.QLabel("💡 Kéo chuột để di chuyển • Ctrl+Scroll để zoom • Double-click để fit")
        self.tips_label.setObjectName("tips")
        status_layout.addWidget(self.tips_label)

        layout.addWidget(self.status_bar)

    def _setup_context_menu(self):
        """#(Thiết lập context menu toàn diện)"""
        self.setContextMenuPolicy(Qt.CustomContextMenu)
        self.customContextMenuRequested.connect(self._show_enhanced_context_menu)

    def _show_enhanced_context_menu(self, position):
        """Hiển thị context menu nâng cao"""
        menu = QtWidgets.QMenu(self)
        menu.setStyleSheet("""
            QMenu {
                background: #2b2b2b;
                color: white;
                border: 1px solid #555;
                border-radius: 4px;
                padding: 4px;
            }
            QMenu::item {
                padding: 8px 20px;
                border-radius: 3px;
            }
            QMenu::item:selected {
                background: #4fc3f7;
            }
            QMenu::separator {
                height: 1px;
                background: #555;
                margin: 4px 0;
            }
        """)

        # Window controls
        menu.addAction("🗕 Thu nhỏ", self.showMinimized)

        if self.isMaximized():
            menu.addAction("🗗 Khôi phục", self.showNormal)
        else:
            menu.addAction("🗖 Phóng to", self.showMaximized)

        menu.addAction("👁️ Ẩn", self.hide)
        menu.addSeparator()

        # Zoom controls
        menu.addAction("🔍+ Phóng to ảnh", self._zoom_in)
        menu.addAction("🔍− Thu nhỏ ảnh", self._zoom_out)
        menu.addAction("📐 Vừa cửa sổ", self._fit_to_window)
        menu.addAction("1:1 Kích thước gốc", self._actual_size)
        menu.addSeparator()

        # Image utilities
        if self.original_pixmap and not self.original_pixmap.isNull():
            menu.addAction("💾 Lưu ảnh...", lambda: self._save_image())
            menu.addAction("📋 Copy ảnh", lambda: QtWidgets.QApplication.clipboard().setPixmap(self.original_pixmap))
            menu.addSeparator()

        # Window utilities
        always_top_text = "📌 Luôn trên ✓" if self.always_top_btn.isChecked() else "📌 Luôn trên"
        menu.addAction(always_top_text, lambda: self.always_top_btn.toggle())
        menu.addAction("🔅 Độ mờ...", self._show_opacity_dialog)
        menu.addSeparator()

        # Close
        menu.addAction("✖ Đóng", self.close)

        menu.exec(self.mapToGlobal(position))

    def _mouse_double_click_event(self, event):
        """#(Double-click để fit to window)"""
        if event.button() == Qt.LeftButton:
            self._fit_to_window()

    def _toggle_always_on_top(self, checked):
        """Toggle always on top functionality"""
        if checked:
            self.setWindowFlags(self.windowFlags() | Qt.WindowStaysOnTopHint)
            self.show()
            self.always_top_btn.setText("📌 Luôn trên ✓")
        else:
            self.setWindowFlags(self.windowFlags() & ~Qt.WindowStaysOnTopHint)
            self.show()
            self.always_top_btn.setText("📌 Luôn trên")

    def _show_opacity_dialog(self):
        """Hiển thị dialog điều chỉnh độ mờ"""
        opacity_dialog = QtWidgets.QDialog(self)
        opacity_dialog.setWindowTitle("Điều chỉnh độ mờ")
        opacity_dialog.resize(320, 140)
        opacity_dialog.setStyleSheet("""
            QDialog {
                background: #f8f9fa;
            }
            QSlider::groove:horizontal {
                border: 1px solid #bbb;
                height: 8px;
                background: white;
                border-radius: 4px;
            }
            QSlider::handle:horizontal {
                background: #2196f3;
                border: 1px solid #1976d2;
                width: 20px;
                border-radius: 10px;
                margin: -6px 0;
            }
        """)

        layout = QtWidgets.QVBoxLayout(opacity_dialog)
        layout.setSpacing(15)

        # Title
        title = QtWidgets.QLabel("🔅 Điều chỉnh độ trong suốt")
        title.setStyleSheet("font-weight: bold; font-size: 14px; color: #2c3e50;")
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)

        # Slider
        slider = QtWidgets.QSlider(Qt.Horizontal)
        slider.setRange(20, 100)
        slider.setValue(int(self.windowOpacity() * 100))

        # Label
        label = QtWidgets.QLabel(f"Độ mờ: {slider.value()}%")
        label.setAlignment(Qt.AlignCenter)
        label.setStyleSheet("font-size: 12px; color: #495057;")

        def update_opacity(value):
            self.setWindowOpacity(value / 100.0)
            label.setText(f"Độ mờ: {value}%")

        slider.valueChanged.connect(update_opacity)

        layout.addWidget(label)
        layout.addWidget(slider)

        # Buttons
        btn_layout = QtWidgets.QHBoxLayout()
        reset_btn = QtWidgets.QPushButton("🔄 Reset")
        reset_btn.clicked.connect(lambda: slider.setValue(100))
        close_btn = QtWidgets.QPushButton("✓ Đóng")
        close_btn.clicked.connect(opacity_dialog.accept)

        btn_layout.addWidget(reset_btn)
        btn_layout.addWidget(close_btn)
        layout.addLayout(btn_layout)

        opacity_dialog.exec()

    def _save_image(self):
        """Lưu ảnh hiện tại ra file"""
        if not self.original_pixmap or self.original_pixmap.isNull():
            return

        file_path, _ = QtWidgets.QFileDialog.getSaveFileName(
            self, "Lưu ảnh",
            f"image_{QtCore.QDateTime.currentDateTime().toString('yyyyMMdd_hhmmss')}.png",
            "PNG files (*.png);;JPG files (*.jpg);;All files (*.*)"
        )

        if file_path:
            if self.original_pixmap.save(file_path):
                QtWidgets.QMessageBox.information(self, "Thành công", f"Đã lưu ảnh: {file_path}")
            else:
                QtWidgets.QMessageBox.warning(self, "Lỗi", "Không thể lưu ảnh")
    def _setup_toolbar(self, layout):
        """#(Tạo toolbar với các nút điều khiển zoom và điều hướng)"""
        toolbar = QtWidgets.QWidget()
        toolbar.setFixedHeight(50)
        toolbar.setStyleSheet("""
            QWidget {
                background: rgba(0, 0, 0, 180);
                border-bottom: 1px solid #333;
            }
            QPushButton {
                background: rgba(255, 255, 255, 20);
                border: 1px solid rgba(255, 255, 255, 50);
                border-radius: 4px;
                color: white;
                font-size: 12px;
                padding: 5px 10px;
                min-width: 60px;
            }
            QPushButton:hover {
                background: rgba(255, 255, 255, 40);
                border: 1px solid rgba(255, 255, 255, 80);
            }
            QPushButton:pressed {
                background: rgba(255, 255, 255, 60);
            }
            QSlider::groove:horizontal {
                border: 1px solid #333;
                height: 6px;
                background: #555;
                border-radius: 3px;
            }
            QSlider::handle:horizontal {
                background: white;
                border: 1px solid #333;
                width: 16px;
                border-radius: 8px;
                margin: -5px 0;
            }
            QLabel {
                color: white;
                font-size: 11px;
            }
        """)

        toolbar_layout = QtWidgets.QHBoxLayout(toolbar)
        toolbar_layout.setContentsMargins(10, 5, 10, 5)

        # Zoom controls
        self.zoom_out_btn = QtWidgets.QPushButton("🔍-")
        self.zoom_out_btn.setToolTip("Thu nhỏ (Ctrl + -)")
        self.zoom_out_btn.clicked.connect(self._zoom_out)

        self.zoom_in_btn = QtWidgets.QPushButton("🔍+")
        self.zoom_in_btn.setToolTip("Phóng to (Ctrl + +)")
        self.zoom_in_btn.clicked.connect(self._zoom_in)

        # Zoom slider
        self.zoom_slider = QtWidgets.QSlider(Qt.Horizontal)
        self.zoom_slider.setRange(int(self.min_zoom * 100), int(self.max_zoom * 100))
        self.zoom_slider.setValue(int(self.current_zoom * 100))
        self.zoom_slider.setFixedWidth(150)
        self.zoom_slider.setToolTip("Kéo để điều chỉnh zoom")
        self.zoom_slider.valueChanged.connect(self._zoom_slider_changed)

        # Fit controls
        self.fit_window_btn = QtWidgets.QPushButton("⚏ Vừa cửa sổ")
        self.fit_window_btn.setToolTip("Vừa khít cửa sổ (Ctrl + 0)")
        self.fit_window_btn.clicked.connect(self._fit_to_window)

        self.actual_size_btn = QtWidgets.QPushButton("📐 Kích thước gốc")
        self.actual_size_btn.setToolTip("Hiển thị kích thước gốc (Ctrl + 1)")
        self.actual_size_btn.clicked.connect(self._actual_size)

        # Close button
        self.close_btn = QtWidgets.QPushButton("❌ Đóng")
        self.close_btn.setToolTip("Đóng (ESC)")
        self.close_btn.clicked.connect(self.close)

        # Add to layout
        toolbar_layout.addWidget(self.zoom_out_btn)
        toolbar_layout.addWidget(self.zoom_in_btn)
        toolbar_layout.addWidget(self.zoom_slider)
        toolbar_layout.addWidget(self.fit_window_btn)
        toolbar_layout.addWidget(self.actual_size_btn)
        toolbar_layout.addStretch()
        toolbar_layout.addWidget(self.close_btn)

        layout.addWidget(toolbar)

    def _setup_status_bar(self, layout):
        """#(Tạo thanh trạng thái hiển thị thông tin ảnh)"""
        self.status_bar = QtWidgets.QLabel()
        self.status_bar.setFixedHeight(25)
        self.status_bar.setStyleSheet("""
            QLabel {
                background: rgba(0, 0, 0, 180);
                color: white;
                padding: 0 10px;
                border-top: 1px solid #333;
                font-size: 11px;
            }
        """)
        self._update_status()
        layout.addWidget(self.status_bar)

    def _setup_shortcuts(self):
        """#(Thiết lập keyboard shortcuts cho các thao tác)"""
        # ESC để đóng
        QtGui.QShortcut(Qt.Key_Escape, self, self.close)

        # Zoom shortcuts
        QtGui.QShortcut(QtGui.QKeySequence("Ctrl++"), self, self._zoom_in)
        QtGui.QShortcut(QtGui.QKeySequence("Ctrl+="), self, self._zoom_in)  # Cho bàn phím US
        QtGui.QShortcut(QtGui.QKeySequence("Ctrl+-"), self, self._zoom_out)
        QtGui.QShortcut(QtGui.QKeySequence("Ctrl+0"), self, self._fit_to_window)
        QtGui.QShortcut(QtGui.QKeySequence("Ctrl+1"), self, self._actual_size)

        # Pan với arrow keys
        QtGui.QShortcut(Qt.Key_Left, self, lambda: self._pan_with_keys(-20, 0))
        QtGui.QShortcut(Qt.Key_Right, self, lambda: self._pan_with_keys(20, 0))
        QtGui.QShortcut(Qt.Key_Up, self, lambda: self._pan_with_keys(0, -20))
        QtGui.QShortcut(Qt.Key_Down, self, lambda: self._pan_with_keys(0, 20))

    def _zoom_in(self):
        """#(Phóng to ảnh với giới hạn max_zoom)"""
        new_zoom = min(self.current_zoom + self.zoom_step, self.max_zoom)
        self._set_zoom(new_zoom)

    def _zoom_out(self):
        """#(Thu nhỏ ảnh với giới hạn min_zoom)"""
        new_zoom = max(self.current_zoom - self.zoom_step, self.min_zoom)
        self._set_zoom(new_zoom)

    def _zoom_slider_changed(self, value):
        """#(Xử lý thay đổi zoom từ slider)"""
        new_zoom = value / 100.0
        self._set_zoom(new_zoom, update_slider=False)

    def _set_zoom(self, zoom_level, update_slider=True):
        """#(Áp dụng mức zoom mới và cập nhật UI)"""
        self.current_zoom = max(self.min_zoom, min(zoom_level, self.max_zoom))

        # Update slider if needed
        if update_slider:
            self.zoom_slider.blockSignals(True)
            self.zoom_slider.setValue(int(self.current_zoom * 100))
            self.zoom_slider.blockSignals(False)

        # Scale pixmap
        if self.original_pixmap and not self.original_pixmap.isNull():
            new_size = self.original_pixmap.size() * self.current_zoom
            scaled_pixmap = self.original_pixmap.scaled(
                new_size, Qt.KeepAspectRatio, Qt.SmoothTransformation
            )
            self.image_label.setPixmap(scaled_pixmap)
            self.image_label.resize(new_size)

        self._update_status()

    def _fit_to_window(self):
        """#(Điều chỉnh zoom để ảnh vừa khít cửa sổ)"""
        if not self.original_pixmap or self.original_pixmap.isNull():
            return

        # Kích thước available (trừ toolbar và status bar)
        available_size = self.scroll_area.viewport().size()
        image_size = self.original_pixmap.size()

        # Tính scale factor để fit
        scale_w = available_size.width() / image_size.width()
        scale_h = available_size.height() / image_size.height()
        scale_factor = min(scale_w, scale_h) * 0.9  # 0.9 để có chút margin

        self._set_zoom(scale_factor)

    def _actual_size(self):
        """#(Hiển thị ảnh ở kích thước gốc 100%)"""
        self._set_zoom(1.0)

    def _pan_with_keys(self, dx, dy):
        """#(Di chuyển ảnh bằng phím mũi tên)"""
        h_scroll = self.scroll_area.horizontalScrollBar()
        v_scroll = self.scroll_area.verticalScrollBar()

        h_scroll.setValue(h_scroll.value() + dx)
        v_scroll.setValue(v_scroll.value() + dy)

    def _mouse_press_event(self, event):
        """#(Xử lý nhấn chuột để bắt đầu pan)"""
        if event.button() == Qt.LeftButton:
            self.is_panning = True
            self.last_pan_point = event.pos()
            self.image_label.setCursor(Qt.ClosedHandCursor)

    def _mouse_move_event(self, event):
        """#(Xử lý di chuyển chuột để pan ảnh)"""
        if self.is_panning and (event.buttons() & Qt.LeftButton):
            delta = event.pos() - self.last_pan_point

            h_scroll = self.scroll_area.horizontalScrollBar()
            v_scroll = self.scroll_area.verticalScrollBar()

            h_scroll.setValue(h_scroll.value() - delta.x())
            v_scroll.setValue(v_scroll.value() - delta.y())

            self.last_pan_point = event.pos()

    def _mouse_release_event(self, event):
        """#(Xử lý thả chuột để kết thúc pan)"""
        if event.button() == Qt.LeftButton:
            self.is_panning = False
            self.image_label.setCursor(Qt.OpenHandCursor)

    def _update_status(self):
        """#(Cập nhật status bar với thông tin chi tiết)"""
        if not self.original_pixmap or self.original_pixmap.isNull():
            return

        original_size = self.original_pixmap.size()
        current_size = self.image_label.size()
        zoom_percent = int(self.current_zoom * 100)

        # Update zoom label in toolbar
        self.zoom_label.setText(f"{zoom_percent}%")

        # Update main status
        main_status = (
            f"🖼️ Gốc: {original_size.width()}×{original_size.height()} | "
            f"Hiển thị: {current_size.width()}×{current_size.height()} | "
            f"Zoom: {zoom_percent}%"
        )

        self.main_status_label.setText(main_status)
    def wheelEvent(self, event):
        """#(Xử lý zoom bằng con lăn chuột với Ctrl)"""
        modifiers = QtWidgets.QApplication.keyboardModifiers()
        if modifiers == Qt.ControlModifier:
            # Zoom với mouse wheel
            delta = event.angleDelta().y()
            if delta > 0:
                self._zoom_in()
            else:
                self._zoom_out()
            event.accept()
        else:
            # Scroll bình thường
            super().wheelEvent(event)

    def keyPressEvent(self, event):
        """#(Xử lý các phím tắt bổ sung)"""
        if event.key() == Qt.Key_Space:
            # Space để fit to window
            self._fit_to_window()
            event.accept()
        elif event.key() == Qt.Key_F:
            # F để fullscreen toggle (nếu cần)
            event.accept()
        else:
            super().keyPressEvent(event)
