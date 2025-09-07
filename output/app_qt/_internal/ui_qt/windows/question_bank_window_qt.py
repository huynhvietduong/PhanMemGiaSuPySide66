from __future__ import annotations
import json
import os
import base64
import io
from typing import List, Dict, Optional
from datetime import datetime

# Import PySide6
from PySide6 import QtCore, QtGui, QtWidgets
from PySide6.QtCore import Qt, QByteArray, QBuffer, QIODevice, Signal
from PySide6.QtWidgets import (
    QMenu, QFileDialog, QMessageBox, QDialog,
    QVBoxLayout, QHBoxLayout, QFormLayout,
    QWidget, QLabel, QLineEdit, QTextEdit,
    QPushButton, QComboBox, QSpinBox,
    QToolBar, QGroupBox, QSplitter,
    QStackedWidget, QScrollArea,
    QTableWidget, QTableWidgetItem,
    QTreeWidget, QTreeWidgetItem,
    QHeaderView, QAbstractItemView
)
from PySide6.QtGui import (
    QKeySequence, QShortcut, QPixmap, QImage,
    QTextDocument, QTextCursor, QColor, QBrush,
    QAction, QIcon
)
from PySide6.QtPrintSupport import QPrintPreviewDialog, QPrinter

class QuestionBankWindowQt(QtWidgets.QWidget):
    # ========== NHÓM 1: KHỞI TẠO VÀ THIẾT LẬP ========== #
    def __init__(self, db_manager, parent=None):
        super().__init__(parent)
        self.db = db_manager
        self.setObjectName("QuestionBankWindowQt")
        self.setWindowTitle("Ngân hàng câu hỏi")
        self.resize(1200, 680)

        # Đảm bảo bảng tồn tại
        self._ensure_tables()

        self.current_question_id: int | None = None
        self.tree_nodes: Dict[str, int] = {}

        root = QtWidgets.QVBoxLayout(self)

        # Tạo toolbar chính
        main_toolbar = QtWidgets.QToolBar()
        main_toolbar.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        main_toolbar.setMovable(False)
        main_toolbar.setStyleSheet("""
            QToolBar { 
                background: #f8f9fa; 
                border: 1px solid #dee2e6; 
                spacing: 8px; 
                padding: 4px;
            }
            QToolBar::separator { 
                background: #dee2e6; 
                width: 2px; 
                margin: 0 4px;
            }
        """)
        root.addWidget(main_toolbar)

        # Nhóm quản lý cây
        toggle_action = main_toolbar.addAction("🌲 Ẩn/Hiện cây")
        toggle_action.triggered.connect(self.toggle_tree_panel)

        manage_action = main_toolbar.addAction("⚙️ Quản lý cây")
        manage_action.triggered.connect(self.open_tree_manager)

        main_toolbar.addSeparator()

        # Nhóm tìm kiếm
        search_widget = QtWidgets.QWidget()
        search_layout = QtWidgets.QHBoxLayout(search_widget)
        search_layout.setContentsMargins(0, 0, 0, 0)

        search_layout.addWidget(QtWidgets.QLabel("🔍"))

        self.search_edit = QtWidgets.QLineEdit()
        self.search_edit.setPlaceholderText("Tìm kiếm câu hỏi...")
        self.search_edit.setMinimumWidth(200)
        self.search_edit.setStyleSheet("padding: 4px; border: 1px solid #ced4da; border-radius: 4px;")
        search_layout.addWidget(self.search_edit)

        main_toolbar.addWidget(search_widget)

        search_action = main_toolbar.addAction("Tìm")
        search_action.triggered.connect(self.search_questions)

        main_toolbar.addSeparator()

        # Thêm nút tạo câu hỏi nổi bật
        add_question_action = main_toolbar.addAction("➕ Thêm câu hỏi")
        add_question_action.triggered.connect(self.open_add_question_dialog)
        # Style cho nút
        add_btn_widget = main_toolbar.widgetForAction(add_question_action)
        if add_btn_widget:
            add_btn_widget.setStyleSheet("""
                QToolButton {
                    background: #28a745;
                    color: white;
                    font-weight: bold;
                    padding: 6px 12px;
                    border-radius: 4px;
                }
                QToolButton:hover {
                    background: #218838;
                }
            """)
        # Nhóm tạo mới
        new_action = main_toolbar.addAction("➕ Tạo mới")
        new_action.triggered.connect(self.new_question)

        template_action = main_toolbar.addAction("📋 Template")
        template_action.triggered.connect(self.show_template_dialog)

        main_toolbar.addSeparator()

        # Nhóm import/export
        import_action = main_toolbar.addAction("📥 Import Word")
        import_action.triggered.connect(self.import_from_word)

        export_action = main_toolbar.addAction("📤 Export Word")
        export_action.triggered.connect(self.export_to_word)

        export_pdf_action = main_toolbar.addAction("📄 Export PDF")
        export_pdf_action.triggered.connect(self.export_to_pdf)

        main_toolbar.addSeparator()

        # Toolbar phụ cho filters
        filter_toolbar = QtWidgets.QToolBar()
        filter_toolbar.setStyleSheet("QToolBar { background: #e9ecef; border: 1px solid #dee2e6; }")
        root.addWidget(filter_toolbar)

        self._create_filter_controls(filter_toolbar)

        # Splitter 3 cột
        split = QtWidgets.QSplitter(Qt.Horizontal)
        root.addWidget(split, 1)

        # --- Cột trái: Cây ---
        left = QtWidgets.QWidget()
        left_l = QtWidgets.QVBoxLayout(left)
        left_l.setContentsMargins(6, 6, 6, 6)

        left_l.addWidget(QtWidgets.QLabel("Cây thư mục"))
        self.tree = QtWidgets.QTreeWidget()
        self.tree.setHeaderHidden(True)
        self.tree.itemSelectionChanged.connect(self.on_tree_select)
        left_l.addWidget(self.tree, 1)

        split.addWidget(left)

        # --- Cột giữa: Danh sách câu hỏi ---
        mid = QtWidgets.QWidget()
        mid_l = QtWidgets.QVBoxLayout(mid)
        mid_l.setContentsMargins(6, 6, 6, 6)

        mid_l.addWidget(QtWidgets.QLabel("Danh sách câu hỏi"))

        # Bảng câu hỏi
        self.q_table = QtWidgets.QTableWidget(0, 8)
        headers = ["☑️", "ID", "Nội dung", "Số đáp án", "Đáp án đúng", "Dạng", "Mức độ", "🏷️"]
        self.q_table.setHorizontalHeaderLabels(headers)

        # Cấu hình resize mode
        header = self.q_table.horizontalHeader()
        header.setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeToContents)
        header.setSectionResizeMode(1, QtWidgets.QHeaderView.ResizeToContents)
        header.setSectionResizeMode(2, QtWidgets.QHeaderView.Stretch)
        header.setSectionResizeMode(3, QtWidgets.QHeaderView.ResizeToContents)
        header.setSectionResizeMode(4, QtWidgets.QHeaderView.ResizeToContents)
        header.setSectionResizeMode(5, QtWidgets.QHeaderView.ResizeToContents)
        header.setSectionResizeMode(6, QtWidgets.QHeaderView.ResizeToContents)
        header.setSectionResizeMode(7, QtWidgets.QHeaderView.ResizeToContents)

        self.q_table.setSortingEnabled(True)
        self.q_table.setContextMenuPolicy(Qt.CustomContextMenu)
        self.q_table.customContextMenuRequested.connect(self.show_table_context_menu)
        self.q_table.setAlternatingRowColors(True)
        self.q_table.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        self.q_table.setSelectionMode(QtWidgets.QAbstractItemView.ExtendedSelection)

        self.q_table.setStyleSheet("""
            QTableWidget {
                gridline-color: #e9ecef;
                background-color: white;
                alternate-background-color: #f8f9fa;
            }
            QTableWidget::item {
                padding: 8px;
                border-bottom: 1px solid #e9ecef;
            }
            QTableWidget::item:selected {
                background-color: #007bff;
                color: white;
            }
        """)

        mid_l.addWidget(self.q_table, 1)
        split.addWidget(mid)

        self.q_table.itemSelectionChanged.connect(self.on_question_select)

        # --- Cột phải: Panel chi tiết với tabs ---
        right_tabs = QtWidgets.QTabWidget()
        right_tabs.setStyleSheet("""
            QTabWidget::pane {
                border: 1px solid #dee2e6;
                background: white;
            }
            QTabBar::tab {
                background: #f8f9fa;
                padding: 8px 16px;
                margin-right: 2px;
                border: 1px solid #dee2e6;
                border-bottom: none;
            }
            QTabBar::tab:selected {
                background: white;
                border-bottom: 1px solid white;
            }
        """)

        # Tab 1: Chỉnh sửa câu hỏi
        edit_tab = QtWidgets.QWidget()
        edit_layout = QtWidgets.QVBoxLayout(edit_tab)
        edit_layout.setContentsMargins(10, 10, 10, 10)

        self._create_edit_tab_content(edit_layout)
        right_tabs.addTab(edit_tab, "✏️ Chỉnh sửa")

        # Tab 2: Xem trước
        preview_tab = QtWidgets.QWidget()
        preview_layout = QtWidgets.QVBoxLayout(preview_tab)
        preview_layout.setContentsMargins(10, 10, 10, 10)

        self._create_preview_tab_content(preview_layout)
        right_tabs.addTab(preview_tab, "👁️ Xem trước")

        # Tab 3: Thống kê
        stats_tab = QtWidgets.QWidget()
        stats_layout = QtWidgets.QVBoxLayout(stats_tab)
        stats_layout.setContentsMargins(10, 10, 10, 10)

        self._create_stats_tab_content(stats_layout)
        right_tabs.addTab(stats_tab, "📊 Thống kê")

        # Tab 4: Lịch sử
        history_tab = QtWidgets.QWidget()
        history_layout = QtWidgets.QVBoxLayout(history_tab)
        history_layout.setContentsMargins(10, 10, 10, 10)

        self._create_history_tab_content(history_layout)
        right_tabs.addTab(history_tab, "📜 Lịch sử")

        split.addWidget(right_tabs)
        split.setSizes([240, 520, 440])

        # Init dữ liệu
        self.refresh_tree()
        self.load_available_subjects()
        self.load_available_grades()

        # Signal cho combobox
        self.subject_cb.currentIndexChanged.connect(self.load_available_topics)
        self.grade_cb.currentIndexChanged.connect(self.load_available_topics)
        self.topic_cb.currentIndexChanged.connect(self.load_available_types)

        # Keyboard shortcuts
        QShortcut("Ctrl+N", self, self.new_question)
        QShortcut("Ctrl+S", self, self.save_question)
        QShortcut("Ctrl+F", self, self.focus_search)
        QShortcut("Delete", self, self.delete_question)
        QShortcut("F5", self, self.refresh_all)

        self.setAcceptDrops(True)
        self.q_table.setDragDropMode(QtWidgets.QAbstractItemView.InternalMove)

        self._setup_tree_management()
    def open_add_question_dialog(self):
        """Mở dialog thêm câu hỏi mới"""
        tree_id = self._current_tree_id()
        if not tree_id:
            QtWidgets.QMessageBox.warning(self, "Chưa chọn thư mục",
                                          "Vui lòng chọn thư mục trong cây để lưu câu hỏi")
            return

        dialog = QuestionEditDialog(self.db, tree_id=tree_id, parent=self)
        if dialog.exec() == QtWidgets.QDialog.Accepted:
            # Refresh danh sách câu hỏi
            self.on_tree_select()

    def _ensure_tables(self):
        """Tạo cấu trúc bảng cơ bản"""
        # Tạo bảng question_bank với đầy đủ cột từ đầu
        self.db.execute_query("""
            CREATE TABLE IF NOT EXISTS question_bank (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                content_text TEXT,
                options TEXT,
                correct TEXT,
                tree_id INTEGER,
                difficulty_level TEXT,
                content_type TEXT DEFAULT 'text',
                content_data BLOB,
                answer_type TEXT DEFAULT 'text',
                answer_data BLOB,
                created_date TEXT DEFAULT CURRENT_TIMESTAMP,
                modified_date TEXT DEFAULT CURRENT_TIMESTAMP
            );
        """)

        # Kiểm tra và thêm cột mới nếu bảng đã tồn tại từ trước
        c = self.db.conn.cursor()
        c.execute("PRAGMA table_info(question_bank)")
        existing_columns = [column[1] for column in c.fetchall()]

        # Thêm cột content_type nếu chưa có
        if 'content_type' not in existing_columns:
            try:
                self.db.execute_query("ALTER TABLE question_bank ADD COLUMN content_type TEXT DEFAULT 'text'")
                print("✅ Đã thêm cột content_type")
            except Exception as e:
                if "duplicate column name" not in str(e).lower():
                    print(f"⚠️ Lỗi thêm cột content_type: {e}")

        # Thêm cột content_data nếu chưa có
        if 'content_data' not in existing_columns:
            try:
                self.db.execute_query("ALTER TABLE question_bank ADD COLUMN content_data BLOB")
                print("✅ Đã thêm cột content_data")
            except Exception as e:
                if "duplicate column name" not in str(e).lower():
                    print(f"⚠️ Lỗi thêm cột content_data: {e}")

        # Thêm cột answer_type nếu chưa có
        if 'answer_type' not in existing_columns:
            try:
                self.db.execute_query("ALTER TABLE question_bank ADD COLUMN answer_type TEXT DEFAULT 'text'")
                print("✅ Đã thêm cột answer_type")
            except Exception as e:
                if "duplicate column name" not in str(e).lower():
                    print(f"⚠️ Lỗi thêm cột answer_type: {e}")

        # Thêm cột answer_data nếu chưa có
        if 'answer_data' not in existing_columns:
            try:
                self.db.execute_query("ALTER TABLE question_bank ADD COLUMN answer_data BLOB")
                print("✅ Đã thêm cột answer_data")
            except Exception as e:
                if "duplicate column name" not in str(e).lower():
                    print(f"⚠️ Lỗi thêm cột answer_data: {e}")

        # Tạo bảng exercise_tree
        self.db.execute_query("""
            CREATE TABLE IF NOT EXISTS exercise_tree (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                parent_id INTEGER,
                name TEXT NOT NULL,
                level TEXT NOT NULL,
                description TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (parent_id) REFERENCES exercise_tree (id)
            )
        """)

        # Tạo bảng question_tags
        self.db.execute_query("""
            CREATE TABLE IF NOT EXISTS question_tags (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                question_id INTEGER,
                tag_name TEXT,
                color TEXT DEFAULT '#3498db',
                created_date TEXT DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (question_id) REFERENCES question_bank(id) ON DELETE CASCADE,
                UNIQUE(question_id, tag_name)
            );
        """)

        # Tạo bảng question_history
        self.db.execute_query("""
            CREATE TABLE IF NOT EXISTS question_history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                question_id INTEGER,
                action_type TEXT,
                old_content TEXT,
                new_content TEXT,
                changed_date TEXT DEFAULT CURRENT_TIMESTAMP,
                user_note TEXT,
                FOREIGN KEY (question_id) REFERENCES question_bank(id) ON DELETE CASCADE
            );
        """)

        # Tạo indexes
        self.db.execute_query("CREATE INDEX IF NOT EXISTS idx_question_tree_id ON question_bank(tree_id)")
        self.db.execute_query("CREATE INDEX IF NOT EXISTS idx_question_tags_question_id ON question_tags(question_id)")
    def _insert_sample_tree_data(self):
        """Thêm dữ liệu mẫu cho cây thư mục"""
        sample_data = [
            (None, "Toán", "Môn", "Môn Toán học"),
            (None, "Lý", "Môn", "Môn Vật lý"),
            (None, "Hóa", "Môn", "Môn Hóa học"),
            (1, "Lớp 10", "Lớp", "Toán lớp 10"),
            (1, "Lớp 11", "Lớp", "Toán lớp 11"),
            (1, "Lớp 12", "Lớp", "Toán lớp 12"),
            (4, "Mệnh đề - Tập hợp", "Chủ đề", "Chương 1"),
            (4, "Hàm số", "Chủ đề", "Chương 2"),
            (7, "Nhận biết", "Mức độ", "Câu hỏi nhận biết cơ bản"),
            (7, "Thông hiểu", "Mức độ", "Câu hỏi thông hiểu"),
            (7, "Vận dụng", "Mức độ", "Câu hỏi vận dụng"),
        ]

        for parent_id, name, level, description in sample_data:
            self.db.execute_query(
                "INSERT INTO exercise_tree (parent_id, name, level, description) VALUES (?, ?, ?, ?)",
                (parent_id, name, level, description)
            )

    # ========== NHÓM 2: TẠO GIAO DIỆN ========== #
    def _create_filter_controls(self, toolbar):
        """Tạo các combobox filter"""
        toolbar.addWidget(QtWidgets.QLabel("Môn:"))
        self.subject_cb = QtWidgets.QComboBox()
        self.subject_cb.setMinimumWidth(120)
        toolbar.addWidget(self.subject_cb)

        toolbar.addWidget(QtWidgets.QLabel("Lớp:"))
        self.grade_cb = QtWidgets.QComboBox()
        self.grade_cb.setMinimumWidth(100)
        toolbar.addWidget(self.grade_cb)

        toolbar.addWidget(QtWidgets.QLabel("Chủ đề:"))
        self.topic_cb = QtWidgets.QComboBox()
        self.topic_cb.setMinimumWidth(150)
        toolbar.addWidget(self.topic_cb)

        toolbar.addWidget(QtWidgets.QLabel("Dạng:"))
        self.type_cb = QtWidgets.QComboBox()
        self.type_cb.setMinimumWidth(120)
        toolbar.addWidget(self.type_cb)

        toolbar.addWidget(QtWidgets.QLabel("Mức độ:"))
        self.level_cb = QtWidgets.QComboBox()
        self.level_cb.addItems(["", "Nhận biết", "Thông hiểu", "Vận dụng", "Vận dụng cao"])
        self.level_cb.setMinimumWidth(120)
        toolbar.addWidget(self.level_cb)

        toolbar.addSeparator()

        filter_btn = toolbar.addAction("🔽 Lọc")
        filter_btn.triggered.connect(self.filter_by_combobox)

        clear_filter_btn = toolbar.addAction("🔄 Xóa lọc")
        clear_filter_btn.triggered.connect(self.clear_filters)
    def _create_edit_tab_content(self, layout):
        """Tạo nội dung tab chỉnh sửa với hỗ trợ đa media"""
        # Container cho nội dung
        self.content_container = QtWidgets.QStackedWidget()

        # Text editor
        self.content_text = QtWidgets.QTextEdit()
        self.content_text.setMinimumHeight(150)
        self.content_container.addWidget(self.content_text)

        # Image viewer cho review
        self.content_image_viewer = ImageViewer()
        self.content_container.addWidget(self.content_image_viewer)

        # PDF viewer cho review
        self.content_pdf_viewer = PDFViewer()
        self.content_container.addWidget(self.content_pdf_viewer)

        layout.addWidget(QtWidgets.QLabel("Nội dung câu hỏi:"))
        layout.addWidget(self.content_container)

        # Tương tự cho đáp án
        layout.addWidget(QtWidgets.QLabel("Đáp án:"))
        self.answer_container = QtWidgets.QStackedWidget()

        self.answer_text = QtWidgets.QTextEdit()
        self.answer_text.setMaximumHeight(100)
        self.answer_container.addWidget(self.answer_text)

        self.answer_image_viewer = ImageViewer()
        self.answer_container.addWidget(self.answer_image_viewer)

        layout.addWidget(self.answer_container)

        # Tags (giữ nguyên)
        tags_group = QtWidgets.QGroupBox("🏷️ Thẻ")
        tags_layout = QtWidgets.QHBoxLayout(tags_group)

        self.tags_edit = QtWidgets.QLineEdit()
        self.tags_edit.setPlaceholderText("Nhập thẻ, phân cách bằng dấu phẩy")
        tags_layout.addWidget(self.tags_edit)

        layout.addWidget(tags_group)

        # Buttons
        buttons_layout = QtWidgets.QHBoxLayout()

        self.btn_edit = QtWidgets.QPushButton("✏️ Sửa câu hỏi")
        self.btn_edit.clicked.connect(self.edit_current_question)
        self.btn_edit.setStyleSheet("QPushButton { background: #007bff; color: white; padding: 8px 16px; }")

        self.btn_delete = QtWidgets.QPushButton("🗑️ Xóa")
        self.btn_delete.clicked.connect(self.delete_question)
        self.btn_delete.setStyleSheet("QPushButton { background: #dc3545; color: white; padding: 8px 16px; }")

        buttons_layout.addWidget(self.btn_edit)
        buttons_layout.addWidget(self.btn_delete)

        layout.addLayout(buttons_layout)
    def _create_preview_tab_content(self, layout):
        """Tạo nội dung tab preview"""
        layout.addWidget(QtWidgets.QLabel("📋 Xem trước câu hỏi:"))

        self.preview_widget = QtWidgets.QTextEdit()
        self.preview_widget.setReadOnly(True)
        self.preview_widget.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 2px solid #e9ecef;
                border-radius: 8px;
                padding: 16px;
                font-size: 14px;
            }
        """)
        layout.addWidget(self.preview_widget)
    def _create_stats_tab_content(self, layout):
        """Tạo nội dung tab thống kê"""
        layout.addWidget(QtWidgets.QLabel("📊 Thống kê ngân hàng câu hỏi:"))

        self.stats_widget = QtWidgets.QTextEdit()
        self.stats_widget.setReadOnly(True)
        layout.addWidget(self.stats_widget)

        update_stats_btn = QtWidgets.QPushButton("🔄 Cập nhật thống kê")
        update_stats_btn.clicked.connect(self.update_statistics)
        layout.addWidget(update_stats_btn)
    def _create_history_tab_content(self, layout):
        """Tạo nội dung tab lịch sử"""
        layout.addWidget(QtWidgets.QLabel("📜 Lịch sử chỉnh sửa:"))

        self.history_table = QtWidgets.QTableWidget(0, 4)
        self.history_table.setHorizontalHeaderLabels(["Thời gian", "Hành động", "Nội dung cũ", "Nội dung mới"])
        header = self.history_table.horizontalHeader()
        header.setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeToContents)
        header.setSectionResizeMode(1, QtWidgets.QHeaderView.ResizeToContents)
        header.setSectionResizeMode(2, QtWidgets.QHeaderView.Stretch)
        header.setSectionResizeMode(3, QtWidgets.QHeaderView.Stretch)

        layout.addWidget(self.history_table)

    # ========== NHÓM 3: QUẢN LÝ CÂY THƯ MỤC ========== #
    def refresh_tree(self):
        """Làm mới cây thư mục"""
        try:
            self.tree.clear()
            self.tree_nodes.clear()

            rows = self.db.execute_query(
                "SELECT id, parent_id, name, level FROM exercise_tree ORDER BY parent_id, level, name",
                fetch='all'
            ) or []

            if not rows:
                self._insert_sample_tree_data()
                rows = self.db.execute_query(
                    "SELECT id, parent_id, name, level FROM exercise_tree ORDER BY parent_id, level, name",
                    fetch='all'
                ) or []

            children: Dict[int | None, list] = {}
            for r in rows:
                children.setdefault(r["parent_id"], []).append(r)

            def build(parent_db_id: int | None, parent_item: QtWidgets.QTreeWidgetItem | None):
                for node in children.get(parent_db_id, []):
                    icon_text = self._get_level_icon(node["level"])
                    item_text = f"{icon_text} {node['name']}"

                    item = QtWidgets.QTreeWidgetItem([item_text])
                    item.setData(0, Qt.UserRole, node["id"])
                    item.setToolTip(0, f"Level: {node['level']}\nID: {node['id']}")

                    self.tree_nodes[str(id(item))] = node["id"]

                    if parent_item is None:
                        self.tree.addTopLevelItem(item)
                    else:
                        parent_item.addChild(item)
                    build(node["id"], item)

            build(None, None)
            self.tree.expandAll()

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể tải cây thư mục: {e}")

    def _get_level_icon(self, level: str) -> str:
        """Trả về icon cho từng level"""
        icons = {
            "Môn": "📚",
            "Lớp": "🎓",
            "Chủ đề": "📖",
            "Dạng": "📝",
            "Mức độ": "⭐"
        }
        return icons.get(level, "📁")

    def on_tree_select(self):
        """Xử lý khi chọn node trên cây"""
        items = self.tree.selectedItems()
        if not items:
            return
        tree_id = items[0].data(0, Qt.UserRole)
        if not tree_id:
            return

        rows = self.db.execute_query("SELECT * FROM question_bank WHERE tree_id=?", (tree_id,), fetch="all") or []
        self._load_question_rows(rows)

    def toggle_tree_panel(self):
        """Ẩn/hiện panel cây"""
        w = self.tree.parentWidget()
        w.setVisible(not w.isVisible())

    def open_tree_manager(self):
        """Mở cửa sổ quản lý cây"""
        QtWidgets.QMessageBox.information(self, "Thông tin", "Chức năng quản lý cây đang phát triển.")
    def _setup_tree_management(self):
        """Thiết lập chức năng quản lý cây thư mục"""
        # Thêm context menu cho tree
        self.tree.setContextMenuPolicy(Qt.CustomContextMenu)
        self.tree.customContextMenuRequested.connect(self._show_tree_context_menu)

        # Thêm double-click để edit
        self.tree.itemDoubleClicked.connect(self._edit_tree_node)

        # Thêm keyboard shortcuts
        self._setup_tree_shortcuts()

    def _setup_tree_shortcuts(self):
        """Thiết lập keyboard shortcuts cho tree"""
        # F2 để edit node được chọn
        edit_shortcut = QShortcut(QKeySequence("F2"), self.tree)
        edit_shortcut.activated.connect(self._edit_selected_tree_node)

        # Delete để xóa node
        delete_shortcut = QShortcut(QKeySequence("Delete"), self.tree)
        delete_shortcut.activated.connect(self._delete_selected_tree_node)

        # Ctrl+N để thêm node mới
        add_shortcut = QShortcut(QKeySequence("Ctrl+Shift+N"), self.tree)
        add_shortcut.activated.connect(self._add_tree_node)

    def _show_tree_context_menu(self, position):
        """Hiển thị context menu cho tree"""
        item = self.tree.itemAt(position)

        menu = QtWidgets.QMenu(self)

        # Thêm node mới
        add_action = menu.addAction("➕ Thêm nhánh mới")
        add_action.triggered.connect(lambda: self._add_tree_node(item))

        if item:  # Nếu click vào node
            menu.addSeparator()

            # Thêm node con
            add_child_action = menu.addAction("📁 Thêm nhánh con")
            add_child_action.triggered.connect(lambda: self._add_child_node(item))

            # Sửa node
            edit_action = menu.addAction("✏️ Sửa tên nhánh")
            edit_action.triggered.connect(lambda: self._edit_tree_node(item))

            # Sao chép node
            copy_action = menu.addAction("📋 Sao chép nhánh")
            copy_action.triggered.connect(lambda: self._copy_tree_node(item))

            menu.addSeparator()

            # Xóa node
            delete_action = menu.addAction("🗑️ Xóa nhánh")
            delete_action.triggered.connect(lambda: self._delete_tree_node(item))

            menu.addSeparator()

            # Thông tin node
            info_action = menu.addAction("ℹ️ Thông tin")
            info_action.triggered.connect(lambda: self._show_node_info(item))

        # Làm mới cây
        menu.addSeparator()
        refresh_action = menu.addAction("🔄 Làm mới")
        refresh_action.triggered.connect(self.refresh_tree)

        # Hiển thị menu
        menu.exec(self.tree.mapToGlobal(position))

    def _add_tree_node(self, parent_item=None):
        """Thêm node mới"""
        try:
            dialog = TreeNodeDialog(self.db, mode="add", parent=self)

            # Nếu có parent item, set làm parent
            parent_id = None
            if parent_item:
                parent_id = parent_item.data(0, Qt.UserRole)
                if parent_id:
                    dialog.set_parent_id(parent_id)

            if dialog.exec() == QtWidgets.QDialog.Accepted:
                # Refresh tree sau khi thêm
                self.refresh_tree()

                # Tìm lại parent item sau khi refresh
                if parent_id:
                    self._expand_node_by_id(parent_id)

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể thêm node: {e}")

    def _add_child_node(self, parent_item):
        """Thêm node con"""
        if not parent_item:
            return

        parent_id = parent_item.data(0, Qt.UserRole)
        if not parent_id:
            return

        try:
            dialog = TreeNodeDialog(self.db, mode="add", parent=self)
            dialog.set_parent_id(parent_id)

            if dialog.exec() == QtWidgets.QDialog.Accepted:
                self.refresh_tree()
                # Tìm lại và expand parent sau khi refresh
                self._expand_node_by_id(parent_id)

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể thêm node con: {e}")

    def _edit_tree_node(self, item):
        """Sửa node"""
        if not item:
            return

        node_id = item.data(0, Qt.UserRole)
        if not node_id:
            return

        try:
            dialog = TreeNodeDialog(self.db, mode="edit", node_id=node_id, parent=self)

            if dialog.exec() == QtWidgets.QDialog.Accepted:
                self.refresh_tree()
                # Tìm lại và select node sau khi refresh
                self._select_node_by_id(node_id)

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể sửa node: {e}")

    def _edit_selected_tree_node(self):
        """Sửa node được chọn"""
        selected_items = self.tree.selectedItems()
        if selected_items:
            self._edit_tree_node(selected_items[0])

    def _copy_tree_node(self, item):
        """Sao chép node"""
        if not item:
            return

        node_id = item.data(0, Qt.UserRole)
        if not node_id:
            return

        try:
            # Lấy thông tin node gốc
            row = self.db.execute_query(
                "SELECT name, level, description, parent_id FROM exercise_tree WHERE id = ?",
                (node_id,), fetch="one"
            )

            if row:
                new_name = f"{row['name']} (Sao chép)"

                # Tạo node mới
                description = row.get('description', '') if row.get('description') else ''

                self.db.execute_query(
                    "INSERT INTO exercise_tree (parent_id, name, level, description) VALUES (?, ?, ?, ?)",
                    (row['parent_id'], new_name, row['level'], description)
                )

                self.refresh_tree()
                QtWidgets.QMessageBox.information(self, "Thành công", f"Đã sao chép '{new_name}'")

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể sao chép node: {e}")

    def _delete_tree_node(self, item):
        """Xóa node với xác nhận"""
        if not item:
            return

        node_id = item.data(0, Qt.UserRole)
        node_name = item.text(0)

        if not node_id:
            return

        try:
            # Kiểm tra node con
            children_count = self.db.execute_query(
                "SELECT COUNT(*) as count FROM exercise_tree WHERE parent_id = ?",
                (node_id,), fetch="one"
            )

            # Kiểm tra câu hỏi trong node
            questions_count = self.db.execute_query(
                "SELECT COUNT(*) as count FROM question_bank WHERE tree_id = ?",
                (node_id,), fetch="one"
            )

            warning_msg = f"Bạn có chắc muốn xóa nhánh '{node_name}'?"

            if children_count and children_count["count"] > 0:
                warning_msg += f"\n\n⚠️ Nhánh này có {children_count['count']} nhánh con."

            if questions_count and questions_count["count"] > 0:
                warning_msg += f"\n⚠️ Nhánh này chứa {questions_count['count']} câu hỏi."
                warning_msg += "\n\nTất cả dữ liệu sẽ bị xóa vĩnh viễn!"

            reply = QtWidgets.QMessageBox.question(
                self, "Xác nhận xóa",
                warning_msg,
                QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No,
                QtWidgets.QMessageBox.No
            )

            if reply == QtWidgets.QMessageBox.Yes:
                # Xóa node và tất cả con
                self.db.execute_query("DELETE FROM exercise_tree WHERE id = ?", (node_id,))
                self.refresh_tree()
                QtWidgets.QMessageBox.information(self, "Thành công", f"Đã xóa nhánh '{node_name}'")

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể xóa node: {e}")

    def _delete_selected_tree_node(self):
        """Xóa node được chọn"""
        selected_items = self.tree.selectedItems()
        if selected_items:
            self._delete_tree_node(selected_items[0])

    def _show_node_info(self, item):
        """Hiển thị thông tin node"""
        if not item:
            return

        node_id = item.data(0, Qt.UserRole)
        if not node_id:
            return

        try:
            # Lấy thông tin node
            node = self.db.execute_query(
                "SELECT * FROM exercise_tree WHERE id = ?",
                (node_id,), fetch="one"
            )

            if node:
                # Đếm số lượng con
                children_count = self.db.execute_query(
                    "SELECT COUNT(*) as count FROM exercise_tree WHERE parent_id = ?",
                    (node_id,), fetch="one"
                )["count"]

                # Đếm số câu hỏi
                questions_count = self.db.execute_query(
                    "SELECT COUNT(*) as count FROM question_bank WHERE tree_id = ?",
                    (node_id,), fetch="one"
                )["count"]

                info_text = f"""
                📁 THÔNG TIN NHÁNH
                
                ID: {node['id']}
                Tên: {node['name']}
                Cấp độ: {node['level']}
                Mô tả: {node.get('description', 'Không có')}
                
                📊 THỐNG KÊ:
                - Số nhánh con: {children_count}
                - Số câu hỏi: {questions_count}
                
                🕐 Ngày tạo: {node.get('created_at', 'Không rõ')}
                """

                QtWidgets.QMessageBox.information(self, "Thông tin nhánh", info_text)

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể lấy thông tin: {e}")

    def _expand_node_by_id(self, node_id):
        """Tìm và expand node theo ID"""
        try:
            root = self.tree.invisibleRootItem()
            self._find_and_expand_recursive(root, node_id)
        except Exception:
            pass

    def _find_and_expand_recursive(self, parent_item, target_id):
        """Đệ quy tìm và expand node"""
        for i in range(parent_item.childCount()):
            child = parent_item.child(i)
            if child and child.data(0, Qt.UserRole) == target_id:
                child.setExpanded(True)
                return True

            if self._find_and_expand_recursive(child, target_id):
                return True

        return False

    def _select_node_by_id(self, node_id):
        """Tìm và select node theo ID"""
        try:
            root = self.tree.invisibleRootItem()
            self._find_and_select_recursive(root, node_id)
        except Exception:
            pass

    def _find_and_select_recursive(self, parent_item, target_id):
        """Đệ quy tìm và select node"""
        for i in range(parent_item.childCount()):
            child = parent_item.child(i)
            if child and child.data(0, Qt.UserRole) == target_id:
                self.tree.setCurrentItem(child)
                return True

            if self._find_and_select_recursive(child, target_id):
                return True

        return False

    def open_tree_manager(self):
        """Mở cửa sổ quản lý cây nâng cao"""
        dialog = TreeManagerDialog(self.db, parent=self)
        dialog.exec()
        self.refresh_tree()
    # ========== NHÓM 4: QUẢN LÝ DANH SÁCH CÂU HỎI ========== #
    def _load_question_rows(self, rows):
        """Load danh sách câu hỏi vào bảng"""
        self.q_table.setRowCount(0)

        for r in rows:
            # Chuyển sqlite3.Row thành dict để dễ xử lý
            row_dict = dict(r)

            checkbox = QtWidgets.QCheckBox()
            checkbox.setChecked(False)

            # Lấy nội dung preview
            content_text = row_dict.get("content_text", "") or ""
            content_preview = content_text[:50].replace("\n", " ").strip()

            row_id = row_dict.get("id", 0)
            tree_id = row_dict.get("tree_id", 0)

            # Lấy loại nội dung (text, image, pdf, word)
            content_type = row_dict.get("content_type", "text")
            type_icon = {
                "text": "📝",
                "image": "🖼️",
                "pdf": "📄",
                "word": "📘"
            }.get(content_type, "📝")

            # Lấy thông tin từ options nếu có
            opts_json = row_dict.get("options", "[]") or "[]"
            try:
                opts = json.loads(opts_json) if opts_json else []
                so_dapan = len(opts)
            except:
                so_dapan = 0

            dap_an = row_dict.get("correct", "-") or "-"

            # Lấy chuỗi dạng/mức độ từ path
            path = self.get_tree_path(tree_id) if tree_id else []
            path_dict = {p["level"]: p["name"] for p in path}
            dang = path_dict.get("Dạng", "-")
            muc_do = path_dict.get("Mức độ", "-")

            # Lấy tags
            tags = self.db.execute_query(
                "SELECT tag_name FROM question_tags WHERE question_id=?",
                (row_id,), fetch="all"
            ) or []
            tags_text = ", ".join([dict(tag)["tag_name"] for tag in tags]) if tags else ""

            # Thêm row vào table
            row_idx = self.q_table.rowCount()
            self.q_table.insertRow(row_idx)

            self.q_table.setCellWidget(row_idx, 0, checkbox)
            self.q_table.setItem(row_idx, 1, QtWidgets.QTableWidgetItem(str(row_id)))
            self.q_table.setItem(row_idx, 2, QtWidgets.QTableWidgetItem(content_preview))
            self.q_table.setItem(row_idx, 3, QtWidgets.QTableWidgetItem(type_icon))  # Loại
            self.q_table.setItem(row_idx, 4, QtWidgets.QTableWidgetItem(str(so_dapan)))
            self.q_table.setItem(row_idx, 5, QtWidgets.QTableWidgetItem(dap_an))
            self.q_table.setItem(row_idx, 6, QtWidgets.QTableWidgetItem(dang))
            self.q_table.setItem(row_idx, 7, QtWidgets.QTableWidgetItem(muc_do))

            tags_item = QtWidgets.QTableWidgetItem(tags_text)
            if tags_text:
                tags_item.setBackground(QtGui.QColor("#e3f2fd"))
            self.q_table.setItem(row_idx, 8, tags_item)
    def on_question_select(self):
        """Load câu hỏi được chọn với hỗ trợ multi-media"""
        items = self.q_table.selectedItems()
        if not items:
            return
        row = items[0].row()
        item_text = self.q_table.item(row, 1).text()

        try:
            qid = int(item_text)
        except (ValueError, IndexError):
            return

        q = self.db.execute_query("SELECT * FROM question_bank WHERE id=?", (qid,), fetch="one")
        if not q:
            return

        self.current_question_id = qid

        # Load nội dung theo loại
        content_type = q.get("content_type", "text")

        if content_type == "text":
            self.content_container.setCurrentWidget(self.content_text)
            self.content_text.setPlainText(q.get("content_text", "") or "")
        elif content_type == "image":
            self.content_container.setCurrentWidget(self.content_image_viewer)
            # Load image từ content_data (BLOB)
            if q.get("content_data"):
                pixmap = QPixmap()
                pixmap.loadFromData(q["content_data"])
                self.content_image_viewer.set_pixmap(pixmap)
        elif content_type == "pdf":
            self.content_container.setCurrentWidget(self.content_pdf_viewer)
            # Load PDF info

        # Tương tự cho đáp án
        answer_type = q.get("answer_type", "text")
        if answer_type == "text":
            self.answer_container.setCurrentWidget(self.answer_text)
            self.answer_text.setPlainText(q.get("correct", "") or "")
        elif answer_type == "image":
            self.answer_container.setCurrentWidget(self.answer_image_viewer)
            if q.get("answer_data"):
                pixmap = QPixmap()
                pixmap.loadFromData(q["answer_data"])
                self.answer_image_viewer.set_pixmap(pixmap)

        # Load tags (giữ nguyên)
        if hasattr(self, 'tags_edit'):
            tags = self.db.execute_query(
                "SELECT tag_name FROM question_tags WHERE question_id=? ORDER BY tag_name",
                (qid,), fetch="all"
            ) or []
            tags_text = ", ".join([tag["tag_name"] for tag in tags])
            self.tags_edit.setText(tags_text)

        # Update preview
        self.update_preview()
    def show_table_context_menu(self, position):
        """Hiển thị context menu cho bảng"""
        if not self.q_table.itemAt(position):
            return

        menu = QtWidgets.QMenu(self)

        edit_action = menu.addAction("✏️ Chỉnh sửa")
        edit_action.triggered.connect(self.on_question_select)

        menu.addSeparator()

        export_menu = menu.addMenu("📤 Xuất")
        export_menu.addAction("Xuất ra Word").triggered.connect(self.export_to_word)
        export_menu.addAction("Xuất ra PDF").triggered.connect(self.export_to_pdf)

        menu.addSeparator()

        delete_action = menu.addAction("🗑️ Xóa")
        delete_action.triggered.connect(self.delete_question)

        menu.exec(self.q_table.mapToGlobal(position))

    def get_tree_path(self, tree_id: int) -> List[dict]:
        """Lấy đường dẫn của node"""
        path = []
        while tree_id:
            row = self.db.execute_query(
                "SELECT id, parent_id, name, level FROM exercise_tree WHERE id=?",
                (tree_id,), fetch="one"
            )
            if row:
                path.insert(0, row)
                tree_id = row["parent_id"]
            else:
                break
        return path

    # ========== NHÓM 5: LƯU/CẬP NHẬT/XÓA CÂU HỎI ========== #
    def save_question(self):
        """Lưu câu hỏi"""
        tree_id = self._current_tree_id()
        if not tree_id:
            QtWidgets.QMessageBox.warning(self, "Chưa chọn thư mục", "Vui lòng chọn vị trí lưu trong cây.")
            return

        content = self.content_text.toPlainText().strip()
        answer = self.answer_text.toPlainText().strip()

        if not content:
            QtWidgets.QMessageBox.warning(self, "Lỗi", "Nội dung câu hỏi không được để trống")
            return

        try:
            if self.current_question_id:
                # Cập nhật
                self.db.execute_query(
                    "UPDATE question_bank SET content_text=?, correct=?, tree_id=?, modified_date=CURRENT_TIMESTAMP WHERE id=?",
                    (content, answer, tree_id, self.current_question_id)
                )
                self._save_question_history(self.current_question_id, "UPDATE", "", content)
                QtWidgets.QMessageBox.information(self, "Thành công", "Đã cập nhật câu hỏi.")
            else:
                # Thêm mới
                new_id = self.db.execute_query(
                    "INSERT INTO question_bank(content_text, correct, tree_id) VALUES (?,?,?)",
                    (content, answer, tree_id)
                )
                self.current_question_id = new_id
                self._save_question_history(new_id, "CREATE", "", content)
                QtWidgets.QMessageBox.information(self, "Thành công", "Đã lưu câu hỏi mới.")

            # Save tags
            self._save_question_tags()

            # Reload
            rows = self.db.execute_query("SELECT * FROM question_bank WHERE tree_id=?", (tree_id,), fetch="all") or []
            self._load_question_rows(rows)

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể lưu: {e}")

    def delete_question(self):
        """Xóa câu hỏi"""
        if not self.current_question_id:
            QtWidgets.QMessageBox.warning(self, "Chưa chọn", "Vui lòng chọn câu hỏi để xóa.")
            return

        if QtWidgets.QMessageBox.question(self, "Xác nhận",
                                          "Bạn có chắc muốn xóa câu hỏi này?") != QtWidgets.QMessageBox.Yes:
            return

        try:
            self.db.execute_query("DELETE FROM question_bank WHERE id=?", (self.current_question_id,))
            self.clear_question_form()

            tree_id = self._current_tree_id()
            if tree_id:
                rows = self.db.execute_query("SELECT * FROM question_bank WHERE tree_id=?", (tree_id,),
                                             fetch="all") or []
                self._load_question_rows(rows)

            QtWidgets.QMessageBox.information(self, "Thành công", "Đã xóa câu hỏi.")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể xóa: {e}")

    def new_question(self):
        """Tạo câu hỏi mới"""
        self.clear_question_form()
        self.content_text.setFocus()

    def clear_question_form(self):
        """Xóa form"""
        self.current_question_id = None
        if hasattr(self, 'content_text'):
            self.content_text.clear()
        if hasattr(self, 'answer_text'):
            self.answer_text.clear()
        if hasattr(self, 'tags_edit'):
            self.tags_edit.clear()

    def _current_tree_id(self) -> int | None:
        """Lấy tree_id hiện tại"""
        items = self.tree.selectedItems()
        if not items:
            return None
        return items[0].data(0, Qt.UserRole)

    def _save_question_tags(self):
        """Lưu tags cho câu hỏi"""
        if not self.current_question_id or not hasattr(self, 'tags_edit'):
            return

        tags_text = self.tags_edit.text().strip()
        if not tags_text:
            return

        # Xóa tags cũ
        self.db.execute_query("DELETE FROM question_tags WHERE question_id=?", (self.current_question_id,))

        # Thêm tags mới
        tag_names = [tag.strip() for tag in tags_text.split(',') if tag.strip()]
        for tag_name in tag_names:
            try:
                self.db.execute_query(
                    "INSERT INTO question_tags(question_id, tag_name) VALUES (?,?)",
                    (self.current_question_id, tag_name)
                )
            except:
                pass

    def _save_question_history(self, question_id, action_type, old_content, new_content):
        """Lưu lịch sử thay đổi"""
        try:
            self.db.execute_query(
                "INSERT INTO question_history(question_id, action_type, old_content, new_content) VALUES (?,?,?,?)",
                (question_id, action_type, old_content, new_content)
            )
        except:
            pass

    def _load_question_history(self, question_id):
        """Load lịch sử thay đổi"""
        if not hasattr(self, 'history_table'):
            return

        history = self.db.execute_query(
            "SELECT * FROM question_history WHERE question_id=? ORDER BY changed_date DESC LIMIT 50",
            (question_id,), fetch="all"
        ) or []

        self.history_table.setRowCount(0)

        for h in history:
            row_idx = self.history_table.rowCount()
            self.history_table.insertRow(row_idx)

            time_str = h.get("changed_date", "")

            self.history_table.setItem(row_idx, 0, QtWidgets.QTableWidgetItem(time_str))
            self.history_table.setItem(row_idx, 1, QtWidgets.QTableWidgetItem(h.get("action_type", "")))

            old_content = (h.get("old_content", "") or "")[:100]
            new_content = (h.get("new_content", "") or "")[:100]

            self.history_table.setItem(row_idx, 2, QtWidgets.QTableWidgetItem(old_content))
            self.history_table.setItem(row_idx, 3, QtWidgets.QTableWidgetItem(new_content))

    def edit_current_question(self):
        """Mở dialog edit câu hỏi hiện tại"""
        if not self.current_question_id:
            QtWidgets.QMessageBox.warning(self, "Chưa chọn câu hỏi",
                                          "Vui lòng chọn câu hỏi để chỉnh sửa")
            return

        tree_id = self._current_tree_id()
        dialog = QuestionEditDialog(self.db, tree_id=tree_id,
                                    question_id=self.current_question_id, parent=self)
        if dialog.exec() == QtWidgets.QDialog.Accepted:
            self.on_tree_select()
    # ========== NHÓM 6: TÌM KIẾM VÀ LỌC ========== #
    def search_questions(self):
        """Tìm kiếm câu hỏi"""
        keyword = (self.search_edit.text() or "").strip().lower()
        if not keyword:
            self.on_tree_select()
            return

        items = self.tree.selectedItems()
        if not items:
            QtWidgets.QMessageBox.warning(self, "Chưa chọn", "Hãy chọn thư mục để tìm.")
            return

        root_id = items[0].data(0, Qt.UserRole)
        all_ids = self.get_all_subtree_ids(root_id)
        if not all_ids:
            return

        placeholders = ",".join(["?"] * len(all_ids))
        query = f"SELECT * FROM question_bank WHERE tree_id IN ({placeholders})"
        rows = self.db.execute_query(query, tuple(all_ids), fetch="all") or []

        # Filter theo keyword
        rows = [r for r in rows if keyword in (r.get("content_text", "") or "").lower()]
        self._load_question_rows(rows)

    def get_all_subtree_ids(self, root_id: int) -> List[int]:
        """Lấy tất cả ID con"""
        ids = [root_id]
        children = self.db.execute_query("SELECT id FROM exercise_tree WHERE parent_id=?", (root_id,),
                                         fetch="all") or []
        for c in children:
            ids.extend(self.get_all_subtree_ids(c["id"]))
        return ids

    def focus_search(self):
        """Focus vào ô tìm kiếm"""
        self.search_edit.setFocus()
        self.search_edit.selectAll()

    def filter_by_combobox(self):
        """Lọc theo combobox"""
        # Simplified filtering logic
        self.on_tree_select()

    def clear_filters(self):
        """Xóa bộ lọc"""
        self.subject_cb.setCurrentIndex(0)
        self.grade_cb.setCurrentIndex(0)
        self.topic_cb.setCurrentIndex(0)
        self.type_cb.setCurrentIndex(0)
        self.level_cb.setCurrentIndex(0)
        self.on_tree_select()

    # ========== NHÓM 7: LOAD DỮ LIỆU COMBOBOX ========== #
    def load_available_subjects(self):
        """Load danh sách môn"""
        rows = self.db.execute_query(
            "SELECT DISTINCT name FROM exercise_tree WHERE level='Môn' ORDER BY name ASC",
            fetch="all"
        ) or []
        self.subject_cb.clear()
        self.subject_cb.addItem("")
        for r in rows:
            self.subject_cb.addItem(r["name"])

    def load_available_grades(self):
        """Load danh sách lớp"""
        rows = self.db.execute_query(
            "SELECT DISTINCT name FROM exercise_tree WHERE level='Lớp' ORDER BY name ASC",
            fetch="all"
        ) or []
        self.grade_cb.clear()
        self.grade_cb.addItem("")
        for r in rows:
            self.grade_cb.addItem(r["name"])

    def load_available_topics(self):
        """Load danh sách chủ đề"""
        subject = self.subject_cb.currentText().strip()
        grade = self.grade_cb.currentText().strip()

        if not subject or not grade:
            self.topic_cb.clear()
            self.type_cb.clear()
            return

        rows = self.db.execute_query("""
            SELECT name FROM exercise_tree 
            WHERE level='Chủ đề' AND parent_id IN (
                SELECT id FROM exercise_tree 
                WHERE name=? AND level='Lớp' AND parent_id IN (
                    SELECT id FROM exercise_tree WHERE name=? AND level='Môn'
                )
            )
        """, (grade, subject), fetch="all") or []

        self.topic_cb.clear()
        self.topic_cb.addItem("")
        for r in rows:
            self.topic_cb.addItem(r["name"])

    def load_available_types(self):
        """Load danh sách dạng"""
        topic = self.topic_cb.currentText().strip()
        if not topic:
            self.type_cb.clear()
            return

        rows = self.db.execute_query("""
            SELECT name FROM exercise_tree
            WHERE level='Dạng' AND parent_id IN (
                SELECT id FROM exercise_tree WHERE level='Chủ đề' AND name=?
            )
        """, (topic,), fetch="all") or []

        self.type_cb.clear()
        self.type_cb.addItem("")
        for r in rows:
            self.type_cb.addItem(r["name"])

    # ========== NHÓM 8: IMPORT/EXPORT ========== #
    def import_from_word(self):
        """Import từ Word"""
        try:
            from docx import Document
        except ImportError:
            QtWidgets.QMessageBox.critical(self, "Lỗi", "Cần cài đặt python-docx: pip install python-docx")
            return

        file_path, _ = QtWidgets.QFileDialog.getOpenFileName(
            self, "Chọn file Word", "", "Word files (*.docx)")
        if not file_path:
            return

        tree_id = self._current_tree_id()
        if not tree_id:
            QtWidgets.QMessageBox.warning(self, "Thiếu thư mục", "Vui lòng chọn nơi lưu câu hỏi.")
            return

        try:
            doc = Document(file_path)
            count = 0

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Simple import - just save as content
                    self.db.execute_query(
                        "INSERT INTO question_bank(content_text, tree_id) VALUES (?,?)",
                        (text, tree_id)
                    )
                    count += 1

            # Reload
            rows = self.db.execute_query("SELECT * FROM question_bank WHERE tree_id=?", (tree_id,), fetch="all") or []
            self._load_question_rows(rows)

            QtWidgets.QMessageBox.information(self, "Thành công", f"Đã import {count} câu hỏi.")

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể import: {e}")

    def export_to_word(self):
        """Export ra Word"""
        tree_id = self._current_tree_id()
        if not tree_id:
            QtWidgets.QMessageBox.warning(self, "Chưa chọn thư mục", "Vui lòng chọn thư mục để xuất.")
            return

        file_path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Lưu file Word", "", "Word files (*.docx)")
        if not file_path:
            return

        try:
            from docx import Document

            doc = Document()
            doc.add_heading('NGÂN HÀNG CÂU HỎI', 0)

            # Thông tin đường dẫn
            path_info = self.get_tree_path(tree_id)
            if path_info:
                path_text = " > ".join([p["name"] for p in path_info])
                doc.add_paragraph(f"Đường dẫn: {path_text}")

            # Lấy câu hỏi
            rows = self.db.execute_query("SELECT * FROM question_bank WHERE tree_id=?", (tree_id,), fetch="all") or []

            for i, row in enumerate(rows, 1):
                doc.add_paragraph(f"Câu {i}: {row.get('content_text', '')}", style='Heading 3')

                if row.get('correct'):
                    doc.add_paragraph(f"Đáp án: {row['correct']}")

                doc.add_paragraph("")  # Dòng trống

            doc.save(file_path)
            QtWidgets.QMessageBox.information(self, "Thành công", f"Đã xuất {len(rows)} câu hỏi ra file Word.")

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể xuất file: {e}")

    def export_to_pdf(self):
        """Export ra PDF"""
        tree_id = self._current_tree_id()
        if not tree_id:
            QtWidgets.QMessageBox.warning(self, "Chưa chọn thư mục", "Vui lòng chọn thư mục để xuất.")
            return

        file_path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Lưu file PDF", "", "PDF files (*.pdf)")
        if not file_path:
            return

        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch

            doc = SimpleDocTemplate(file_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []

            # Tiêu đề
            title = Paragraph("NGÂN HÀNG CÂU HỎI", styles['Title'])
            story.append(title)
            story.append(Spacer(1, 0.2 * inch))

            # Lấy câu hỏi
            rows = self.db.execute_query("SELECT * FROM question_bank WHERE tree_id=?", (tree_id,), fetch="all") or []

            for i, row in enumerate(rows, 1):
                question_para = Paragraph(f"<b>Câu {i}:</b> {row.get('content_text', '')}", styles['Normal'])
                story.append(question_para)

                if row.get('correct'):
                    answer_para = Paragraph(f"<b>Đáp án:</b> {row['correct']}", styles['Normal'])
                    story.append(answer_para)

                story.append(Spacer(1, 0.2 * inch))

            doc.build(story)
            QtWidgets.QMessageBox.information(self, "Thành công", f"Đã xuất {len(rows)} câu hỏi ra file PDF.")

        except ImportError:
            QtWidgets.QMessageBox.critical(self, "Lỗi", "Cần cài đặt reportlab: pip install reportlab")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể xuất PDF: {e}")

    # ========== NHÓM 9: TEMPLATE ========== #
    def show_template_dialog(self):
        """Hiển thị dialog template"""
        QtWidgets.QMessageBox.information(self, "Thông tin", "Chức năng template đang phát triển.")

    # ========== NHÓM 10: TAGS ========== #
    def add_new_tag(self):
        """Thêm tag mới"""
        if not hasattr(self, 'tags_edit') or not self.current_question_id:
            return

        tags_text = self.tags_edit.text().strip()
        if not tags_text:
            return

        tag_names = [tag.strip() for tag in tags_text.split(',') if tag.strip()]

        added_count = 0
        for tag_name in tag_names:
            try:
                self.db.execute_query(
                    "INSERT INTO question_tags(question_id, tag_name) VALUES (?,?)",
                    (self.current_question_id, tag_name)
                )
                added_count += 1
            except:
                pass

        if added_count > 0:
            self.tags_edit.clear()
            tree_id = self._current_tree_id()
            if tree_id:
                rows = self.db.execute_query("SELECT * FROM question_bank WHERE tree_id=?", (tree_id,),
                                             fetch="all") or []
                self._load_question_rows(rows)

    # ========== NHÓM 11: PREVIEW VÀ THỐNG KÊ ========== #
    def update_preview(self):
        """Cập nhật preview"""
        if not hasattr(self, 'preview_widget'):
            return

        content = self.content_text.toPlainText() if hasattr(self, 'content_text') else ""
        answer = self.answer_text.toPlainText() if hasattr(self, 'answer_text') else ""

        html = f"""
        <div style="font-family: Arial, sans-serif;">
            <h3 style="color: #2c3e50;">📝 Câu hỏi</h3>
            <p style="background: #f8f9fa; padding: 15px; border-left: 4px solid #007bff;">
                {content or '<em>Chưa có nội dung...</em>'}
            </p>
            <h4 style="color: #2c3e50;">✅ Đáp án</h4>
            <p style="background: #e8f5e9; padding: 15px;">
                {answer or '<em>Chưa có đáp án...</em>'}
            </p>
        </div>
        """
        self.preview_widget.setHtml(html)

    def update_statistics(self):
        """Cập nhật thống kê"""
        if not hasattr(self, 'stats_widget'):
            return

        try:
            # Tổng số câu hỏi
            total = self.db.execute_query("SELECT COUNT(*) as count FROM question_bank", fetch="one")["count"]

            # Thống kê theo mức độ
            level_stats = self.db.execute_query("""
                SELECT e.name, COUNT(q.id) as count 
                FROM question_bank q 
                JOIN exercise_tree e ON e.id = q.tree_id 
                WHERE e.level = 'Mức độ'
                GROUP BY e.name
                ORDER BY count DESC
            """, fetch="all") or []

            # Tạo HTML thống kê
            stats_html = f"""
            <div style="font-family: Arial, sans-serif;">
                <h2 style="color: #2c3e50;">📊 Thống kê</h2>
                <p><strong>Tổng số câu hỏi:</strong> {total}</p>

                <h3>Phân bố theo mức độ:</h3>
                <ul>
            """

            for stat in level_stats:
                percentage = (stat["count"] / total * 100) if total > 0 else 0
                stats_html += f"<li>{stat['name']}: {stat['count']} ({percentage:.1f}%)</li>"

            stats_html += "</ul></div>"

            self.stats_widget.setHtml(stats_html)

        except Exception as e:
            self.stats_widget.setHtml(f"<p style='color: red;'>Lỗi: {e}</p>")

    def refresh_all(self):
        """Làm mới toàn bộ"""
        self.refresh_tree()
        self.load_available_subjects()
        self.load_available_grades()
        self.on_tree_select()
        self.update_statistics()
class ImageViewer(QtWidgets.QWidget):
    """Widget hiển thị và xử lý ảnh"""

    def __init__(self):
        super().__init__()
        layout = QtWidgets.QVBoxLayout(self)

        # Scroll area để hiển thị ảnh lớn
        scroll = QtWidgets.QScrollArea()
        self.image_label = QtWidgets.QLabel()
        self.image_label.setAlignment(Qt.AlignCenter)
        self.image_label.setScaledContents(False)  # Giữ tỷ lệ ảnh
        scroll.setWidget(self.image_label)
        scroll.setWidgetResizable(True)
        layout.addWidget(scroll)

        # Toolbar zoom
        toolbar = QtWidgets.QToolBar()
        toolbar.addAction("🔍+").triggered.connect(lambda: self.zoom(1.2))
        toolbar.addAction("🔍-").triggered.connect(lambda: self.zoom(0.8))
        toolbar.addAction("💯").triggered.connect(self.fit_to_window)
        toolbar.addAction("📋 Paste").triggered.connect(self.paste_image)
        layout.addWidget(toolbar)

        self.current_pixmap = None
        self.scale_factor = 1.0

        # Enable drop
        self.setAcceptDrops(True)

    def set_image(self, image):
        """Set QImage từ clipboard"""
        if isinstance(image, QtGui.QImage):
            self.current_pixmap = QtGui.QPixmap.fromImage(image)
            self.display_image()
            return True
        return False

    def set_pixmap(self, pixmap):
        """Set QPixmap"""
        if isinstance(pixmap, QtGui.QPixmap):
            self.current_pixmap = pixmap
            self.display_image()
            return True
        return False

    def display_image(self):
        if self.current_pixmap and not self.current_pixmap.isNull():
            # Scale ảnh nhưng giữ tỷ lệ
            scaled = self.current_pixmap.scaled(
                self.current_pixmap.size() * self.scale_factor,
                Qt.KeepAspectRatio,
                Qt.SmoothTransformation
            )
            self.image_label.setPixmap(scaled)

    def paste_image(self):
        """Paste ảnh trực tiếp trong viewer"""
        clipboard = QtWidgets.QApplication.clipboard()
        mime_data = clipboard.mimeData()

        if mime_data.hasImage():
            image = clipboard.image()
            if not image.isNull():
                self.set_image(image)
                print("✅ Đã paste ảnh vào viewer")
        else:
            print("❌ Không có ảnh trong clipboard")

    def dragEnterEvent(self, event):
        """Xử lý kéo thả file"""
        if event.mimeData().hasImage() or event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dropEvent(self, event):
        """Xử lý thả file"""
        mime_data = event.mimeData()

        if mime_data.hasImage():
            self.set_image(mime_data.imageData())
        elif mime_data.hasUrls():
            for url in mime_data.urls():
                file_path = url.toLocalFile()
                if file_path.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                    pixmap = QtGui.QPixmap(file_path)
                    if not pixmap.isNull():
                        self.set_pixmap(pixmap)
                        break

    # Thêm vào class ImageViewer hiện tại của bạn:

    def fit_to_window(self):
        """Reset zoom về 100%"""
        self.scale_factor = 1.0
        self.display_image()

    def zoom(self, factor):
        """Zoom ảnh"""
        if not self.current_pixmap:
            return

        self.scale_factor *= factor
        # Giới hạn zoom từ 10% đến 500%
        self.scale_factor = max(0.1, min(5.0, self.scale_factor))
        self.display_image()

    def paste_image(self):
        """Paste ảnh từ clipboard"""
        clipboard = QtWidgets.QApplication.clipboard()
        if clipboard.mimeData().hasImage():
            image = clipboard.image()
            if not image.isNull():
                self.set_image(image)

    def clear_image(self):
        """Xóa ảnh"""
        self.current_pixmap = None
        self.image_label.clear()
class PDFViewer(QtWidgets.QWidget):
    """Widget hiển thị PDF (placeholder - cần thêm thư viện PDF)"""

    def __init__(self):
        super().__init__()
        layout = QtWidgets.QVBoxLayout(self)

        self.info_label = QtWidgets.QLabel("📄 PDF Viewer")
        self.info_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.info_label)

        # Toolbar cho PDF
        toolbar = QtWidgets.QToolBar()
        self.page_spin = QtWidgets.QSpinBox()
        self.page_spin.setMinimum(1)
        toolbar.addWidget(QtWidgets.QLabel("Trang:"))
        toolbar.addWidget(self.page_spin)
        toolbar.addAction("⬅️").triggered.connect(self.prev_page)
        toolbar.addAction("➡️").triggered.connect(self.next_page)
        layout.addWidget(toolbar)

        self.pdf_path = None
        self.current_page = 1
        self.total_pages = 1

    def load_pdf(self, file_path):
        """Load PDF file"""
        self.pdf_path = file_path
        self.info_label.setText(f"📄 {os.path.basename(file_path)}")
        # TODO: Implement PDF rendering với PyMuPDF hoặc pdf2image

    def prev_page(self):
        if self.current_page > 1:
            self.current_page -= 1
            self.page_spin.setValue(self.current_page)

    def next_page(self):
        if self.current_page < self.total_pages:
            self.current_page += 1
            self.page_spin.setValue(self.current_page)
class LaTeXInputDialog(QtWidgets.QDialog):
    """Dialog nhập công thức LaTeX"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("∑ Nhập công thức LaTeX")
        self.resize(500, 300)

        layout = QtWidgets.QVBoxLayout(self)

        layout.addWidget(QtWidgets.QLabel("Nhập công thức LaTeX:"))

        self.latex_edit = QtWidgets.QTextEdit()
        self.latex_edit.setPlaceholderText(r"\frac{a}{b} hoặc \int_{0}^{1} x^2 dx")
        layout.addWidget(self.latex_edit)

        # Preview (nếu có thư viện render LaTeX)
        self.preview_label = QtWidgets.QLabel("Preview sẽ hiển thị ở đây")
        self.preview_label.setAlignment(Qt.AlignCenter)
        self.preview_label.setStyleSheet("border: 1px solid #ccc; padding: 10px;")
        layout.addWidget(self.preview_label)

        # Buttons
        button_layout = QtWidgets.QHBoxLayout()
        ok_btn = QtWidgets.QPushButton("OK")
        ok_btn.clicked.connect(self.accept)
        cancel_btn = QtWidgets.QPushButton("Hủy")
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(cancel_btn)
        button_layout.addWidget(ok_btn)
        layout.addLayout(button_layout)

    def get_latex(self):
        return self.latex_edit.toPlainText()
class QuestionEditDialog(QtWidgets.QDialog):
    def __init__(self, db_manager, tree_id=None, question_id=None, parent=None):
        super().__init__(parent)
        self.db = db_manager
        self.tree_id = tree_id
        self.question_id = question_id
        self.content_type = "text"  # text, image, pdf, word
        self.answer_type = "text"
        self.content_data = None
        self.answer_data = None

        self.setup_ui()
        if question_id:
            self.load_question_data()

    def setup_ui(self):
        self.setWindowTitle("➕ Thêm câu hỏi mới" if not self.question_id else "✏️ Chỉnh sửa câu hỏi")
        self.resize(900, 700)

        layout = QtWidgets.QVBoxLayout(self)

        # Toolbar cho chọn loại nội dung
        toolbar = QtWidgets.QToolBar()
        toolbar.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)

        # Nhóm nút cho Câu hỏi
        toolbar.addWidget(QtWidgets.QLabel("📝 Thêm câu hỏi:"))

        text_action = toolbar.addAction("📝 Văn bản")
        text_action.triggered.connect(lambda: self.add_content("text"))

        image_action = toolbar.addAction("🖼️ Ảnh")
        image_action.triggered.connect(lambda: self.add_content("image"))

        pdf_action = toolbar.addAction("📄 PDF")
        pdf_action.triggered.connect(lambda: self.add_content("pdf"))

        word_action = toolbar.addAction("📘 Word")
        word_action.triggered.connect(lambda: self.add_content("word"))

        toolbar.addSeparator()

        latex_action = toolbar.addAction("∑ LaTeX")
        latex_action.triggered.connect(self.insert_latex)

        layout.addWidget(toolbar)

        # Splitter cho câu hỏi và đáp án
        splitter = QtWidgets.QSplitter(Qt.Vertical)

        # Phần 1: Câu hỏi
        question_group = QtWidgets.QGroupBox("📋 NỘI DUNG CÂU HỎI")
        question_layout = QtWidgets.QVBoxLayout(question_group)

        # Widget hiển thị nội dung (có thể là text, image, pdf viewer)
        self.content_widget = QtWidgets.QStackedWidget()

        # Text editor cho văn bản
        self.text_editor = QtWidgets.QTextEdit()
        self.text_editor.installEventFilter(self)
        self.text_editor.setAcceptRichText(True)
        self.text_editor.setPlaceholderText("Nhập nội dung câu hỏi hoặc dán ảnh (Ctrl+V)...")
        self.content_widget.addWidget(self.text_editor)

        # Image viewer
        self.image_viewer = ImageViewer()
        self.content_widget.addWidget(self.image_viewer)

        # PDF viewer
        self.pdf_viewer = PDFViewer()
        self.content_widget.addWidget(self.pdf_viewer)

        question_layout.addWidget(self.content_widget)
        splitter.addWidget(question_group)

        # Phần 2: Đáp án (có thể ẩn/hiện)
        self.answer_group = QtWidgets.QGroupBox("✅ ĐÁP ÁN")
        self.answer_group.setCheckable(True)
        self.answer_group.setChecked(True)
        answer_layout = QtWidgets.QVBoxLayout(self.answer_group)

        # Toolbar cho đáp án
        answer_toolbar = QtWidgets.QToolBar()
        answer_toolbar.addWidget(QtWidgets.QLabel("Thêm đáp án:"))

        ans_text_action = answer_toolbar.addAction("📝 Văn bản")
        ans_text_action.triggered.connect(lambda: self.add_answer("text"))

        ans_image_action = answer_toolbar.addAction("🖼️ Ảnh")
        ans_image_action.triggered.connect(lambda: self.add_answer("image"))

        answer_layout.addWidget(answer_toolbar)

        # Widget hiển thị đáp án
        self.answer_widget = QtWidgets.QStackedWidget()

        self.answer_text_editor = QtWidgets.QTextEdit()
        self.answer_text_editor.setMaximumHeight(150)
        self.answer_widget.addWidget(self.answer_text_editor)

        self.answer_image_viewer = ImageViewer()
        self.answer_widget.addWidget(self.answer_image_viewer)

        answer_layout.addWidget(self.answer_widget)
        splitter.addWidget(self.answer_group)

        layout.addWidget(splitter)

        # Phần 3: Tags
        tags_group = QtWidgets.QGroupBox("🏷️ TAGS")
        tags_layout = QtWidgets.QHBoxLayout(tags_group)

        self.tags_edit = QtWidgets.QLineEdit()
        self.tags_edit.setPlaceholderText("Nhập tags: môn, lớp, chủ đề, mức độ (phân cách bằng dấu phẩy)")
        tags_layout.addWidget(self.tags_edit)

        layout.addWidget(tags_group)

        # Buttons
        button_layout = QtWidgets.QHBoxLayout()

        save_btn = QtWidgets.QPushButton("💾 Lưu")
        save_btn.setStyleSheet("QPushButton { background: #28a745; color: white; padding: 10px 30px; }")
        save_btn.clicked.connect(self.save_question)

        cancel_btn = QtWidgets.QPushButton("❌ Hủy")
        cancel_btn.setStyleSheet("QPushButton { background: #6c757d; color: white; padding: 10px 30px; }")
        cancel_btn.clicked.connect(self.reject)

        button_layout.addStretch()
        button_layout.addWidget(cancel_btn)
        button_layout.addWidget(save_btn)

        layout.addLayout(button_layout)

        # Enable paste từ clipboard
        self.setup_clipboard()

    def setup_clipboard(self):
        """Xử lý paste ảnh từ clipboard"""
        shortcut = QShortcut(QKeySequence.Paste, self)  # ✅ cross-platform (Ctrl+V / Cmd+V)
        shortcut.activated.connect(self.paste_from_clipboard)

    def paste_from_clipboard(self):
        """Dán nội dung (ưu tiên ảnh) từ clipboard vào nội dung câu hỏi"""
        try:
            cb = QtWidgets.QApplication.clipboard()
            md = cb.mimeData()

            def qimage_from_clipboard():
                # 1) Ảnh thuần
                if md.hasImage():
                    img = cb.image()
                    if not img.isNull():
                        return img
                # 2) Fallback pixmap (Windows hay dùng)
                pm = cb.pixmap()
                if not pm.isNull():
                    return pm.toImage()
                # 3) Ảnh nhúng base64 trong HTML
                if md.hasHtml():
                    import re, base64
                    html = md.html()
                    m = re.search(r'data:image/(png|jpeg|jpg);base64,([A-Za-z0-9+/=]+)', html, re.I)
                    if m:
                        fmt = m.group(1).upper()
                        ba = QtCore.QByteArray.fromBase64(m.group(2).encode('ascii'))
                        img = QtGui.QImage()
                        if img.loadFromData(ba, fmt):
                            return img
                return None

            image = qimage_from_clipboard()

            if image:
                # Chuyển sang chế độ ảnh và hiển thị
                self.content_type = "image"
                self.content_widget.setCurrentWidget(self.image_viewer)
                self.image_viewer.set_image(image)

                # Lưu tạm bytes PNG để save xuống DB khi bấm Lưu
                ba = QtCore.QByteArray()
                buff = QtCore.QBuffer(ba)
                buff.open(QtCore.QIODevice.WriteOnly)
                image.save(buff, "PNG")
                self.content_data = bytes(ba)

                QtWidgets.QMessageBox.information(self, "Thành công", "Đã dán ảnh từ clipboard!")
                return

            # Nếu không có ảnh, dán text bình thường vào editor (nếu đang ở chế độ text)
            if md.hasText():
                text = md.text()
                if self.content_widget.currentWidget() != self.text_editor:
                    # Nếu đang ở màn ảnh, chuyển về text
                    self.add_content("text")
                cursor = self.text_editor.textCursor()
                cursor.insertText(text)
                return

            QtWidgets.QMessageBox.information(self, "Thông báo", "Clipboard không có nội dung phù hợp để dán.")
        except Exception as e:
            QtWidgets.QMessageBox.warning(self, "Lỗi", f"Không thể dán: {e}")
    def eventFilter(self, obj, event):
        """Bắt sự kiện keyboard cho text editor"""
        if obj == self.text_editor and event.type() == QtCore.QEvent.KeyPress:
            if event.matches(QtGui.QKeySequence.Paste):
                self.paste_from_clipboard()
                return True
        return super().eventFilter(obj, event)
    def add_content(self, content_type):
        """Thêm nội dung theo loại"""
        if content_type == "text":
            self.content_type = "text"
            self.content_widget.setCurrentWidget(self.text_editor)

        elif content_type == "image":
            file_path, _ = QtWidgets.QFileDialog.getOpenFileName(
                self, "Chọn ảnh", "", "Images (*.png *.jpg *.jpeg *.gif *.bmp)")
            if file_path:
                self.content_type = "image"
                pixmap = QPixmap(file_path)
                self.image_viewer.set_pixmap(pixmap)
                self.content_widget.setCurrentWidget(self.image_viewer)

        elif content_type == "pdf":
            file_path, _ = QtWidgets.QFileDialog.getOpenFileName(
                self, "Chọn PDF", "", "PDF Files (*.pdf)")
            if file_path:
                self.content_type = "pdf"
                self.pdf_viewer.load_pdf(file_path)
                self.content_widget.setCurrentWidget(self.pdf_viewer)

        elif content_type == "word":
            QtWidgets.QMessageBox.information(self, "Thông báo",
                                              "Chức năng import Word đang phát triển")

    def add_answer(self, answer_type):
        """Thêm đáp án theo loại"""
        if answer_type == "text":
            self.answer_type = "text"
            self.answer_widget.setCurrentWidget(self.answer_text_editor)

        elif answer_type == "image":
            file_path, _ = QtWidgets.QFileDialog.getOpenFileName(
                self, "Chọn ảnh đáp án", "", "Images (*.png *.jpg *.jpeg *.gif *.bmp)")
            if file_path:
                self.answer_type = "image"
                pixmap = QPixmap(file_path)
                self.answer_image_viewer.set_pixmap(pixmap)
                self.answer_widget.setCurrentWidget(self.answer_image_viewer)

    def insert_latex(self):
        """Chèn công thức LaTeX"""
        dialog = LaTeXInputDialog(self)
        if dialog.exec() == QtWidgets.QDialog.Accepted:
            latex_code = dialog.get_latex()
            if self.content_type == "text":
                cursor = self.text_editor.textCursor()
                cursor.insertText(f"$${latex_code}$$")

    def save_question(self):
        """Lưu câu hỏi vào database"""
        try:
            if not self.tree_id:
                QtWidgets.QMessageBox.warning(self, "Thiếu thư mục", "Vui lòng chọn vị trí lưu trong cây.")
                return

            # Chuẩn bị dữ liệu nội dung
            content_text = None
            content_data = None
            if self.content_type == "text":
                content_text = (self.text_editor.toPlainText() or "").strip()
                if not content_text:
                    QtWidgets.QMessageBox.warning(self, "Thiếu nội dung", "Nội dung câu hỏi không được để trống.")
                    return
            elif self.content_type == "image":
                # ưu tiên self.content_data (đã set khi dán); nếu chưa có, lấy từ viewer
                if self.content_data is None and self.image_viewer.current_pixmap:
                    ba = QtCore.QByteArray()
                    buff = QtCore.QBuffer(ba)
                    buff.open(QtCore.QIODevice.WriteOnly)
                    self.image_viewer.current_pixmap.toImage().save(buff, "PNG")
                    self.content_data = bytes(ba)
                content_data = self.content_data

            # Chuẩn bị dữ liệu đáp án (đơn giản: text hoặc ảnh)
            answer_text = None
            answer_data = None
            if self.answer_type == "text":
                answer_text = (self.answer_text_editor.toPlainText() or "").strip()
            elif self.answer_type == "image" and self.answer_image_viewer.current_pixmap:
                ba = QtCore.QByteArray()
                buff = QtCore.QBuffer(ba)
                buff.open(QtCore.QIODevice.WriteOnly)
                self.answer_image_viewer.current_pixmap.toImage().save(buff, "PNG")
                answer_data = bytes(ba)

            # Insert/Update
            if self.question_id:
                self.db.execute_query(
                    """UPDATE question_bank
                       SET content_text=?, content_type=?, content_data=?,
                           answer_type=?, answer_data=?,
                           correct=?,
                           tree_id=?, modified_date=CURRENT_TIMESTAMP
                       WHERE id=?""",
                    (content_text, self.content_type, content_data,
                     self.answer_type, answer_data,
                     answer_text,  # dùng cột 'correct' để lưu đáp án text ngắn
                     self.tree_id, self.question_id)
                )
            else:
                new_id = self.db.execute_query(
                    """INSERT INTO question_bank
                       (content_text, content_type, content_data,
                        answer_type, answer_data, correct, tree_id)
                       VALUES (?,?,?,?,?,?,?)""",
                    (content_text, self.content_type, content_data,
                     self.answer_type, answer_data, answer_text, self.tree_id)
                )
                self.question_id = new_id

            QtWidgets.QMessageBox.information(self, "Thành công", "Đã lưu câu hỏi.")
            self.accept()
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể lưu: {e}")


class TreeNodeDialog(QtWidgets.QDialog):
    def __init__(self, db_manager, mode="add", node_id=None, parent=None):
        super().__init__(parent)
        self.db = db_manager
        self.mode = mode  # "add" hoặc "edit"
        self.node_id = node_id
        self.parent_id = None

        self._setup_dialog()
        self._build_ui()
        self._load_data()

    def _setup_dialog(self):
        """Thiết lập dialog"""
        if self.mode == "add":
            self.setWindowTitle("➕ Thêm nhánh mới")
        else:
            self.setWindowTitle("✏️ Sửa nhánh")

        self.setModal(True)
        self.resize(450, 400)

    def _build_ui(self):
        """Xây dựng giao diện"""
        layout = QtWidgets.QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)

        # Header
        header = QtWidgets.QLabel()
        if self.mode == "add":
            header.setText("➕ Thêm nhánh mới vào cây thư mục")
        else:
            header.setText("✏️ Chỉnh sửa thông tin nhánh")

        header.setStyleSheet("""
            QLabel {
                font-size: 16px;
                font-weight: bold;
                color: #2E86AB;
                padding: 15px;
                background-color: #f8f9fa;
                border-radius: 8px;
                border: 1px solid #dee2e6;
            }
        """)
        layout.addWidget(header)

        # Form container
        form_container = QtWidgets.QWidget()
        form_container.setStyleSheet("""
            QWidget {
                background-color: white;
                border-radius: 8px;
                border: 1px solid #e1e5e9;
            }
        """)

        form_layout = QtWidgets.QFormLayout(form_container)
        form_layout.setSpacing(12)
        form_layout.setContentsMargins(20, 20, 20, 20)

        # Parent selection (chỉ hiện khi thêm)
        if self.mode == "add":
            self.parent_combo = QtWidgets.QComboBox()
            self.parent_combo.addItem("(Không có parent - Cấp gốc)", None)
            self._load_parent_options()

            parent_label = QtWidgets.QLabel("📁 Nhánh cha:")
            parent_label.setStyleSheet("font-weight: 500; color: #495057;")
            form_layout.addRow(parent_label, self.parent_combo)

        # Tên nhánh
        self.name_edit = QtWidgets.QLineEdit()
        self.name_edit.setPlaceholderText("Nhập tên nhánh...")

        name_label = QtWidgets.QLabel("📝 Tên nhánh:")
        name_label.setStyleSheet("font-weight: 500; color: #495057;")
        form_layout.addRow(name_label, self.name_edit)

        # Cấp độ
        self.level_combo = QtWidgets.QComboBox()
        self.level_combo.addItems(["Môn", "Lớp", "Chủ đề", "Dạng", "Mức độ"])

        level_label = QtWidgets.QLabel("📊 Cấp độ:")
        level_label.setStyleSheet("font-weight: 500; color: #495057;")
        form_layout.addRow(level_label, self.level_combo)

        # Mô tả
        self.description_edit = QtWidgets.QTextEdit()
        self.description_edit.setMaximumHeight(100)
        self.description_edit.setPlaceholderText("Nhập mô tả chi tiết...")

        desc_label = QtWidgets.QLabel("📄 Mô tả:")
        desc_label.setStyleSheet("font-weight: 500; color: #495057;")
        form_layout.addRow(desc_label, self.description_edit)

        # Style cho form inputs
        input_style = """
            QLineEdit, QComboBox, QTextEdit {
                padding: 10px;
                border: 2px solid #e1e5e9;
                border-radius: 6px;
                font-size: 13px;
                background-color: white;
            }
            QLineEdit:focus, QComboBox:focus, QTextEdit:focus {
                border-color: #2E86AB;
                outline: none;
                background-color: #f8fbff;
            }
        """

        self.name_edit.setStyleSheet(input_style)
        self.level_combo.setStyleSheet(input_style)
        self.description_edit.setStyleSheet(input_style)

        if hasattr(self, 'parent_combo'):
            self.parent_combo.setStyleSheet(input_style)

        layout.addWidget(form_container)

        # Buttons
        button_layout = QtWidgets.QHBoxLayout()
        button_layout.setContentsMargins(0, 10, 0, 0)

        cancel_btn = QtWidgets.QPushButton("❌ Hủy")
        cancel_btn.setFixedSize(100, 40)
        cancel_btn.clicked.connect(self.reject)
        cancel_btn.setStyleSheet("""
            QPushButton {
                background-color: #6c757d;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 13px;
            }
            QPushButton:hover {
                background-color: #5a6268;
            }
        """)

        # Save button
        if self.mode == "add":
            save_btn = QtWidgets.QPushButton("➕ Thêm")
            save_btn.setStyleSheet("""
                QPushButton {
                    background-color: #28a745;
                    color: white;
                    border: none;
                    padding: 10px 20px;
                    border-radius: 6px;
                    font-weight: 600;
                    font-size: 13px;
                }
                QPushButton:hover {
                    background-color: #218838;
                }
            """)
        else:
            save_btn = QtWidgets.QPushButton("💾 Lưu")
            save_btn.setStyleSheet("""
                QPushButton {
                    background-color: #007bff;
                    color: white;
                    border: none;
                    padding: 10px 20px;
                    border-radius: 6px;
                    font-weight: 600;
                    font-size: 13px;
                }
                QPushButton:hover {
                    background-color: #0056b3;
                }
            """)

        save_btn.setFixedSize(100, 40)
        save_btn.clicked.connect(self.accept)
        save_btn.setDefault(True)

        button_layout.addStretch()
        button_layout.addWidget(cancel_btn)
        button_layout.addSpacing(10)
        button_layout.addWidget(save_btn)

        layout.addWidget(QtWidgets.QWidget())
        layout.layout().addLayout(button_layout)

        # Focus vào name edit
        self.name_edit.setFocus()

    def _load_parent_options(self):
        """Load danh sách parent có thể chọn"""
        if self.mode != "add":
            return

        try:
            rows = self.db.execute_query(
                "SELECT id, name, level FROM exercise_tree ORDER BY level, name",
                fetch="all"
            ) or []

            for row in rows:
                # Nếu đang edit, không cho chọn chính nó làm parent
                if self.mode == "edit" and row["id"] == self.node_id:
                    continue

                display_text = f"{row['name']} ({row['level']})"
                self.parent_combo.addItem(display_text, row["id"])

        except Exception as e:
            print(f"Lỗi load parent options: {e}")

    def set_parent_id(self, parent_id):
        """Đặt parent được chọn"""
        self.parent_id = parent_id

        if self.mode == "add" and hasattr(self, 'parent_combo'):
            # Tìm và chọn parent trong combo
            for i in range(self.parent_combo.count()):
                if self.parent_combo.itemData(i) == parent_id:
                    self.parent_combo.setCurrentIndex(i)
                    break

    def _load_data(self):
        """Load dữ liệu nếu đang edit"""
        if self.mode != "edit" or not self.node_id:
            return

        try:
            row = self.db.execute_query(
                "SELECT name, level, description FROM exercise_tree WHERE id = ?",
                (self.node_id,), fetch="one"
            )

            if row:
                self.name_edit.setText(row["name"] or "")

                # Set level
                level = row["level"] or "Môn"
                index = self.level_combo.findText(level)
                if index >= 0:
                    self.level_combo.setCurrentIndex(index)

                # Description
                description = row.get('description', '') or ''
                self.description_edit.setPlainText(description)

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Không thể tải dữ liệu: {e}")

    def accept(self):
        """Xử lý khi người dùng nhấn Save/Add"""
        if not self._validate_input():
            return

        name = self.name_edit.text().strip()
        level = self.level_combo.currentText()
        description = self.description_edit.toPlainText().strip()

        try:
            if self.mode == "add":
                # Thêm node mới
                parent_id = None
                if hasattr(self, 'parent_combo'):
                    parent_id = self.parent_combo.currentData()
                elif self.parent_id:
                    parent_id = self.parent_id

                self.db.execute_query(
                    "INSERT INTO exercise_tree (parent_id, name, level, description) VALUES (?, ?, ?, ?)",
                    (parent_id, name, level, description)
                )

                QtWidgets.QMessageBox.information(
                    self, "Thành công",
                    f"Đã thêm nhánh '{name}' thành công!"
                )

            else:
                # Cập nhật node
                self.db.execute_query(
                    "UPDATE exercise_tree SET name = ?, level = ?, description = ? WHERE id = ?",
                    (name, level, description, self.node_id)
                )

                QtWidgets.QMessageBox.information(
                    self, "Thành công",
                    f"Đã cập nhật nhánh '{name}' thành công!"
                )

            super().accept()

        except Exception as e:
            QtWidgets.QMessageBox.critical(
                self, "Lỗi database",
                f"Không thể lưu dữ liệu:\n{str(e)}"
            )

    def _validate_input(self):
        """Validate dữ liệu đầu vào"""
        name = self.name_edit.text().strip()

        if not name:
            QtWidgets.QMessageBox.warning(
                self, "Lỗi",
                "Tên nhánh không được để trống!"
            )
            self.name_edit.setFocus()
            return False

        if len(name) > 100:
            QtWidgets.QMessageBox.warning(
                self, "Lỗi",
                "Tên nhánh không được quá 100 ký tự!"
            )
            self.name_edit.setFocus()
            return False

        return True
# ========== CLASS QUẢN LÝ CÂY NÂNG CAO ========== #
class TreeManagerDialog(QtWidgets.QDialog):
    def __init__(self, db_manager, parent=None):
        super().__init__(parent)
        self.db = db_manager
        self.setWindowTitle("⚙️ Quản lý cây thư mục")
        self.setModal(True)
        self.resize(800, 600)
        self._build_ui()
        self._load_tree_data()

    def _build_ui(self):
        """Xây dựng giao diện"""
        layout = QtWidgets.QVBoxLayout(self)

        # Toolbar
        toolbar = QtWidgets.QToolBar()
        toolbar.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)

        add_action = toolbar.addAction("➕ Thêm nhánh")
        add_action.triggered.connect(self._add_node)

        edit_action = toolbar.addAction("✏️ Sửa nhánh")
        edit_action.triggered.connect(self._edit_node)

        delete_action = toolbar.addAction("🗑️ Xóa nhánh")
        delete_action.triggered.connect(self._delete_node)

        toolbar.addSeparator()

        export_action = toolbar.addAction("📤 Xuất cấu trúc")
        export_action.triggered.connect(self._export_structure)

        layout.addWidget(toolbar)

        # Tree view
        self.tree_table = QtWidgets.QTreeWidget()
        self.tree_table.setHeaderLabels(["Tên", "Cấp độ", "Số câu hỏi", "Mô tả"])
        self.tree_table.setColumnWidth(0, 250)
        self.tree_table.setColumnWidth(1, 100)
        self.tree_table.setColumnWidth(2, 100)

        layout.addWidget(self.tree_table)

        # Buttons
        button_layout = QtWidgets.QHBoxLayout()

        close_btn = QtWidgets.QPushButton("Đóng")
        close_btn.clicked.connect(self.accept)

        button_layout.addStretch()
        button_layout.addWidget(close_btn)

        layout.addLayout(button_layout)

    def _load_tree_data(self):
        """Load dữ liệu cây"""
        # Implementation tương tự refresh_tree nhưng hiển thị trong table
        pass

    def _add_node(self):
        """Thêm node mới"""
        dialog = TreeNodeDialog(self.db, mode="add", parent=self)
        if dialog.exec() == QtWidgets.QDialog.Accepted:
            self._load_tree_data()

    def _edit_node(self):
        """Sửa node được chọn"""
        # Implementation
        pass

    def _delete_node(self):
        """Xóa node được chọn"""
        # Implementation
        pass

    def _export_structure(self):
        """Xuất cấu trúc cây"""
        # Implementation
        pass